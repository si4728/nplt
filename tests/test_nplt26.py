import subprocess
import sys
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock, patch

import nplt26
from PIL import Image
from docx import Document


class Nplt26UtilityTests(unittest.TestCase):
    @staticmethod
    def make_response(content, content_type="", apparent_encoding=None):
        response = nplt26.requests.Response()
        response.status_code = 200
        response._content = content
        if content_type:
            response.headers["Content-Type"] = content_type
        response.encoding = nplt26.requests.utils.get_encoding_from_headers(
            response.headers
        )
        if apparent_encoding:
            response._content_consumed = True
            response.apparent_encoding = apparent_encoding
        return response

    def test_remove_right_number_removes_trailing_digits(self):
        self.assertEqual(nplt26.remove_right_number("report123"), "report")

    def test_remove_right_number_handles_empty_string(self):
        self.assertEqual(nplt26.remove_right_number(""), "")

    def test_skip_costsaving_mode_accepts_lowercase_content_type(self):
        headers = {"content-type": "application/pdf"}
        self.assertTrue(nplt26.skip_costsavingMode(headers))

    def test_skip_costsaving_mode_detects_attachment(self):
        headers = {
            "Content-Type": "text/plain",
            "content-disposition": "attachment; filename=report.txt",
        }
        self.assertTrue(nplt26.skip_costsavingMode(headers))

    def test_skip_costsaving_mode_allows_html(self):
        headers = {"Content-Type": "text/html; charset=utf-8"}
        self.assertFalse(nplt26.skip_costsavingMode(headers))

    def test_year_list_contains_2022_and_current_year(self):
        self.assertIn("2022", nplt26.year_list)
        self.assertIn(str(nplt26.datetime.now().year), nplt26.year_list)

    def test_extract_extension_filter_list(self):
        self.assertEqual(
            nplt26.Extract_extion_fd("pdf, docx, hwp"),
            ["pdf", "docx", "hwp"],
        )

    def test_extract_clean_tags_keeps_selected_tags(self):
        html = '<div><a href="/">Home</a><span>Text</span></div>'
        self.assertEqual(
            nplt26.extract_clean_tags(html, ["div", "a"]),
            ["<div>", "<a>", "</a>", "</div>"],
        )

    def test_merge_dicts_as_tuples_uses_common_keys(self):
        result = nplt26.merge_dicts_as_tuples(
            {"h1": 2, "h2": 3}, {"h1": 1, "h3": 4}
        )
        self.assertEqual(result, {"h1": (2, 1)})

    def test_extract_text_removes_script_content(self):
        soup = nplt26.bs(
            "<html><script>ignore()</script><p>Hello world</p></html>",
            "html.parser",
        )
        self.assertEqual(nplt26.extract_text(soup), "Hello world")

    def test_parse_html_response_uses_declared_utf8_charset(self):
        response = self.make_response(
            "<title>온브랜딩</title>".encode("utf-8"),
            "text/html; charset=utf-8",
        )
        text, soup = nplt26.parse_html_response(response)
        self.assertIn("온브랜딩", text)
        self.assertEqual(soup.title.string, "온브랜딩")

    def test_parse_html_response_uses_declared_cp949_charset(self):
        response = self.make_response(
            "<title>한글 페이지</title>".encode("cp949"),
            "text/html; charset=cp949",
        )
        text, soup = nplt26.parse_html_response(response)
        self.assertIn("한글 페이지", text)
        self.assertEqual(soup.title.string, "한글 페이지")

    def test_parse_html_response_uses_apparent_encoding_without_charset(self):
        response = Mock()
        response.headers = {"Content-Type": "text/html"}
        response.apparent_encoding = "cp949"
        response.encoding = None
        response.text = "<title>감지된 한글</title>"
        text, soup = nplt26.parse_html_response(response)
        self.assertEqual(response.encoding, "cp949")
        self.assertIn("감지된 한글", text)
        self.assertEqual(soup.title.string, "감지된 한글")

    def test_small_and_positive_ignores_negative_value(self):
        self.assertEqual(nplt26.smallANDpositive(-1, 7), 7)

    def test_smallest_and_positive_returns_smallest_positive(self):
        self.assertEqual(nplt26.smallestANDpositive(-1, 8, 3, 0), 3)

    def test_smallest_and_positive_returns_minus_one_without_positive(self):
        self.assertEqual(nplt26.smallestANDpositive(-1, 0, -3), -1)

    def test_should_skip_href_rejects_javascript_this_reference(self):
        self.assertTrue(nplt26.should_skip_href("this.openPage()"))

    def test_get_rgb_space_handles_missing_file(self):
        self.assertEqual(
            nplt26.get_rgb_space("missing-image-for-test.png"),
            [0, 0, 0, 0, 0],
        )

    def test_get_rgb_space_counts_dominant_channels(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            image_path = Path(temp_dir) / "rgb.png"
            image = Image.new("RGB", (4, 1))
            image.putdata(
                [
                    (255, 0, 0),
                    (0, 255, 0),
                    (0, 0, 255),
                    (128, 128, 128),
                ]
            )
            image.save(image_path)
            self.assertEqual(
                nplt26.get_rgb_space(image_path),
                [1, 1, 1, 1, 4],
            )

    def test_make_pi_chart_uses_temp_directory(self):
        chart_path = nplt26.make_piChart(
            ["25 used", "75 unused"],
            "Test chart",
            "test_nplt26_chart.png",
        )
        self.assertEqual(
            Path(chart_path),
            nplt26.TEMP_DIR / "test_nplt26_chart.png",
        )
        self.assertTrue(Path(chart_path).is_file())
        Path(chart_path).unlink()

    def test_color_analysis_charts_are_included_in_report(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            red_path = temp_path / "red.png"
            blue_path = temp_path / "blue.png"
            red_image = Image.new("RGB", (8, 8), "red")
            blue_image = Image.new("RGB", (8, 8), "blue")
            red_image.putpixel((0, 0), (255, 128, 0))
            blue_image.putpixel((0, 0), (0, 128, 255))
            red_image.save(red_path)
            blue_image.save(blue_path)

            with (
                patch.object(nplt26, "IMAGE_DIR", temp_path),
                patch.object(nplt26, "DEFAULT_REPORT_PATH", str(temp_path) + "/"),
                patch.object(nplt26, "baseUrl", "example.com"),
                patch.object(nplt26, "report_list", []),
            ):
                bar_path = nplt26.make_image_analysis_bar(
                    [red_path, blue_path]
                )
                palette_path = nplt26.make_dominant_color_chart(
                    [red_path, blue_path]
                )
                nplt26.progress_make(2, "Website Color Analysis", "")
                nplt26.progress_make(5, palette_path, "")
                nplt26.progress_make(5, bar_path, "")
                nplt26.report_write("color-report.docx", "unused.png")

            document = Document(temp_path / "color-report.docx")
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Website Color Analysis", text)
            self.assertEqual(len(document.inline_shapes), 2)

    def test_extract_font_colors_normalizes_css_values(self):
        css = """
        body { color: #fff; background-color: #000; }
        h1 { color: rgb(255, 0, 0); }
        p { color: navy !important; }
        a { color: var(--link-color); }
        """
        self.assertEqual(
            nplt26.extract_font_colors(css),
            {"#FFFFFF": 1, "#FF0000": 1, "#000080": 1},
        )

    def test_parsing_fontlist_handles_css_shorthand_and_face(self):
        source = """
        body { font-family: "Pretendard", Arial, sans-serif !important; }
        h1 { font: italic 700 16px/1.5 "Noto Sans KR", sans-serif; }
        <font face="Nanum Gothic, serif">text</font>
        """
        self.assertEqual(
            nplt26.select_font(nplt26.parsing_fontlist(source)),
            [
                "Arial",
                "Nanum Gothic",
                "Noto Sans KR",
                "Pretendard",
                "sans-serif",
                "serif",
            ],
        )

    def test_report_images_fit_word_content_width(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "wide.png"
            Image.new("RGB", (2400, 800), "purple").save(image_path)
            with (
                patch.object(nplt26, "DEFAULT_REPORT_PATH", str(temp_path) + "/"),
                patch.object(nplt26, "baseUrl", "example.com"),
                patch.object(nplt26, "report_list", [[5, str(image_path)]]),
            ):
                nplt26.report_write("wide-report.docx", "unused.png")

            document = Document(temp_path / "wide-report.docx")
            section = document.sections[0]
            content_width = (
                section.page_width - section.left_margin - section.right_margin
            )
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertLessEqual(
                document.inline_shapes[0].width,
                content_width,
            )

    def test_report_write_handles_missing_image_entry(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            with (
                patch.object(nplt26, "DEFAULT_REPORT_PATH", str(temp_path) + "/"),
                patch.object(nplt26, "baseUrl", "example.com"),
                patch.object(nplt26, "report_list", [[5, "missing-chart.png"]]),
            ):
                nplt26.report_write("missing-image-report.docx", "unused.png")

            document = Document(temp_path / "missing-image-report.docx")
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("image unavailable", text)

    def test_tree_map_separates_internal_paths_and_external_domains(self):
        edges = [
            ("/", "/bbs/view.php?id=1"),
            ("/", "https://www.example.com/content/page.html"),
            ("/", "http://news.example.net/article/1"),
            ("/", "https://map.naver.com/p/search/test"),
        ]
        visited = {
            "https://example.com/bbs/view.php?id=1": 1,
            "http://www.example.com/bbs/view.php?id=2": 1,
            "https://example.com/content/page.html": 1,
        }

        paths, counts, external = nplt26.build_tree_map_data(
            edges,
            visited,
            "www.example.com",
            simple_mode=True,
        )

        self.assertIn("/bbs", paths)
        self.assertIn("/content", paths)
        self.assertNotIn("https://map.naver.com", paths)
        self.assertEqual(counts["/bbs"], 2)
        self.assertEqual(counts["/content"], 1)
        self.assertEqual(external, {"map.naver.com": 1, "news.example.net": 1})

    def test_font_color_chart_is_included_in_report(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            with (
                patch.object(nplt26, "IMAGE_DIR", temp_path),
                patch.object(nplt26, "DEFAULT_REPORT_PATH", str(temp_path) + "/"),
                patch.object(nplt26, "baseUrl", "example.com"),
                patch.object(nplt26, "report_list", []),
            ):
                chart_path = nplt26.make_font_color_chart(
                    {"#112233": 4, "#FFFFFF": 2}
                )
                nplt26.progress_make(2, "Website Color Analysis", "")
                nplt26.progress_make(5, chart_path, "")
                nplt26.report_write("font-color-report.docx", "unused.png")

            document = Document(temp_path / "font-color-report.docx")
            self.assertEqual(len(document.inline_shapes), 1)

    @patch("nplt26.HTTP_SESSION.get")
    def test_save_favicon_keeps_single_png_extension(self, get):
        image_bytes = BytesIO()
        Image.new("RGB", (16, 16), "red").save(image_bytes, format="PNG")
        response = Mock()
        response.content = image_bytes.getvalue()
        response.raise_for_status.return_value = None
        get.return_value = response

        with tempfile.TemporaryDirectory() as temp_dir:
            with patch.object(nplt26, "IMAGE_DIR", Path(temp_dir)):
                saved = Path(
                    nplt26.save_Favicon(
                        "https://example.com/icon/favicon-16x16.png?version=1"
                    )
                )

        self.assertEqual(saved.name, "favicon-16x16.png")
        self.assertNotIn(".png.png", saved.name)

    def test_report_write_uses_structured_favicon_data(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            favicon_path = temp_path / "favicon.png"
            Image.new("RGB", (16, 16), "blue").save(favicon_path)

            with (
                patch.object(nplt26, "DEFAULT_REPORT_PATH", str(temp_path) + "/"),
                patch.object(nplt26, "baseUrl", "example.com"),
                patch.object(nplt26, "report_list", []),
            ):
                nplt26.add_favicon_report(
                    "https://example.com/favicon.png",
                    favicon_path,
                )
                nplt26.report_write("favicon-report.docx", "unused.png")

            document = Document(temp_path / "favicon-report.docx")
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn(
                "Favicon information: https://example.com/favicon.png",
                text,
            )
            self.assertEqual(len(document.inline_shapes), 1)

    def test_report_write_repairs_legacy_double_png_path(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            favicon_path = temp_path / "favicon-16x16.png"
            Image.new("RGB", (16, 16), "green").save(favicon_path)
            legacy_path = temp_path / "favicon-16x16.png.png"

            legacy_entry = (
                "Favicon information: "
                "http://www.onbranding.co.kr/favicon-16x16.png,"
                f"{legacy_path}"
            )
            with (
                patch.object(nplt26, "DEFAULT_REPORT_PATH", str(temp_path) + "/"),
                patch.object(nplt26, "baseUrl", "www.onbranding.co.kr"),
                patch.object(nplt26, "report_list", [[7, legacy_entry]]),
            ):
                nplt26.report_write("legacy-favicon-report.docx", "unused.png")

            document = Document(temp_path / "legacy-favicon-report.docx")
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertNotIn("favicon image unavailable", text)

    @patch("nplt26.socket.getaddrinfo")
    def test_validate_public_url_blocks_private_address(self, getaddrinfo):
        getaddrinfo.return_value = [
            (2, 1, 6, "", ("127.0.0.1", 80)),
        ]
        with self.assertRaisesRegex(ValueError, "Non-public destination"):
            nplt26.validate_public_url("http://example.test")

    @patch("nplt26.socket.getaddrinfo")
    def test_validate_public_url_accepts_public_address(self, getaddrinfo):
        getaddrinfo.return_value = [
            (2, 1, 6, "", ("93.184.216.34", 80)),
        ]
        self.assertTrue(nplt26.validate_public_url("http://example.com"))

    @patch("nplt26.fetch_text_status")
    def test_configure_robots_blocks_disallowed_path(self, fetch_status):
        fetch_status.return_value = {
            "url": "https://example.com/robots.txt",
            "status": "available",
            "status_code": 200,
            "text": "User-agent: *\nDisallow: /private\n",
            "error": "",
        }
        nplt26.configure_robots("https://example.com", respect=True)
        self.assertFalse(nplt26.robots_allows("https://example.com/private/a"))
        self.assertTrue(nplt26.robots_allows("https://example.com/public"))

    def test_robots_status_message_reports_file_availability(self):
        self.assertEqual(
            nplt26.robots_status_message(
                {
                    "status": "available",
                    "status_code": 200,
                }
            ),
            "robots.txt is available. [HTTP 200]",
        )
        self.assertEqual(
            nplt26.robots_status_message({"status": "not_found"}),
            "robots.txt was not found. [HTTP 404]",
        )

    @patch("nplt26.HTTP_SESSION.get")
    @patch("nplt26.validate_public_url")
    def test_fetch_text_status_treats_empty_200_as_available(
        self, validate_url, get
    ):
        response = Mock()
        response.status_code = 200
        response.headers = {"Content-Type": "text/plain; charset=utf-8"}
        response.text = ""
        response.raise_for_status.return_value = None
        get.return_value = response

        result = nplt26.fetch_text_status("https://example.com/robots.txt")

        self.assertEqual(result["status"], "available")
        self.assertEqual(result["status_code"], 200)
        self.assertEqual(result["text"], "")

    @patch("nplt26.fetch_text")
    def test_sitemap_information_returns_text_or_none(self, fetch_text):
        fetch_text.return_value = "<urlset></urlset>"
        self.assertEqual(
            nplt26.getsitemapInformation("https://example.com"),
            "<urlset></urlset>",
        )
        fetch_text.return_value = None
        self.assertIsNone(
            nplt26.getsitemap2Information("https://example.com")
        )

    @patch("nplt26.mysql.connector.connect")
    def test_get_lastnumber_handles_connection_failure(self, connect):
        connect.side_effect = nplt26.mysql.connector.Error("connection failed")
        self.assertEqual(nplt26.get_lastnumber(), 0)

    @patch("nplt26.mysql.connector.connect")
    def test_db_cursor_commits_and_closes(self, connect):
        connection = Mock()
        cursor = Mock()
        connection.cursor.return_value = cursor
        connection.is_connected.return_value = True
        connect.return_value = connection

        with nplt26.get_db_cursor() as active_cursor:
            self.assertIs(active_cursor, cursor)

        connection.commit.assert_called_once()
        cursor.close.assert_called_once()
        connection.close.assert_called_once()

    def test_abnormal_url_allows_normal_query_urls(self):
        self.assertFalse(nplt26.abnormal_url("https://www.youtube.com/watch?v=abc"))
        self.assertFalse(nplt26.abnormal_url("https://twitter.com/intent/tweet?url=x"))
        self.assertTrue(nplt26.abnormal_url("https://example.com/login?next=home"))

    def test_sns_platform_detection_normalizes_hosts(self):
        self.assertEqual(
            nplt26.identify_sns_platform("https://m.facebook.com/acme"),
            "Facebook",
        )
        self.assertEqual(
            nplt26.identify_sns_platform("https://youtube.com/watch?v=1"),
            "YouTube",
        )
        self.assertEqual(
            nplt26.identify_sns_platform("https://youtu.be/abc"),
            "YouTube",
        )
        self.assertIsNone(nplt26.identify_sns_platform("https://fakefacebook.com"))

    def test_record_sns_link_tracks_types_pages_and_unique_urls(self):
        original_list = nplt26.list_sns
        original_details = nplt26.sns_details
        try:
            nplt26.list_sns = {}
            nplt26.sns_details = {}
            nplt26.record_sns_link(
                "https://www.youtube.com/watch?v=abc&utm_source=test",
                "https://example.com/page1",
            )
            nplt26.record_sns_link(
                "https://youtu.be/xyz",
                "https://example.com/page2",
            )

            self.assertEqual(nplt26.list_sns["YouTube"], 2)
            self.assertEqual(nplt26.sns_details["YouTube"]["types"]["content"], 2)
            self.assertNotIn("profile", nplt26.sns_details["YouTube"]["types"])
            self.assertEqual(len(nplt26.sns_details["YouTube"]["pages"]), 2)
        finally:
            nplt26.list_sns = original_list
            nplt26.sns_details = original_details

    def test_report_write_renders_sns_table_with_hyperlink(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            rows = [
                [
                    "YouTube",
                    "content:1",
                    "1",
                    "1",
                    "https://youtube.com/watch?v=abc",
                ]
            ]
            with (
                patch.object(nplt26, "DEFAULT_REPORT_PATH", str(temp_path) + "/"),
                patch.object(nplt26, "baseUrl", "example.com"),
                patch.object(nplt26, "report_list", []),
            ):
                nplt26.progress_make(2, "SNS link information", "")
                nplt26.progress_make_table(
                    ["Platform", "Type counts", "Pages", "Unique URLs", "Representative URL"],
                    rows,
                )
                nplt26.report_write("sns-report.docx", "unused.png")

            document = Document(temp_path / "sns-report.docx")
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(document.tables[0].cell(1, 0).text, "YouTube")
            self.assertIn(
                "hyperlink",
                document.tables[0].cell(1, 4).paragraphs[0]._p.xml,
            )

    def test_counting_esg_word_is_case_and_space_normalized(self):
        original_count = nplt26.esg_count
        try:
            nplt26.esg_count = {}
            found = nplt26.counting_esg_word(
                "Our Climate   Change policy improves 지속 가능성."
            )
            self.assertIn("CLIMATE CHANGE", found)
            self.assertTrue(
                "지속가능성" in found or "지속 가능성" in found
            )
            self.assertEqual(nplt26.esg_count["CLIMATE CHANGE"], 1)
        finally:
            nplt26.esg_count = original_count

    def test_import_has_no_console_output(self):
        workspace = Path(__file__).resolve().parents[1]
        result = subprocess.run(
            [sys.executable, "-c", "import nplt26"],
            cwd=workspace,
            capture_output=True,
            text=True,
            check=False,
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertEqual(result.stdout, "")
        self.assertEqual(result.stderr, "")


if __name__ == "__main__":
    unittest.main()
