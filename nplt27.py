# #23616  os: 47    https://www.kimm.re.kr/tour_visit/index/year/3944/month/02
# 23617  os: 47    https://www.kimm.re.kr/tour_visit/index/year/3944/month/03
# 23618  os: 47    https://www.kimm.re.kr/tour_visit/index/year/3944/month/04
# 23619  os: 47    https://www.kimm.re.kr/tour_visit/index/year/3944/month/05
# 23620  os: 47    https://www.kimm.re.kr/tour_visit/index/year/3944/month/06
## 2024.8.23 ing
# 351    os: 108   http://www.dtaq.re.kr/ko/notice/?mode=view&article_no=157534&board_wrapper=/ko/notice/press.jsp&pager.offset=310&board_no=26
# 19     os: 6     http://www.dtaq.re.kr/ko/state/info08.jsp
# skip
# 20     os: 7     http://www.dtaq.re.kr/ko/state/info08.jsp?mode=view&article_no=146881&board_wrapper=/ko/state/info08.jsp&pager.offset=0&board_no=6&default:category_id=170
#                  https://www.dtaq.re.kr/ko/state/info07.jsp?mode=view&article_no=140263&board_wrapper=%2Fko%2Fstate%2Finfo07.jsp&pager.offset=0&board_no=6&default:category_id=169
#21     os: 7     http://www.dtaq.re.kr/ko/state/info07.jsp#
# <a   href="?mode=view&amp;article_no=146881&amp;board_wrapper=%2Fko%2Fstate%2Finfo08.jsp&amp;pager.offset=0&amp;board_no=6&amp;default:category_id=170">
#############################################################################
# 프로그램명 nplt.exe (np_link_tracking)
# 주요기능:
#  nplt12.py add wordcloud function 
# 
# 사용법 nplt.exe -url "1)" -cost "2)" -tree "3)" -wl "4)"
# 1) 대상페이지 주소 
# 2) 페이지 추적시 이미지, 동영상, 문서 등에 대한 full download 테스트 실시 여부 
#    문자 1 또는 Yes 입력시 download 테스트 skip, 그오는 fulldownload 시작, 
#    기본 값은 Yes 설정되어 있습니다.
# 3) 정보구조도를 tree 형식으로 출력할때, 마지막 file level까지 출력(full mode)할지 
#    directory 단위까지 출력(simple mode)할지 결정
#    1 또는 Simple 입력시 simple 모드, 기본값은 full mode 입니다.
# 4) 페이지에 사용된 단어를 빈도수로 카운트하여 보여줄때 많은 순으로 몇개까지 출력할지
#    결정하는 값 숫자로입력 
#    기본값은 20 입니다.
# 5) nplt.exe -url https://5d29083470654.site123.me -cost Yes -wl 10
#    https://5d29083470654.site123.me 를 대상으로
#    이미지, 동영상은 download 를 제외하고 (cost=Yes)
#    전체 페이지에서 사용된 빈도수가 많은 단어를 10개 까지 출력
#    최종 출력물은 5d29083470654.site123.me.docx 에 저장됩니다.
#  
#     nplt18, nplt19 차이점 통신판매업 홈페이지 처리 루틴 추가 2022.2.6
#############################################################################
#-*-coding: utf-8
from nturl2path import url2pathname
import pdb #Debugging Usage
import argparse
from contextlib import contextmanager
import http
import ipaddress
import json
import os
import re
import socket
import sys
import tempfile
import time
import unicodedata
from datetime import datetime
from pathlib import Path
import urllib.error
import urllib.parse
import urllib.request
import urllib.robotparser
import html
import warnings
from datetime import datetime
from urllib.parse import urlparse
from urllib.parse import urljoin

import docx  # #pip install python-docx
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import colors as matplotlib_colors

#import MeCab  # ## 한글형태 분석  2019/9/21  https://joyhong.tistory.com/127
import networkx as nx  # #networkx 2.5
import requests
###import whois  # # python-whois 0.17.1        whois_function.py 로 교체 2021.11.28
from anytree import Node, RenderTree
from bs4 import BeautifulSoup as bs
from docx.shared import Inches, Pt, RGBColor
#from requests.packages.urllib3.exceptions import InsecureRequestWarning  #20240803
import warnings
import urllib3
#from requests.packages.urllib3.exceptions import InsecureRequestWarning

import mysql.connector
from mysql.connector import Error
from collections import defaultdict   ## merge_dicts(d1, d2)  20240721

from wordcloud import WordCloud     ##### 2020.7.8
import nltk    ####mMeCab 대체예정
#nltk.download() #first time
from nltk.tokenize import sent_tokenize
from nltk.tokenize import word_tokenize  ## nltk.download() on the GUI

import numpy as np
import cv2
from sklearn.cluster import KMeans
import nplt_forbiddenword
import nplt_whois2

#mMeCab = MeCab.Tagger()     ### 한글형태 분석  2019/9/21

sys.setrecursionlimit(100000)
#warnings.simplefilter('ignore',InsecureRequestWarning)
#Global Variable
timeout_duration = 100   ####2700 line 20240721 , request time out
REQUEST_TIMEOUT = (10, timeout_duration)
VERIFY_TLS = os.getenv("NPLT_VERIFY_TLS", "true").lower() not in {"0", "false", "no"}
ALLOW_PRIVATE_NETWORKS = (
    os.getenv("NPLT_ALLOW_PRIVATE_NETWORKS", "false").lower()
    in {"1", "true", "yes"}
)
if not VERIFY_TLS:
    warnings.simplefilter("ignore", urllib3.exceptions.InsecureRequestWarning)

BASE_DIR = Path(__file__).resolve().parent
REPORT_DIR = Path(os.getenv("NPLT_REPORT_PATH", BASE_DIR / "report"))
TEMP_DIR = Path(os.getenv("NPLT_TEMP_PATH", BASE_DIR / "temp"))
IMAGE_DIR = Path(
    os.getenv("NPLT_IMAGE_PATH", Path(tempfile.gettempdir()) / "nplt_image")
)

DEFAULT_PATH = str(BASE_DIR)
DEFAULT_REPORT_PATH = str(REPORT_DIR) + os.sep
DEFAULT_TEMP_PATH = str(TEMP_DIR) + os.sep
DEFAULT_IMAGE_PATH = str(IMAGE_DIR) + os.sep

# 디렉토리가 존재하지 않으면 생성
def ensure_output_directories():
    for directory in (REPORT_DIR, TEMP_DIR, IMAGE_DIR):
        directory.mkdir(parents=True, exist_ok=True)

DB_CONFIG = {
    "host": os.getenv("NPLT_DB_HOST", "localhost"),
    "port": int(os.getenv("NPLT_DB_PORT", "3306")),
    "database": os.getenv("NPLT_DB_NAME", "nplt"),
    "user": os.getenv("NPLT_DB_USER", "root"),
    "password": os.getenv("NPLT_DB_PASSWORD", "P@ssw0rd"),
}
host_address = DB_CONFIG["host"]
#host_address = 'localhost' #'203.250.81.142'
counter = 0#
costSavingMode = True
fullTreeMode = False
webPageColorAnalysis = True
keyWordList = 50
MinLinkLimit = 5
skipWordMode = True
caseSensitiveMode = False
http_https = "http"
Debug_mode = False
Search_string = ""
rptLang = 2 ##  1= KR, 2=EN
web_builder ="0"        # web builder tool def identify_website_builder
RestricedRestCycleCount = 99  #//동일 url 이나 parameter 가 다른 경우 ex) bbs 호출 제한
RestricedOnlineLinkCount = 49  # online salse site 인경우 동일 path 제한
Redirection_level = 0
forbiddenWordEnable = False

siteMap_flag = False

baseUrl2 =""
baseUrl  = ""

scanWebList = []
scanWebSet = set()

brokenLink = []
forbiddenList = []
imgCount = 0
imgAlt = 0
scriptCount = 0
styleCount = 0
cssCount = 0
cssCount2 = 0


frameCount = 0
iframeCount = 0
anchorCount = 0
anchorTCount = 0
FaviconUrl = "zz"
dbConnection = 0
dbConnection_id = 0
StopLine = 0

G = nx.Graph()
H = nx.Graph()

visitLinkDict = dict()
visitUrlDict = dict()
onlineLink = dict()
node_colors ="black"
FULLURL = False
print_toggle = 1
urltimelist = []
head_table = []
report_list = []
meta_list = []
img_list = []
img_list2 = []
#img_string = ""
#img_string2 = ""
lang_list = []
list_title = []                 ##check dupulicate title  []= [[url,title]] 2029.11.19
list_function = []
list_script =[]     ### [[name,path,type]]
list_sns = dict()
sns_details = {}
list_html5_tag=dict()
list_domain = []
list_plugin = []
list_search = []
list_flash = []     ### replace flashCount
list_para = set()
#list_font = set()
metaListset = set()
list_302 = set()
extion_count = dict()
word_count = dict()
incollect_path_list=set()
esg_count = dict()
html_tag=dict() ## check if remove 20241130
start_tag=dict()
end_tag=dict()

AddFavoriteSet=set()
AddFavoriteCount=0

keyword_path = "/"
list_record = []
list_page = []
list_header = []
list_Search_string = []
frameSetinformation = 0
list_img_analysis = set()
font_color_count = {}
input_url = "/"
yearUrlskip = True
year_list = [str(year) for year in range(2000, datetime.now().year + 1)]

tmp_cccc = set()
File_Download=[]

requests.headers = {
    'Accept-Language':'kr',
   ## 'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:57.0) Gecko/20100101 Firefox/57.0'
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}
HTTP_SESSION = requests.Session()
HTTP_SESSION.headers.update(requests.headers)
ROBOTS_PARSER = None
RESPECT_ROBOTS = True
ROBOTS_TEXT = None
ROBOTS_INFO = {
    "url": "",
    "status": "not_checked",
    "status_code": None,
    "text": None,
    "error": "",
}


def normalize_site_root(url):
    parsed = urlparse(url if "://" in url else f"http://{url}")
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError(f"Unsupported URL: {url}")
    port = f":{parsed.port}" if parsed.port else ""
    return f"{parsed.scheme}://{parsed.hostname}{port}"


def validate_public_url(url, allow_private=None):
    allow_private = ALLOW_PRIVATE_NETWORKS if allow_private is None else allow_private
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError(f"Only http/https URLs are supported: {url}")
    if parsed.username or parsed.password:
        raise ValueError("Credentials in URLs are not allowed")

    try:
        addresses = {
            info[4][0]
            for info in socket.getaddrinfo(
                parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80)
            )
        }
    except socket.gaierror as error:
        raise ValueError(f"Unable to resolve host {parsed.hostname}: {error}") from error

    if not allow_private:
        for address in addresses:
            ip = ipaddress.ip_address(address)
            if not ip.is_global:
                raise ValueError(f"Non-public destination is blocked: {address}")
    return True


def decode_response_text(response):
    content_type = response.headers.get("Content-Type", "")
    if "charset=" not in content_type.lower():
        response.encoding = response.apparent_encoding or "utf-8"
    return response.text


def parse_html_response(response):
    text = decode_response_text(response)
    return text, bs(text, "html.parser")


def fetch_text_status(url):
    try:
        validate_public_url(url)
        response = HTTP_SESSION.get(
            url,
            timeout=REQUEST_TIMEOUT,
            verify=VERIFY_TLS,
        )
        if response.status_code == 404:
            return {
                "url": url,
                "status": "not_found",
                "status_code": 404,
                "text": None,
                "error": "",
            }
        response.raise_for_status()
        return {
            "url": url,
            "status": "available",
            "status_code": response.status_code,
            "text": decode_response_text(response),
            "error": "",
        }
    except (requests.RequestException, ValueError) as error:
        return {
            "url": url,
            "status": "error",
            "status_code": None,
            "text": None,
            "error": str(error),
        }


def fetch_text(url):
    result = fetch_text_status(url)
    if result["status"] == "available":
        return result["text"]
    return None


def robots_status_message(info):
    status = info.get("status")
    if status == "available":
        status_code = info.get("status_code")
        return f"robots.txt is available. [HTTP {status_code}]"
    if status == "not_found":
        return "robots.txt was not found. [HTTP 404]"
    if status == "error":
        return f"robots.txt availability check failed: {info.get('error', '')}"
    return "robots.txt availability was not checked."


def configure_robots(start_url, respect=True):
    global ROBOTS_PARSER, ROBOTS_TEXT, ROBOTS_INFO, RESPECT_ROBOTS
    RESPECT_ROBOTS = respect
    root = normalize_site_root(start_url)
    robots_url = urljoin(root + "/", "robots.txt")
    ROBOTS_INFO = fetch_text_status(robots_url)
    ROBOTS_TEXT = ROBOTS_INFO.get("text")
    parser = urllib.robotparser.RobotFileParser()
    parser.set_url(robots_url)
    parser.parse((ROBOTS_TEXT or "").splitlines())
    ROBOTS_PARSER = parser
    return parser


def getrobotsInformation(url):
    root = normalize_site_root(url)
    result = fetch_text_status(urljoin(root + "/", "robots.txt"))
    if result["status"] == "available":
        return result["text"]
    return None


def robots_allows(url):
    if not RESPECT_ROBOTS or ROBOTS_PARSER is None:
        return True
    return ROBOTS_PARSER.can_fetch(requests.headers["User-Agent"], url)

online_score = 0
online_business = "통신판매"
online_word_list = ["예금주", "가격", "회원", "후기", "할인", "특가", "판매", "배송", "반품", "결재", "은행계좌"]


nplt_esg_word = ["ESG", "ENVIRONMENT", "CLIMATE CHANGE", "CARBON EMISSIONS", "POLLUTION", "RENEWABLE ENERGY", 
"WASTE MANAGEMENT", "WATER CONSERVATION", "DEFORESTATION", "BIODIVERSITY", "ECOSYSTEM SERVICES", "CIRCULAR ECONOMY", 
"SOCIAL", "HUMAN RIGHTS", "LABOR PRACTICES", "CHILD LABOR", "FORCED LABOR", "DIVERSITY AND INCLUSION", 
"HEALTH AND SAFETY", "COMMUNITY ENGAGEMENT", "PRODUCT SAFETY", "CONSUMER PROTECTION", "DATA PRIVACY", "GOVERNANCE", 
"CORPORATE GOVERNANCE","ANTI-CORRUPTION", "BOARD DIVERSITY", "EXECUTIVE COMPENSATION", "TRANSPARENCY", "ACCOUNTABILITY", 
"RISK MANAGEMENT", "COMPLIANCE", "STAKEHOLDER ENGAGEMENT", "ETHICS", "ESG RATING", "ESG INTEGRATION", 
"ESG INVESTING", "SUSTAINABLE INVESTING", "IMPACT INVESTING", "GREENWASHING", "DIVESTMENT", "SHAREHOLDER ENGAGEMENT", 
"PROXY VOTING", "ACTIVISM", "SUSTAINABILITY", "CORPORATE SOCIAL RESPONSIBILITY", "TRIPLE BOTTOM LINE", "SHARED VALUE", 
"STAKEHOLDER CAPITALISM", "GREEN ECONOMY", "LOW-CARBON ECONOMY", "CLIMATE RESILIENCE", "JUST TRANSITION", "ECO-FRIENDLY", 
"ECOSYSTEM", "ECOLOGY", "GREEN BUILDING", "SUSTAINABLE AGRICULTURE", "GREEN PRODUCTS", "CARBON NEUTRALITY", 
"RESOURCE EFFICIENCY", "ENVIRONMENTAL IMPACT ASSESSMENT", "OCEAN CONSERVATION", "RECYCLING", "AIR POLLUTION", "SUSTAINABLE CONSUMPTION", 
"SUSTAINABLE FUTURE", "GREEN TRANSPORTATION", "GREEN TECHNOLOGY", "CARBON FOOTPRINT", "SUSTAINABLE DEVELOPMENT", "ENVIRONMENTAL PROTECTION", 
"GLOBAL WARMING", "ENVIRONMENTAL RESPONSIBILITY", "GREEN LIFESTYLE", "SOCIAL RESPONSIBILITY", "LABOR CONDITIONS", "SOCIAL CONTRIBUTION", 
"SOCIAL INVESTMENT", "SOCIAL VALUE", "EDUCATIONAL OPPORTUNITIES", "HEALTH AND SAFETY", "SOCIAL ENTREPRENEURSHIP", "PHILANTHROPY", 
"SOCIAL COLLABORATION", "LABOR DISPUTES", "VOCATIONAL TRAINING", "CORPORATE GOVERNANCE", "ETHICAL MANAGEMENT", "BOARD OF DIRECTORS", 
"SHAREHOLDER VALUE", "RESPONSIBLE MANAGEMENT", "INTERNAL CONTROLS", "AUDIT COMMITTEE", "STAKEHOLDER TRANSPARENCY", "CODE OF ETHICS", 
"FINANCIAL REPORTING TRANSPARENCY", "GOVERNANCE ASSESSMENT", "ETHICAL INVESTING", "CORPORATE POLICIES", "BOARD INDEPENDENCE", "DISCLOSURE", 
"GOVERNANCE IMPROVEMENT", "STAKEHOLDER UNDERSTANDING", "INSIDER TRADING MONITORING", "GOVERNANCE TRANSPARENCY", "BOARD ACCOUNTABILITY", "GOVERNANCE EDUCATION", 
"VOTING RIGHTS", "LEGAL COMPLIANCE", "GOVERNANCE SUITABILITY", "CORPORATE ETHICS", "ACCOUNTING TRANSPARENCY", "STAKEHOLDER DIALOGUE", 
"STAKEHOLDER PARTICIPATION", "STAKEHOLDER CONSULTATION", "STAKEHOLDER COMMUNICATION", "STAKEHOLDER ENGAGEMENT PROCESS", "STAKEHOLDER REQUIREMENTS", "STAKEHOLDER FEEDBACK", 
"STAKEHOLDER ENGAGEMENT MODEL", "STAKEHOLDER MANAGEMENT", "STAKEHOLDER AWARENESS", "STAKEHOLDER STRATEGY", "STAKEHOLDER RELATIONS", "STAKEHOLDER EXPECTATIONS", 
"STAKEHOLDER TRUST", "STAKEHOLDER EMPATHY", "STAKEHOLDER INTERACTION", "STAKEHOLDER-CENTRIC", "STAKEHOLDER INCLUSIVENESS", "STAKEHOLDER SIGNIFICANCE", 
"STAKEHOLDER PRIORITIZATION", "SOCIAL ISSUES", "SOCIAL IMPACT", "SUSTAINABLE BUSINESS", "CORPORATE CITIZENSHIP", "ENVIRONMENTAL TRANSPARENCY", 
"ETHICAL CONDUCT", "STAKEHOLDER ENGAGEMENT METHODS", "GOVERNANCE ACCOUNTABILITY", "SOCIAL VALUE CREATION", "STAKEHOLDER NEEDS", "ENVIRONMENTAL SUSTAINABILITY", 
"CARBON NEUTRALIZATION", "GOVERNANCE EFFECTIVENESS", "SOCIAL IMPACT ASSESSMENT", "SOCIAL ORGANIZATION", "SUSTAINABLE MANAGEMENT", "SOCIAL LEADERSHIP", 
"CORPORATE POLITICAL CONTRIBUTIONS", "GOVERNANCE RESILIENCE", "CORPORATE AUDITS", "SOCIAL OPPORTUNITIES", "ETHICAL DECISION-MAKING", "SUSTAINABLE INNOVATION", 
"CORPORATE TRANSPARENCY", "GOVERNANCE ENHANCEMENT PROCESS", "ENVIRONMENTAL ADAPTABILITY", "GOVERNANCE LEADERSHIP", "ENVIRONMENTAL ACCOUNTABILITY", "SUSTAINABLE ECONOMY", 
"SOCIAL VALUE ASSESSMENT", "GOVERNANCE IMPROVEMENT MEASURES", "CORPORATE TRANSPARENCY ENHANCEMENT", "SUSTAINABLE BUSINESS MODEL", "EIA: ENVIRONMENTAL IMPACT ASSESSMENT", "SOCIAL IMPACT INVESTING", 
"CORPORATE PROFITS AND SOCIAL RESPONSIBILITY", "SOCIAL IMPACT ASSESSMENT INDICATORS", "CORPORATE ENVIRONMENTAL RESPONSIBILITY", "GOVERNANCE STRENGTHENING", "ENVIRONMENTAL CONSERVATION MEASURES", "SUSTAINABLE CONSUMER BEHAVIOR", 
"GOVERNANCE STREAMLINING", "SOCIAL VALUE NEGOTIATION", "ENVIRONMENTAL RESPONSIBILITY ENHANCEMENT", "SOCIAL REPORT", "SOCIAL SUSTAINABILITY", "GOVERNANCE ADJUSTMENT", 
"SOCIAL VALUE ENHANCEMENT", "SUSTAINABLE BUSINESS STRATEGY", "ENVIRONMENTAL CONSERVATION EFFORTS", "GOVERNANCE TRANSPARENCY ENHANCEMENT", "SOCIAL COEXISTENCE", "ENVIRONMENTAL FORECAST", 
"SOCIAL IMPACT MINIMIZATION", "GOVERNANCE INNOVATION", "ENVIRONMENTAL SUSTAINABILITY ASSESSMENT", "SUSTAINABLE BUSINESS MODELING", "SOCIAL ISSUES MANAGEMENT", "CORPORATE SOCIAL ROLE", 
"GOVERNANCE TRANSFORMATION", "ENVIRONMENTAL PROTECTION POLICIES", "SOCIAL IMPACT MEASUREMENT", "CORPORATE VALUATION", "SUSTAINABLE CONSUMER TRENDS", "ENVIRONMENTAL INNOVATION", 
"SOCIAL VALUE SHARING", "GOVERNANCE OVERHAUL", "SOCIAL IMPACT OPTIMIZATION", "SUSTAINABLE BUSINESS ADVANCEMENT", "ACTIVE ENVIRONMENTAL PARTICIPATION", "SOCIAL TRANSPARENCY", 
"CORPORATE ENVIRONMENTAL ACCOUNTABILITY", "GOVERNANCE INNOVATION MODEL", "SOCIAL IMPACT ANALYSIS", "SUSTAINABLE MANAGEMENT STRATEGY", "ENVIRONMENTAL ISSUE RESOLUTION", "SOCIAL IMPACT ASSESSMENT MODEL", 
"CORPORATE SOCIAL ACCOUNTABILITY", "GOVERNANCE EMPHASIS", "SOCIAL VALUE CONTRIBUTION", "CORPORATE ENVIRONMENTAL POLICIES", "SUSTAINABLE BUSINESS COMPETITIVENESS", "ENVIRONMENTAL RESPONSIBILITY POLICIES", 
"SOCIAL IMPACT ENHANCEMENT", "CORPORATE SOCIAL ACTIONS", "GOVERNANCE INNOVATION APPROACHES", "CO-CREATION OF SOCIAL VALUE", "CORPORATE SUSTAINABLE DEVELOPMENT", "SUSTAINABLE BUSINESS VALUE", 
"ENVIRONMENTAL IMPACT MINIMIZATION", "SOCIAL IMPACT ASSESSMENT TOOLS", "CORPORATE SOCIAL IMPACT", "GOVERNANCE INNOVATION STRATEGY", "INCREASE IN SOCIAL VALUE", "SUSTAINABLE BUSINESS PROCESSES", 
"ENVIRONMENTAL RESPONSIBILITY ACCOUNTING", "SOCIAL IMPACT ASSESSMENT GUIDELINES", "CORPORATE SUSTAINABILITY REPORT", "기후 변화", "탄소 배출", 
"오염", "재생 에너지", "폐기물 관리", "물 절약", "삼림 벌채", "생물 다양성", "지속가능", 
"생태계 서비스", "순환 경제", "인권", "노동 관행", "아동 노동", "강제 노동", 
"다양성 및 포용성", "건강 및 안전", "지역사회 참여", "제품 안전", "소비자 보호", "데이터 프라이버시", 
"지배구조", "기업 지배 구조", "반부패", "이사회 다양성", "경영진 보상", "투명성", 
"책임성", "위험 관리", "규정 준수", " 이해관계자 참여", "윤리", "ESG 평가", "이해관계자",
"ESG 통합", "ESG 투자", "지속가능한 투자", "임팩트 투자", "그린워싱", "투자 철회", 
"주주 참여", "위임 투표", "활동주의", "지속가능성", "기업 사회 책임", "3중 기준", 
"공유 가치", " 이해관계자 자본주의", "녹색 경제", "저탄소 경제", "기후 회복력", "공정한 전환", 
"친환경", "지속 가능성", "생태계", "에코 로지", "녹색 건축", "지속 가능한 농업", 
"친환경 제품", "탄소 중립", "자원 효율성", "환경 영향 평가", "해양 보호", "재활용", 
"대기 오염", "물 보존", "지속 가능한 소비", "지속 가능한 미래", "친환경 운송 수단", "친환경 기술", 
"탄소 발자국", "지속 가능한 개발", "환경 보호", "지구 온난화", "환경 책임", "친환경 라이프스타일", 
 "사회적 책임", "다양성과 포용성", "노동 조건", "지역 사회 기여", "사회적 공헌", 
"사회적 투자", "사회적 가치", "교육 기회", "건강과 안전", "사회적 창업", "기부", 
"사회적 협력", "노동 갈등", "직업 교육", "지배 구조", "윤리적 경영", "이사회", 
"주주 가치", "책임경영", "내부 통제", "감사위원회", "이해관계자 참여", "이해관계자 투명성", 
"부패 방지", "윤리 강령", "재무 보고서 투명성", "지배 구조 평가", "윤리적 투자", "사내 규정", 
"이사회 독립성", "대외 공개", "지배 구조 개선", "이해관계자 이해", "내부자 거래 감시", "지배 구조 투명성", 
"이사회 책임", "지배 구조 교육", "투표 권리", "법률적 준수", "지배 구조 적정성", "회사 윤 리", 
"회계 투명성", "이해관계자 대화", "이해관계자 의견 수렴", "이해관계자 소통", "이해관계자 참여 프로세스", " 이해관계자 요구사항", 
"이해관계자 피드백", "이해관계자 참여 모델", "이해관계자 관리", "이해관계자 인식", "이해관계자 전략", "이해관계자 관계", 
"이해관계자 기대", "이해관계자 신뢰", "이해관계자 공감", "이해관계자 간의 교류", "이해관계자 중심", "이해관계자 포괄성", 
"이해관계자 중요성", "이해관계자 우선순위", "사회적 이슈", "사회적 영향", "지속 가능한 비즈니스", "기업 시민권", 
"환경 투명성", "윤리적 행동", "이해관계자 참여 방법", "지배 구조 책임", "사회적 가치 창출", "이해관계자 니즈", 
"환경 지속 가능성", "기업 윤리", "탄소 중립화", "지배 구조 효과성", "사회적 영향 평가", "사회적 조직", 
"지속 가능한 경영", "사회적 리더십", "기업 정치 기부", "지배 구조 적응력", "기업 감사", "사회적 기회", 
"윤리적 의사결정", "지속 가능한 혁신", "기업 투명성", "지배 구조 향상 프로세스", "환경 적응력", "지배 구조 리더십", 
"환경 책임감", "지속 가능한 경제", "사회적 가치 평가", "환경 책임성", "지배 구조 개선책", "기업 투명성 제고", 
"지속 가능한 비즈니스 모델", "지배 구조 책임성", "환경영향평가", "사회적 영향 투자", "기업 이윤과 사회적 책임", "사회적 영향 평가 지표", 
"기업 환경 책임", "지배 구조 강화", "환경 보전 조치", "지속 가능한 소비 행동", "지배 구조 효율화", "사회적 가치 협상", 
"환경  책임 강화", "사회적 리포트", "사회적 지속 가능성", "지배 구조 조정", "사회적 가치 증진", "지속 가능한 비즈니스 전략", 
"환경 보호 노력", "지배 구조 투명성 강화", "사회적 상생", "환경 예측", "사회적 영향 최소화", "지배 구조 혁신", 
"환경 지속 가능성 평가", "지속 가능한 비즈니스 모델링", "사회적 이슈 관리", "기업 사회적 역할", "지배 구조 변화", "환경 보호 정책", 
"사회적 영향 측정", "기업 가치평가", "지속 가능한 소비 트렌드", "환경 혁신", "사회적 가치 공유", "지배 구조 정비", 
"사회적 영향 최적화", "지속 가능한 비즈니스 증진", "환경 적극 참여", "사회적 투명성", "기업의 환경 책임", "지배 구조 혁신 모델", 
"사회적 영향 분석", "기업의 환경적 책임감", "지속 가능한 경영 전략", "환경 문제 해결", "사회적 영향 평가 모델", "기업의 사회적 책임감", 
"지배 구조 강조", "사회적 가치 공헌", "기업의 환경 정책", "지속 가능한 비즈니스 경쟁력", "환경 책임 정책", "사회적 영향 증진", 
"기업 사회적 행동", "지배 구조 혁신 방안", "사회적 가치 공 동 창출", "기업의 지속 가능한 발전", "지속 가능한 비즈니스 가치", "환경 영향 최소화", 
"사회적 영향 평가 도구", "기업의 사회적 영향", "지배 구조 혁신 전략", "사회적 가치 증가", "기업의 환경적 책임", "지속 가능한 비즈니스 프로세스", 
"환경적 책임 회계", "사회적 영향 평가 지침", "기업의 지속 가능성 보고서"]

skip_ext_list = list(dict.fromkeys([
"zip", "jal", "pdf", "bin", "css", "class", "dat", "dot","hwp",
"psd", "mp4", "avi", "swf", "mpg", "mov", "vob", "ai",
"alz", "dll", "exe", "apk",
"wav", "mp3", "wma", "asf", "asx", "fl3", "m4p", "caf", "ra", "wax", "movie",
"jpg", "jpge", "png", "gif", "jpeg", "pic", "bmp", "tif", "ico", "psp", "eps", "fla",
"acl","ade","asd","cnv","crtx","doc","docm","docx","dotm","dotx","grv","h1q","iaf",
"maf","mam","maq","mar","mat","maw","mda","mdb","mde","mdt","mdw","mpd","mpp","mpt","mso",
"oab","obi","oft","olm","one","onepkg","ops","ost","pa","pip","pot","potm","potx","ppa","ppam","pps","ppsm","ppsx",
"ppt","pptm","pptx","prf","pst","pub","puz","rpmsg","sldm","sldx","slk","snp","svd","thmx","vdx","vsd","vsdx","vss","vst",
"vsx","vtx","wbk","wll","xar","xl","xla","xlam","xlb","xlc","xll","xlm","xls","xlsb","xlsm","xlsx","xlt","xltm","xltx","xlw","xsf","xslb","xsn"
]))

skip_word_list2 = ["HOME", "메인" , "하위분류", "바로가기", "TEL", "ADDRESS", "있습니다", "자료실", "로그인", "페이지", "않습니다", "없습니다",
"사이트", "발신자", "이메일", "글쓰기","아이디","글쓴이", "있는", "스크립트","있다", "것으로", "있습", "있음", "다음", "이전", "위해", "통해", "사항", "대한", "적응", "높은", "경우", "의한", "따른", 
"없습", "글쓴", "파악", "없다고", "요청", "대해", "있으며", "이유", "아래", "같은", "경우에", "귀하", "기타",
"위한", "다양한", "통하여", "위하여", "않으므로", "주십", "이미지", "갤러리", "검색어", "테이블", "타이틀", "글쓴이", "제목", "다음글", "이전글", "작성", "메인메뉴",
"마우스", "게시판", "스크롤", "리스트", "한다시요", "입력", "일치", "삭제할", "댓글", "버튼", "상단", "하단", "시작", "스크립트", "팝업",
"THAT", "MORE", "HAVE", "THIS", "BLOG", "CONTACT", "WITH", "WHAT", "VALLEY","SILICON","LOCAL", "YOUR", "FALSE", "ALTER", "PRAGMA",
"항목", "목록", "첨부", "영역", "메뉴", "돋음", "열지", "않음","이용동의", "이용동", "닫기","사용",
"일때", "즐겨찾기", "가기", "어떻소", "시켜놓은", "빠른", "가려지", "먼저", "나중", "새로운", 
"PARTNER", "THE", "IN", "FOR", "LTD", "TEL", "WHAT", "HOW", "WHEN", "FROM", "THAT", "THRE", "NOT", "BOOTSTRAP", "HTML5", "XHTML",
"YOU", "RESERVED", "EMAIL", "ALL", "GMAIL", "NET", "COM", "NEW", "AND", "RIGHTS","FAX", "GOTHIC", "STYLESHEET",
"SHARE", "WIDTH", "FONT", "CENTER", "BUTTON", "COLOR", "TEXT", "MARGIN", "SIZE", "BORDER", "ALIGN","TOP", "TOTAL", "CANCEL",
"LEFT", "HEIGHT", "NONE", "SOLID", "BOX", "TABLE", "ARIAL", "ABSOLUTE", "POSITION", "DISPLAY", "RELATIVE", "전역변수",
"HOVER", "TABLE","TABLEM", "BLOCK", "GNB", "BACKGROUND", "DECORATION", "URL", "ENDIF", "IMAGEHOVER", "CANONICAL",
"PHP","HTTP", "HTTPS", "HTML", "VIEW", "TITLE", "NBSP", "TYPE", "BGCOLOR","JAVASCRIPT", "CSCRIPT", "VAR", "NAVBAR", "DOCTYPE",
"ELSE", "IF", "THEN", "DATE", "COPYRIGHT", "UPPER", "DOWN", "PATHNAME", "ADDCLASS", "SUBMENU", "CSS", "CSS3", "ENDIF", "XML",
"IDXNO","LOGIN","MENU","NEWS","COOKIES","TERM", "MODOO", "WBODY", "FUNCTION", "SELOBJ", "SUBMENU1", "DOCTYPE", "XHTML", "HTML",
"IMGOBJ","JSHOST", "RETURN", "TAGNAME", "WINDOW", "WINRESULT", "PADDING", "ONCONTEXTMENU",
 "GETELEMENTBYID", "PUBLIC", "PREVIOUS", "COLLAPSE", "DEPTH3", "DEPTH2", "FOOTER",
 "NAVIGATION", "WEBMAIL", "SITEMAP", "REFRESH",
## winx.com
"NORMAL", "IMPORTANT", "INDEXABLE", "STRICT", "MOBILE", "ITALIC", "SELECTABLE", "UNAVAILABLE", "ISLOADABLE", 
"ENUMERABLE", "SUBSCRIPTIONS", "REGULAR", "DISABLED", "ACTIVE", "UPGRADE","CONST", "VISIBLE", "STATIC" , "OBJECT",
"USESSRSEO", "SECTIONTITLE", "BOOKINGS", "UPSCALEMETHODVALUE", "ANONYMOUS", "VERTICAL", "HORIZONTAL", "SYMBOL",
"SOCIALHOME", "BASIC","HCENTER","PROFILE","CALENDAR","TYPEERROR","FORMS","FULLSCREEN", "ERROR", "BOOKING",
"ISVISIBLE", "WEIGHT","ISVISIBLE","DYNAMICPAGELINK", "USERID","USERAGENT","RESIZEOBSERVER","URLTEMPLATE","REGEXP",
"UNDEFINED","HELVETICA","UPSCALE", "WRITABLE","PROMISE","CONFIGURABLE","OVERFLOW","ALIGNMENT", "UWPZU",
"INTERSECTIONOBSERVER","OWETE", "SAFARI", "WALLET","ACCOUNT","INITIAL", "CURRENTCNT","SCROLLY",
## winx.com end
"DECORATION", "RIGHT", "CORP", "COPYRIGHT", "TEST", "WHITE","HIGH","THEFORM","ARRAY","WINNAME", "THEURL", "IMAGE", "SIMSUN",
"SCRIPT", "ALERT","ONCLICK","RESULT","XSS", "IFRAME", "JQUERY", "PAGE", "MOUSE", "CDATA", "XML", "GETCOOKIE", "NEWWINDOW",
"LAYER",  "CLEARFIX", "SECTION", "START", "NUMBER", "MASTER", "HEADER", "FRONT", "SYSTEM",
"SLIDEBG1", "SCREEN", "SUBMIT",
"COUNT", "HIDDEN", "LABEL", "MODAL", "CONTENT", "ACTION", "EMOTICON", "TOGGLE"]


remove_suffix_list1 = ["이", "가", "을", "를", "에", "아", "야", "와", "과", "는", "의"] 
remove_suffix_list2 = ["께서", "에서", "에서", "한테", "으로", "이게", "으면", "이런", "통한", "따라"]
remove_suffix_list3 = ["입니다", "하세요", "주세요"]
remove_excetion1 = ["디스플레이"]

skip_prefix_list = ["IMAGE", "SUBMEMU", "MENU" ]


## https://blog.naver.com/PostView.nhn?blogId=zzangdol57&logNo=30172759123
                                                        

word_filter = [".", "=", "-", "_", "'", '"', "/", "+", "@", "[", "]", "<",">","SITEMAP", "MOUSE", "PAGE", "|"]

sns_domain_list =["x.com", "www.facebook.com", "www.whatsapp.com", "www.qq.com", "www.wechat.com", "qzone.qq.com",
"www.tumblr.com", "www.instagram.com", "twitter.com", "tieba.baidu.com", "www.skype.com", "www.viber.com",
"www.weibo.com", "line.me", "www.snapchat.com", "www.yy.com", "vk.com", "www.pinterest.co.kr",
"www.linkedin.com", "telegram.org", "www.reddit.com", "www.taringa.net", "foursquare.com", "www.renren.com",
"www.tagged.com", "badoo.com", "myspace.com", "www.stumbleupon.com", "the-dots.com", "www.skyrock.com",
"www.snapfish.com", "www.reverbnation.com", "www.flixster.com", "www.care2.com", "www.cafemom.com",
"www.ravelry.com", "go.nextdoor.com", "www.wayn.com", "www.cellufun.com", "www.youtube.com", "vine.co",
"www.classmates.com", "www.myheritage.co.kr", "kr.viadeo.com", "www.xing.com", "xanga.com", "www.livejournal.com",
"www.friendster.com", "www.funnyordie.com", "www.gaiaonline.com", "weheartit.com", "www.buzznet.com",
"www.deviantart.com", "www.flickr.com", "www.meetme.com", "www.meetup.com", "mixi.jp", "www.douban.com", "blog.naver.com", 
"www.spreely.com", "discordapp.com", "www.kakaocorp.com", "pf.kakao.com", "story.kakao.com", "youtu.be" ]

SNS_PLATFORMS = {
    "Facebook": {"facebook.com"},
    "Instagram": {"instagram.com"},
    "YouTube": {"youtube.com", "youtu.be"},
    "X/Twitter": {"x.com", "twitter.com"},
    "LinkedIn": {"linkedin.com"},
    "Naver Blog": {"blog.naver.com"},
    "Naver Band": {"band.us"},
    "Naver Cafe": {"cafe.naver.com"},
    "Naver Post": {"post.naver.com"},
    "Naver TV": {"tv.naver.com"},
    "Kakao Channel": {"channel.kakao.com", "pf.kakao.com"},
    "Kakao": {"kakaocorp.com", "story.kakao.com", "open.kakao.com"},
    "TikTok": {"tiktok.com"},
    "Threads": {"threads.net"},
    "WhatsApp": {"whatsapp.com"},
    "Messenger": {"messenger.com", "m.me"},
    "Pinterest": {"pinterest.co.kr", "pinterest.com"},
    "Telegram": {"telegram.org", "t.me"},
    "Reddit": {"reddit.com"},
    "Discord": {"discordapp.com", "discord.com"},
    "Line": {"line.me"},
    "Snapchat": {"snapchat.com"},
    "Twitch": {"twitch.tv"},
    "Vimeo": {"vimeo.com"},
    "Medium": {"medium.com"},
    "WeChat": {"wechat.com", "weixin.qq.com", "mp.weixin.qq.com"},
    "Weibo": {"weibo.com"},
    "Douyin": {"douyin.com"},
}

SNS_SHARE_PATHS = (
    "/share",
    "/sharer",
    "/intent/",
    "/sharing/",
    "/send",
)

SNS_EMBED_PATHS = (
    "/embed/",
    "/plugins/",
)

html5_elements = [
    'article', 'aside', 'audio', 'canvas', 'datalist', 'details', 'dialog',
    'figcaption', 'figure', 'footer', 'header', 'main', 'mark', 'meter',
    'nav', 'output', 'progress', 'section', 'summary', 'time', 'video',         
     'source', 'track', 'embed',  'bdi', 'wbr', 'template'
]

html_tag_list = ["html", "head", "body", "title", "p", "a",
                "h1", "h2", "h3", "h4", "h5", "h6", 
                "blockquote", "pre", "div", "nav", "style", "span", 
                "section", "article", "header", "footer",
                "aside", "main", "script", "font", 
                "table", "thead",  "tbody", "tfoot", "tr", "th", "td", "dl", "dd",
                "form", "label", "fieldset", "legend", "ul", "ol", "li",
"audio", "video", "figure", "figcaption"]

p3p_list = [ 
    ["compact-access", ["NOI", "ALL", "CAO", "IDC" , "OTI", "NON"]],
    ["compact-disputes", ["DSP"]],
    ["compact-remedies", ["COR", "MON", "LAW"]],
    ["compact-non-identifiable", ["NID"]],
    ["compact-purpose", ["CUR", "ADM", "DEV", "TAI", "PSA", "PSD", "IVA", "IVD", "CON", "HIS", "TEL", "OTP"]],
    ["creq", ["a", "i", "o"]],
    ["compact-recipient", ["OUR", "DEL", "SAM", "UNR", "PUB", "OTR"]],
    ["compact-retention", ["NOR", "STP", "LEG", "BUS", "IND"]],
    ["compact-category", ["PHY", "ONL", "UNI", "PUR", "FIN", "COM", "NAV", "INT", "DEM", "CNT", "STA", "POL", "HEA", "PRE", "LOC", "GOV", "OTC"]],
    ["compact-test", ["TST"]]
]

MessageList=[
                [0, "한글", "English",""],
                [1, "기본정보", "Genenal Information",""],
                [2, "날자: ", "Date: ","cdate_c"],
                [3, "사이트: ", "Site: ", "url_c"],
                [4, "IP: ", "IP: ", "ip_c"],
                [5, "Head 정보", "Header Information",""],
                [6, "페이지별 정보", "Page Information",""],
                [7, "1) 전체 스캔시간: ", "1) total scan time: ","scan_time"],
                [8, "2) 평균 스캔시간: ", "2) average scan time for a page of this site: ","average_scan_time"],
                [9, "3) 전체 페이지: ", "3) total pages in this site: ","pages"],
                [10,"4) BBS 형식 프로그램: ", "4) bbs style program in this site: ","bbs_pages"],
                [11,"5) 페이지 통계: ", "5) Max and Min page in this site: ",""],
                [12, "    - 제일 긴   페이지: ", "    - Max contents length: ","max_content_length"],
                [13, "    - 제일 짧은 페이지: ", "    - Min contents length: ","min_content_length"],
                [14, "    - 응답이 제일 늦은 페이지: ", "    - Max response time: ","max_response_time"],
                [15, "    - 응답이 제일 빠른 페이지: ", "    - Min response time: ","min_response_time"],
                [16, "오류링크", "Broken Links","count_broken_link"],
                [17, "모든 링크가 정상입니다",  "none Boren links",""],
                [18, "Meta Tag 분석", "Meta Tag Information",""],
                [19, "모바일 기기를 지원하는 Tag를 사용하고 있습니다.", "Mobile device friendly site.","mobile_tag"],
                [20, "모바일 기기를 지원하는 Tag 정보가 없습니다.", "Mobile device friendly tag does not exist.",""],
                [21, "Meta tag를 충분히 활용하지 않고 있습니다.", "Meta tag data is not enough.",""],
                [22, "Meta tag는 검색엔진이 웹페이지 분석시 매우 중요하게 활용하는 정보입니다.", "Meta tags are HTML elements that provide information about a web page for search engines and website visitors.",""],
                [23, "Meta tag[keyword, description, robots...]등의 정보를 재확인하여 개선하십시요", "Please update meta [keyword, description, robots..] tag on your websie.",""],
                [24, "정보", "Information",""],
                [25, "없습니다.", "Not exists.",""],
                [26, "Sitemap를 제공하나 XML 또는 HTML형식이 아닙니다.", "This site providing the Sitemap but it is not XML or HTML file.",""],
                [27, "이 사이트에는 sitemap.xml 또는 Sitemap.html이 없습니다.", "No Sitemap.html or Sitemap.xml is available on this site.",""],
                [28, "이 사이트에는 Sitemap.html이 존재합니다.", "Sitemap.html is available on this site.",""],
                [29, "이 사이트에는 Sitemap.xml이 존재 합니다.", "Sitemap.xml is available on this site.",""],
                [30, "사용된 이미지 태그 수: ", "count of image tag: ","count_images"],
                [31, "이미지 태그에 ALT 속성이 적용된 갯수: ",  "count of alt attribute in the image tag: ","count_alt_image"],
                [32, "이미지 태그에 ALT속성이 적용된 비율: ", "Utilization of alt attribute in image tag: ",""],
                [33, "사용된 이미지 파일 종류 ", "Image extention ","image_type"],
                [34, "H1, H2, H3테그를 활용하여 주요한 키워드를 강조 하십시요. 검색엔진은 이 부분을 중요시 여깁니다.", "Search engines use h1, h2, and h3 tags for SEO purposes.",""],
                [35, " 회 사용되었습니다.", " times used in this site","headline_tag"],
                [36, "H1, H2 또는 H3 태그 사용을 발견 할 수 없습니다.", "No h1, h2, or h3 tags were detected on this site.",""],
                [37, "h1,h2 또는 h3 태그를 사용하여 주요 키워드를 강조 하십시요.", "We recommend using h1 and h2 tags for each page's key keywords.",""],
                [38, "   사용된 Flash 파일수: ", "   Flash file count: ","count_flash"],
                [39, "웹사이트에 flash 파일 사용을 권장 하지 않습니다.", "We do not recommend using flash file on your website.",""],
                [40, "Flash 파일이 없습니다 (Good)", "Can not found Flash file in this (Good).",""],
                [41, "사용된 script tag 수: ", "count of script tag: ","count_script"],
                [42, "사용된 stylesheet tag 수: ", "count of stylesheet tag: ",""],
                [43, "사용된 stylesheet file 수: ", "count of stylesheet file: ",""],
                [44, "웹사이트에 사용된 키워드 순위",  " Frequency word list in website",""],
                [45, "웹사이트 노드 Graph", "Node graph",""],
                [46, "--- Cost 절약 모드에서 시행중(멀티미디어 파일에 대한 속도 점검은 생략합니다).",
                "--- run on cost saving mode (most of multimedia data was not include for calculated)",""],
                [47,"6) 페이지 링크 통계: ", "5) Most and least linked page in this site: ",""],
                [48, " - 가장많이 링크된 페이지", " - The most linked page on website","most_linked"],
                [49, " - 가장적게 링크된 페이지", " - The least linked page on website","least_linked"],
                [50,"각 페이지별 타이틀 속성", "Title informatio of each pages",""],
                [51,"참고) 페이지별 TITLE은 다르게 기술하는 것을 권합니다.", "Note) We recommend that describe the TITLE per page differently.",""],
                [52,"참고) 타 웹사이트 속도.", "Note) The speed of other Websites.",""],
                [53, " 외부 사이트 이미지 수: ", " The count of image of other website: ",""],
                [54, "사용된 Javascript file 수: ", "count of javascript file: ",""],
                [55, "Google 검색 결과", "the Result of google.com search" ,"google"],
                [56, "Frame 사용현황", "Information of Frmae tag",""],
                [57, " frame tag 사용수: ", "Count of frame tag: ","count_frame"],
                [58, " iframe tag 사용수: ", "Count of iframe tag: ","count_iframe"],
                [59, "Font 목록", "Font List",""],
                [60, " - 내부링크 페이지 수: ", " - The count of Inner-link: ","count_in_linked"],
                [61, " - 외부링크 페이지 수: ", " - The count of Outer-link: ","count_out_linked"],
                [62, "플러그인 정보", "PlugIn Information",""],
                [63, " - 플러그인 tag 사용횟수: ", " - The count of plugin tag: ","count_plugin"],
                [64, "SNS 사이트 연결 정보", "SNS link information",""],
                [65, " - SNS 연결정보가 없습니다", " - Can not found SNS link information",""],
                [66, "사용된 Anchor 태그 수: ", "count of anchor tag: ","count_anchor"],
                [67, "anchor 태그에 Text 값이 있는 갯수: ",  "count of Text value in the anchor tag: ","count_text_anchor"],
                [68, "anchor 태그에 text 적용 비율: ", "Utilization of Text attribute in anchor tag: ",""],
                [69, "start date time", "start date time","sdate_d"],
                [70, "end date time", "end date time", "edate_d"],
                [71, "개인정보보호정책(P3P): ", "P3P-Platform for Privacy Preferences Project: ", "p3p"],
                [72, "개인정보보호정책(P3P)이 없습니다.", "P3P-Platform for Privacy Preferences Project is not defined.", "p3px"],
                [73, " - 이사이트는 <frame>, <iframe>, <embed> or <object>내부에서 페이지를 호출할 수 있으므로 Clickjacking 공격이 가능함.", " - Sites can use this to avoid click-jacking attacks, by ensuring that their content is not embedded into other sites.","xo"],
                [74, "    - iframe은 보안에 취약합니다. 사용을 제한하십시요.", "    - iframes are vulnerable to security, We are recommended not to use it.", "im"],
                [75, "    - 평균 페이지 Scan 시간(nplt 인덱스): ", "    - Average Page Scan Time(nplt Index): ", "apst"],  ## standard
                [76, "    - 평균 페이지 (nplt 인덱스): ", "    - The Average Page(nplt Index): ", "ap"],  ## standard
                [77, "- Broken Link 갯수(nplt 인덱스): ", "- The Average number of Broken Link(nplt Index): ", "apbl"],  ## standard
                [78, "    - 평균 Frame 갯수(nplt 인덱스): ", "    - The Average number of Frame(nplt Index): ", "anf"],  ## standard
                [79, "    - 평균 Google link수(nplt 인덱스): ", "    - The Average number of link on Google(nplt Index): ", "ang"],  ## standard
                [80, "    - Mobile 태그 사용비율(nplt 인덱스): ", "    - Mobile tag usage rate in Website(nplt Index): ", "mtn"],  ## standard
                [81, "    - Image 내 ALT 태그 사용비율(nplt 인덱스): ", "    - ALT tag usage rate in Image(nplt Index): ", "ani"],  ## standard
                [82, "      - H1 tag 사용비율(nplt 인덱스): ", "    - H1 tag usage rate(nplt Index): ", "h1t"],  ## standard
                [83, "      - H2 tag 사용비율(nplt 인덱스): ", "      - H2 tag usage rate(nplt Index): ", "h2t"],  ## standard
                [84, "      - H3 tag 사용비율(nplt 인덱스): ", "      - H3 tag usage rate(nplt Index): ", "h3t"],  ## standard
                [85, "    - Flash 사용비율(nplt 인덱스): ", "    - Flash tag usage rate(nplt Index): ", "ftu"],  ## standard
                [86, "    - Domain 유지보수 주기가 2년이상이 아닙니다. 개선사항: ", "      - Domain Update term less than 2 years, Please change policy: ", "dut"],
                [87, "AddFavorite 명령어 수: ", "The Count of AddFavorite function: ", "AddFavorite"],
                [88, "Web 사이트 작성 도구: ", "Website builder tool: ", "wbt"],
                [89, "웹사이트에 사용된 ESG 키워드 순위",  " Frequency ESG word list in website",""],
                [90, "HTML5 사용",  "HTML5 tag used",""]
            ]


def Extract_extion_fd(str):
    return [part.strip() for part in str.split(",")]

def extract_clean_tags(html: str, tag_list: list) -> list:
    tag_pattern = r"</?([a-zA-Z][a-zA-Z0-9]*)[^>]*>"
    extracted_tags = []

    for match in re.finditer(tag_pattern, html):
        tag_name = match.group(1)  # 태그 이름만 추출
        full_tag = match.group(0)  # 전체 태그 (<tag> 또는 </tag>)
        
        if tag_name in tag_list:  # 지정된 태그 리스트에 포함된 경우만 추가
            if full_tag.startswith("</"):  # 종료 태그
                extracted_tags.append(f"</{tag_name}>")
            else:  # 시작 태그
                extracted_tags.append(f"<{tag_name}>")

    return extracted_tags

def merge_dicts_as_tuples(dict_a, dict_b):  
    common_keys = dict_a.keys() & dict_b.keys()
    merged_dict = {key: (dict_a[key], dict_b[key]) for key in common_keys}    
    return merged_dict


def extract_text(content):
    # 모든 <script> 태그 제거
    for script in content.find_all("script"):
        script.decompose()  # 태그 제거

    # 텍스트 추출 (공백으로 구분)
    text = content.get_text(separator=' ', strip=True)
    return text

def smallANDpositive(x,y):
    if x < 0:
        return y
    if y < 0:
        return x
    if x < y:
        return x
    else:
        return y

def smallestANDpositive_Old(a,b,c,d,e,f,g,h,i):
    f1 = smallANDpositive(a,b)
    f2 = smallANDpositive(c,d)
    f3 = smallANDpositive(e,f)
    f4 = smallANDpositive(g,h)
    f5 = smallANDpositive(f1,i)
    f23 = smallANDpositive(f2,f3)
    f45 = smallANDpositive(f4,f5)    
    return smallANDpositive(f23,f45)


def should_skip_href(href_string):
    #   href 문자열이 특정 조건을 만족하는지 검사하여 건너뛸지 여부를 반환
     # 'this.'가 포함되어 있으면 자바스크립트 함수 호출로 간주하고 건너뜀
    if "this." in href_string:
        return True

    # 문자열 길이가 3 미만인 경우 건너뜀
    if len(href_string) < 3:
        return True

    # href 바로 뒤에 '.'이 오는 경우 건너뜀 (비정상 링크일 가능성)
    if href_string[4:5] == ".":
        return True

    # href에 쉼표가 포함된 경우 건너뜀
    if "," in href_string:
        return True

    return False

def smallestANDpositive(*args):
    # 양수인 값들 중 가장 작은 값을 반환, 양수 값이 없으면 -1 반환
    positive_vals = [val for val in args if val > 0]
    return min(positive_vals) if positive_vals else -1

def extract_s_list(hstring):
    ss1 = "http://" 
    ss2 = "https://"
    fff = '"'
    zzz = "'"
    bbb = " "
    nnn = "&"
    mmm = ";"
    lll = "<"
    rrr = ")"
    qqq = ">"
    kkk = "["
    start =1
    end=len(hstring)
    s_list=[]

    while(1):        
        x = hstring.find(ss1, start, end)
        if x < 0:
            x = hstring.find(ss2, start, end)
        if x > 0:
            y = hstring.find(fff, x+1, end)
            z = hstring.find(zzz, x+1, end)
            u = hstring.find(bbb, x+1, end)
            n = hstring.find(nnn, x+1, end)
            m = hstring.find(mmm, x+1, end)
            l = hstring.find(lll, x+1, end)
            r = hstring.find(rrr, x+1, end)
            q = hstring.find(qqq, x+1, end)
            k = hstring.find(kkk, x+1, end)        
            y = smallestANDpositive(y,z,u,n,m,l,r,q,k)
        
        if x > 0 and y > 0:
            s_list.append([x,y])
        else:
            break
        start = x + 1
    return s_list

def get_p3p_label(p3pin):   
    p3p_dict = dict()
    for kword in p3pin:
        for L11, L12 in p3p_list:
            ckword = kword
            if len(kword)==4:
                ckword=kword[:-1]
            if ckword in L12:
                new_flag = 1
                for k1, v1 in p3p_dict.items():
                    if k1 == L11:
                        p3p_dict[k1] = v1 + ", " + kword
                        new_flag = 0
                if new_flag == 1:
                    p3p_dict[L11] = kword
    return p3p_dict          

def merge_dicts(d1, d2):
    merged = defaultdict(int)
    for key, value in d1.items():
        merged[key] += value
    for key, value in d2.items():
        merged[key] += value
    return dict(merged)

def check_html5_elements(soup):
    found_elements = {element: len(soup.find_all(element)) for element in html5_elements if soup.find_all(element)}
    return found_elements

def remove_right_number(w):
    if not w:
        return w
    w = w[:-1]
    while w and w[-1].isdigit():
        w = w[:-1]
    return w

def make_Search_string(ss):
    sl = []

    if ss.find(".txt") > 0:
        f = open(ss, "r")
        for keylist in f:
            keylist = keylist.replace("\n","")
            sl.append(keylist.upper())
            #print("....", keylist)
        f.close()
    else:
        sl.append(ss)        
    return sl
    
def convertToBinaryData(filename):
    # Convert digital data to binary format
    with open(filename, 'rb') as file:
        binaryData = file.read()
    return binaryData

@contextmanager
def get_db_cursor():
    connection = None
    cursor = None
    try:
        connection = mysql.connector.connect(**DB_CONFIG)
        cursor = connection.cursor()
        yield cursor
        connection.commit()
    except Exception:
        if connection is not None:
            connection.rollback()
        raise
    finally:
        if cursor is not None:
            cursor.close()
        if connection is not None and connection.is_connected():
            connection.close()


def get_lastnumber(cursor=None):
    def fetch_next_id(active_cursor):
        active_cursor.execute("SELECT MAX(id) FROM basic")
        row = active_cursor.fetchone()
        return ((row[0] if row else None) or 0) + 1

    try:
        if cursor is not None:
            return fetch_next_id(cursor)
        with get_db_cursor() as active_cursor:
            return fetch_next_id(active_cursor)
    except mysql.connector.Error as error:
        print(f"Failed to read next database id: {error}")
        return 0

def getStandardIndex():
    retRecord = [0,0,0,0,0]
    try:
        with get_db_cursor() as cursor:
            sql = (
                "SELECT average_scan_time, pages, count_broken_link, "
                "count_frame, google, mobile_tag, alt_image, hl1, hl2, hl3, flash "
                "FROM standard ORDER BY id DESC LIMIT 1"
            )
            cursor.execute(sql)
            row = cursor.fetchone()
            if row:
                retRecord = row
    except mysql.connector.Error as error:
        print(f"Failed selecting from standard MySQL table: {error}")
    return retRecord

def insertImage(id_num, url_address, photo):
    try:
        with get_db_cursor() as cursor:
            empPicture = convertToBinaryData(photo)
            sql = "INSERT INTO graph (id, url, graph) VALUES(%s, %s, %s)"
            cursor.execute(sql, (id_num, url_address, empPicture))
        return True
    except mysql.connector.Error as error:
        print(f"Failed inserting BLOB data into MySQL table: {error}")
        return False

def sorted_count_items(values):
    return sorted(values.items(), key=lambda item: (-item[1], item[0]))

def insert_count_rows(cursor, table_name, id_num, key_field, count_field, values, max_rows=None):
    recno = 0
    for key, count in sorted_count_items(values):
        if isinstance(key, str) and len(key) > 49:
            key = key[:49]
        sql = f"INSERT INTO {table_name} (id, {key_field}, {count_field}) VALUES(%s, %s, %s)"
        cursor.execute(sql, (id_num, key, count))
        recno += 1
        if max_rows is not None and recno >= max_rows:
            break

def report_to_db(photo):
    retValue = 1
    try:
        with get_db_cursor() as cursor:
            return _report_to_db(cursor, photo)
    except mysql.connector.Error as error:
        print(f"Failed inserting report data into MySQL table: {error}")
        return 0


def _report_to_db(cursor, photo):
        retValue = 1
        #empPicture = convertToBinaryData(photo)
        xlen = len(list_record)
        xcount = 0
        tlst = []
        s_para = ''
        sql_field = 'id, url_i, pid, '
        sql_para = '%s, %s, %s,'

        last_id = get_lastnumber(cursor)
        if last_id <= 0:
            raise mysql.connector.Error("Failed to allocate next database id")
        tlst.append(last_id)
        tlst.append(input_url)
        tlst.append(dbConnection_id)

        for xfield in list_record:
            sql_field = sql_field + xfield[0] 
            sql_para = sql_para + "%s"
            if xfield[0] in ['max_cl_url', 'min_cl_url', 'min_rt_url', 'max_rt_url']:
                if len(xfield[1]) > 254:
                    xfield[1] = xfield[1][:254]
            tlst.append(xfield[1])
            xcount = xcount + 1
            if xcount < xlen:
                sql_field = sql_field + ', '                
                sql_para = sql_para + ', '
        sql_args = tuple(tlst)

        sql = 'INSERT INTO basic (' + sql_field + ') VALUES(' + sql_para + ')'    
        print(sql, sql_args)
        result = cursor.execute(sql,sql_args)

        # cursor = connection.cursor()
        # print(photo)
        # empPicture = convertToBinaryData(photo)

        # sql = 'INSERT INTO graph (id, url, graph) VALUES(%s, %s, %s)'    
        # args = (last_id, baseUrl, empPicture, )
        # result = cursor.execute(sql,args)
        # connection.commit()

        ### ESG
        if len(esg_count) > 0:
            insert_count_rows(cursor, "esg", last_id, "keyword", "wcount", esg_count, keyWordList)
        
        if len(word_count) > 0:            
            insert_count_rows(cursor, "word", last_id, "keyword", "wcount", word_count, keyWordList)

        #####headdr information
        tlst = []
        s_para = ''
        sql_field = 'id, '
        sql_para = '%s, '
        xlen = len(list_header)
        xcount = 0

        tlst.append(last_id)

        if xlen > 0:
            for xlist in list_header:
                sql_field = sql_field + xlist[0] 
                sql_para = sql_para + "%s"
                tlst.append(xlist[1])
                xcount = xcount + 1
                if xcount < xlen:
                    sql_field = sql_field + ', '                
                    sql_para = sql_para + ', '
            sql_args = tuple(tlst)

            sql = 'INSERT INTO head (' + sql_field + ') VALUES(' + sql_para + ')'    
            print(sql, sql_args)
            result = cursor.execute(sql,sql_args)
 
        if len(list_sns) > 0:
            insert_count_rows(cursor, "sns", last_id, "sns_url", "sns_cnt", list_sns)

        tlst = []
        tlst.append(last_id)
        tlst.append(input_url)
        s_para = ''
        sql_field = 'id, url_i, '
        sql_para = '%s, %s, '
        xcount = 0
        xlen = len(list_domain)

        if xlen > 0:
            for xlist in list_domain:
                sql_field = sql_field + xlist[0] 
                sql_para = sql_para + "%s"
                tlst.append(str(xlist[1]))
                xcount = xcount + 1
                if xcount < xlen:
                    sql_field = sql_field + ', '                
                    sql_para = sql_para + ', '
            sql_args = tuple(tlst)
            sql = 'INSERT INTO domain (' + sql_field + ') VALUES(' + sql_para + ')'    
            print(sql, sql_args)
            result = cursor.execute(sql,sql_args)
            
        if len(forbiddenList) > 0:
            for fblist in forbiddenList:
                sql = 'INSERT INTO forbidden (id, url, fbword) VALUES(%s, %s, %s)'                 
                args = (last_id, fblist[0], fblist[1])
                print(sql, sql_args)
                result = cursor.execute(sql,args)

        if len(list_script)  > 0:   
            tmp_js_css=set() 
            for url,y,z in list_script:
                q = url.find("?")
                if q > 0:
                    url = url[:q]
                tmp_js_css.add(url)
            for url in tmp_js_css:
                sql = 'INSERT INTO js_css (id, url) VALUES(%s, %s)'                 
                args = (last_id, url)
                print(sql, args)
                result = cursor.execute(sql,args)
        return retValue

###insertBLOB(11, "http://www.hello.com", "c:/Users/user/Pictures/Capture.jpg")

def Debug_w(instr):
    if Debug_mode == True:
        print(instr)

def save_Favicon(url):
    from PIL import Image

    ensure_output_directories()
    parsed = urlparse(url)
    filename = Path(parsed.path).name or "favicon"
    source_path = IMAGE_DIR / filename
    r = HTTP_SESSION.get(
        url, allow_redirects=True, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
    )
    r.raise_for_status()
    source_path.write_bytes(r.content)

    with Image.open(source_path) as image:
        image.verify()
        image_format = image.format

    if image_format == "PNG" and source_path.suffix.lower() == ".png":
        return str(source_path)

    output_path = source_path.with_suffix(".png")
    with Image.open(source_path) as image:
        image.save(output_path, format="PNG")
    return str(output_path)


def save_url2image(url):
  from PIL import Image
  try:
    # Extract filename from URL
    filename = url.split('/')[-1]
    _, ext = os.path.splitext(filename)  # Extract file extension

    # Create save path
    b=baseUrl.replace(".", "_")
    save_path = IMAGE_DIR / f"{b}_{filename}"
    response = HTTP_SESSION.get(
        url, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
    )

    if response.status_code == 200 and "image" in response.headers.get("Content-Type", ""):
            # 파일 저장
            with open(save_path, "wb") as file:
                file.write(response.content)
            #print(f"다운로드 성공: {save_path}")
    else:
            print(f"이미지가 아닌 응답을 받음 [{url}]: {response.status_code}, {response.headers.get('Content-Type', '')}")

    # Verify image
    with Image.open(save_path) as img:
        img.verify()

    return str(save_path)
  except Exception as e:
    print(f"Error downloading image: {e}")
    return None
  
def image_copy(list_images):
    img_disk=set()
    for x in list_images:
        if "?" in x:
            continue
        fn = save_url2image(x)
        if fn:
            img_disk.add(fn)
        #print(img_disk)
       
    #print("**555**", img_disk)
    return img_disk

def get_limited_number(n, r):
    ret = n + r
    ret = ret / r
    ret = int(ret) * r
    return ret

def get_rgb_space(imgfile):
    image_path = Path(imgfile)
    if not image_path.is_file():
        return [0, 0, 0, 0, 0]
    img_color = cv2.imread(str(image_path))
    if img_color is None:
        return [0, 0, 0, 0, 0]
    rgb_color = cv2.cvtColor(img_color, cv2.COLOR_BGR2RGB)
    red = rgb_color[:, :, 0]
    green = rgb_color[:, :, 1]
    blue = rgb_color[:, :, 2]
    red_mask = (red > green) & (red > blue)
    green_mask = (green > red) & (green > blue)
    blue_mask = (blue > red) & (blue > green)

    color_R = int(np.count_nonzero(red_mask))
    color_G = int(np.count_nonzero(green_mask))
    color_B = int(np.count_nonzero(blue_mask))
    pix_T = int(rgb_color.shape[0] * rgb_color.shape[1])
    color_N = pix_T - color_R - color_G - color_B

    rgbn_colors = [color_R, color_G, color_B, color_N, pix_T]
    return rgbn_colors

def make_image_analysis_bar(file_list):    
    color_B = 0
    color_G = 0
    color_R = 0
    color_N = 0
    pixel_T = 0

    for fn in file_list:
        cR, cG, cB, cN, pT = get_rgb_space(fn)
        color_R = color_R + cR
        color_G = color_G + cG
        color_B = color_B + cB
        color_N = color_N + cN
        pixel_T = pixel_T + pT
                   
    if pixel_T == 0:
        return None

    bc = str(color_B/pixel_T)
    gc = str(color_G/pixel_T)
    rc = str(color_R/pixel_T)
    nc = str(color_N/pixel_T)

    bar_width = 500
    bar_height = 100
    text_string = "R.G.B.N = " + rc[:4] + ", " + gc[:4] + ", " + bc[:4] + ", " + nc[:4]
    font = cv2.FONT_HERSHEY_TRIPLEX
    text_color = (127,128,126)
    location=(10, int(bar_height * 0.8))

    bar_width = 500
    bar_height = 100
    bar_T = bar_width * bar_height

    bar_B = int((color_B/pixel_T) * bar_T)
    bar_G = int((color_G/pixel_T) * bar_T)
    bar_R = int((color_R/pixel_T) * bar_T)
    bar_N = int((color_N/pixel_T) * bar_T)

    image = np.zeros((bar_height, bar_width,3), np.uint8)

    bar_G = get_limited_number(bar_G, bar_width)
    bar_B = get_limited_number(bar_B, bar_width)
    bar_R = get_limited_number(bar_R, bar_width)

    image[:] = (220,220,220)
    bar_T=0

    for i in range(bar_height):
        for j in range(bar_width):
            bar_T = bar_T + 1
            if bar_G > bar_T:
                image[i,j] = (0,255,0)
                continue
            if bar_G + bar_B > bar_T:
                image[i,j] = (0,0,255)
                continue
            if bar_R + bar_B + bar_G > bar_T:
                image[i,j] = (255,0,0)
                continue
            if bar_R + bar_B + bar_G < bar_T:
                break
    b=baseUrl.replace(".", "_")
    save_path = IMAGE_DIR / f"_bar{b}.png"
    rgb_img = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    cv2.putText(rgb_img, text_string, location, font, 0.4, text_color, 1, cv2.LINE_AA)
    cv2.imwrite(str(save_path), rgb_img)
    # for x in range(1,30):
    #     print("************************************")
    return str(save_path)


def make_dominant_color_chart(file_list):
    all_pixels = []
    random_generator = np.random.default_rng(42)
    for image_file in file_list:
        image = cv2.imread(str(Path(image_file)))
        if image is None:
            continue
        image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        height, width = image_rgb.shape[:2]
        scale = min(1.0, 256.0 / max(height, width))
        if scale < 1.0:
            image_rgb = cv2.resize(
                image_rgb,
                (
                    max(1, round(width * scale)),
                    max(1, round(height * scale)),
                ),
                interpolation=cv2.INTER_AREA,
            )
        pixels = image_rgb.reshape(-1, 3)
        if len(pixels) > 5000:
            indexes = random_generator.choice(
                len(pixels),
                size=5000,
                replace=False,
            )
            pixels = pixels[indexes]
        all_pixels.append(pixels)

    if not all_pixels:
        return None

    pixels = np.vstack(all_pixels)
    if len(pixels) > 100000:
        indexes = random_generator.choice(
            len(pixels),
            size=100000,
            replace=False,
        )
        pixels = pixels[indexes]
    n_clusters = min(5, len(pixels))
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    labels = kmeans.fit_predict(pixels)
    colors = np.round(kmeans.cluster_centers_).astype(int)
    counts = np.bincount(labels, minlength=n_clusters)
    order = np.argsort(counts)[::-1]
    colors = colors[order]
    percentages = counts[order] * 100.0 / counts.sum()

    site_name = baseUrl.replace(".", "_")
    save_path = IMAGE_DIR / f"{site_name}_CCAI.jpg"
    fig, ax = plt.subplots(figsize=(7, 4))
    positions = np.arange(n_clusters)
    ax.bar(
        positions,
        percentages,
        color=colors / 255.0,
        edgecolor="black",
    )
    ax.set_xticks(
        positions,
        ["#{:02X}{:02X}{:02X}".format(*color) for color in colors],
    )
    ax.set_ylabel("Pixel ratio (%)")
    ax.set_title(f"Dominant Colors Across Images - {baseUrl}")
    fig.tight_layout()
    fig.savefig(save_path)
    plt.close(fig)
    return str(save_path)


def normalize_css_color(value):
    value = value.strip().lower()
    if not value or value.startswith(("var(", "inherit", "initial", "unset")):
        return None
    value = value.split("!important", 1)[0].strip()
    rgb_match = re.fullmatch(
        r"rgba?\(\s*(\d+(?:\.\d+)?)\s*,\s*(\d+(?:\.\d+)?)\s*,"
        r"\s*(\d+(?:\.\d+)?)(?:\s*,\s*[\d.]+)?\s*\)",
        value,
    )
    if rgb_match:
        channels = [
            max(0, min(255, round(float(channel))))
            for channel in rgb_match.groups()
        ]
        return "#{:02X}{:02X}{:02X}".format(*channels)
    try:
        return matplotlib_colors.to_hex(value, keep_alpha=False).upper()
    except ValueError:
        return None


def extract_font_colors(css_text):
    colors = {}
    declarations = re.findall(
        r"(?<![-\w])color\s*:\s*([^;}{]+)",
        css_text or "",
        flags=re.IGNORECASE,
    )
    for declaration in declarations:
        color = normalize_css_color(declaration)
        if color:
            colors[color] = colors.get(color, 0) + 1
    return colors


def merge_color_counts(target, source):
    for color, count in source.items():
        target[color] = target.get(color, 0) + count
    return target


def make_font_color_chart(color_counts):
    if not color_counts:
        return None

    ranked = sorted(
        color_counts.items(),
        key=lambda item: item[1],
        reverse=True,
    )[:12]
    labels = [item[0] for item in ranked]
    counts = [item[1] for item in ranked]
    save_path = IMAGE_DIR / f"{baseUrl.replace('.', '_')}_FONT_COLORS.png"

    fig, ax = plt.subplots(figsize=(8, 4))
    positions = np.arange(len(labels))
    ax.bar(positions, counts, color=labels, edgecolor="black")
    ax.set_xticks(positions, labels, rotation=45, ha="right")
    ax.set_ylabel("CSS declaration count")
    ax.set_title(f"Declared Font/Text Colors - {baseUrl}")
    fig.tight_layout()
    fig.savefig(save_path)
    plt.close(fig)
    return str(save_path)
        
def skip_costsavingMode(yy):
    normalized_headers = {key.lower(): value for key, value in yy.items()}
    ct = normalized_headers.get("content-type", "")
    if ct:
        if check_string(ct.lower(), ["hwp", "msword", "pdf", "octet-stream","video", "mp4",
                                    "vnd.ms", "zip", "multipart", "flash", "download",
                                    "audio", "mp3", "image", "file","jpg",
                                    "jpge", "midi", "mpeg", "force-download", "octet-stream", 
                                    "x-shockwave-flash", "json",
                                    "x-ms-wma",
                                    "vnd.apple.installer+xml", "vnd.oasis.opendocument.presentation", 
                                    "vnd.oasis.opendocument.spreadsheet", "vnd.oasis.opendocument.text", 
                                    "ogg", "vnd.ms-powerpoint", "x-rar-compressed", "rtf", "x-sh", "svg+xml", 
                                    "x-shockwave-flash", "x-tar", "tiff", "x-font-ttf", 
                                    "vnd.visio", "x-wav", "webm", "webp", "x-font-woff", 
                                    "vnd.ms-excel", "vnd.mozilla.xul+xml", "3gpp", "x-7z-compressed", "unknown"]):
            return True
        
        content_disposition = normalized_headers.get("content-disposition", "")
        if content_disposition:
            if "attachment" in content_disposition.lower():
                return True
    return False

def multiple_url(addr):
    qlist = addr.split("?")
    if len(qlist) > 2:
        if qlist[1] == qlist[2]:
            print("Warning url ?: ", addr)
            return True
    
    x = addr.find("'")
    if x > 0:
        addr = addr[:x]
    x = addr.find('"')
    if x > 0:
        addr = addr[:x]
    
    #print("*********************", addr, "**************", len(addr))
    try:
        v_url = urlparse(addr)
        v_path = v_url[2]
    except:
        return False

    slist = v_path.split("/")
    if len(slist) > 5:
        sset = set()
        for x in slist:
            sset.add(x)
        if len(sset) + 2 < len(slist):
            print("Warring url /: ", addr, baseUrl)
            return True

    return False

def cmp_domain(x,y):  ##cur_url , baseUrl
    if "//" not in x:
        return 3
    x2=x.split(".")
    y2=y.split(".")

    yl = len(y2)
    xl = len(x2)
    fi = xl - yl

    cflag = 0
    for i in range(0,yl):    
        if y2[i] == x2[i+fi]:
            cflag = cflag + 1
    return(cflag)

def get_path(p1):
    p_url = urlparse(p1)
    v_url = p_url[0] + "://" + p_url[1] + p_url[2]

    return v_url

def getHttp_Https(u):
    if "https".lower() in u:
        return "https"
    else:
        return "http"

def test_12Byte(word):  ## test 1-byte or 2-byte word.
    for x in word:
        v = ord(x)
        if (v > 255):            
            if v > 60000:
                return -1
            if v < 10000:
                return -1
            return 2
        elif (v < 32):
            return -1
    return 1

def replaceMultiple(mString, charSet, nChar):
    for tchar in mString :
        if tchar in charSet :
            mString = mString.replace(tchar, nChar)    
    return  mString  

def check_string(string, substring_list):
    if string is None or substring_list is None:
        return False  
    for substring in substring_list:
        if substring in string:
            return True
    return False

def check_string2(string, substring_list):
    return check_string(string, substring_list)

def check_Prefixstring(string, substing_list):
    for substring in substing_list:
        slength = len(substring)
        if len(string) < slength:
            continue
        if string[:slength] == substring:
            return True
        #print(substring, string, string[:slength])
    return False

def check_string_Length(string, substring_list, sl):
    if len(string) > sl:
        return True
    for substring in substring_list:
        if substring in string:
            return True
    return False

def check_stringAll(string, substring_list):
    list_tmp=[]
    for substring in substring_list:
        if substring in string:            
            list_tmp.append(substring)
    return list_tmp

def normalize_keyword_text(value):
    value = unicodedata.normalize("NFKC", str(value or ""))
    value = re.sub(r"\s+", " ", value)
    return value.strip().casefold()

def normalized_esg_keywords():
    keywords = []
    seen = set()
    for keyword in nplt_esg_word:
        clean_keyword = unicodedata.normalize("NFKC", str(keyword)).strip()
        clean_keyword = re.sub(r"\s+", " ", clean_keyword)
        if not clean_keyword:
            continue
        normalized = normalize_keyword_text(clean_keyword)
        if normalized in seen:
            continue
        seen.add(normalized)
        keywords.append((clean_keyword, normalized))
    keywords.sort(key=lambda item: len(item[1]), reverse=True)
    return keywords

def contains_english_word_boundary(text, keyword):
    if not re.fullmatch(r"[a-z0-9][a-z0-9 -]*", keyword):
        return keyword in text
    pattern = r"(?<![a-z0-9])" + re.escape(keyword) + r"(?![a-z0-9])"
    return re.search(pattern, text) is not None

def counting_esg_word(txt):     
     ls = set()
     normalized_text = normalize_keyword_text(txt)
     for original, normalized in normalized_esg_keywords():
        if contains_english_word_boundary(normalized_text, normalized):
            esg_count[original] = esg_count.get(original,0) + 1
            ls.add(original)
     return ls

def normalize_domain_host(host_or_url, strip_common_subdomain=False):
    if not host_or_url:
        return ""
    value = str(host_or_url).strip()
    parsed = urlparse(value if "://" in value else f"http://{value}")
    host = parsed.hostname or value.split("/")[0]
    host = host.strip().strip(".").lower()
    if strip_common_subdomain and host.startswith(("www.", "m.", "mobile.")):
        host = host.split(".", 1)[1]
    return host

def normalize_sns_host(host_or_url):
    return normalize_domain_host(host_or_url, strip_common_subdomain=True)

def host_matches_domain(host, domain):
    host = normalize_sns_host(host)
    domain = normalize_sns_host(domain)
    return host == domain or host.endswith("." + domain)

def identify_sns_platform(url_or_host):
    host = normalize_sns_host(url_or_host)
    for platform, domains in SNS_PLATFORMS.items():
        if any(host_matches_domain(host, domain) for domain in domains):
            return platform
    legacy_domains = {normalize_sns_host(domain) for domain in sns_domain_list}
    if host in legacy_domains:
        return host
    return None

def classify_sns_url(url):
    parsed = urlparse(url if "://" in str(url) else f"http://{url}")
    path = (parsed.path or "").lower()
    host = normalize_sns_host(parsed.hostname or parsed.netloc)
    if any(marker in path for marker in SNS_SHARE_PATHS):
        return "share"
    if any(marker in path for marker in SNS_EMBED_PATHS):
        return "embed"
    if parsed.query and any(key in parsed.query.lower() for key in ("url=", "u=", "text=")):
        return "share"
    if host == "youtu.be" and path not in ("", "/"):
        return "content"
    if path in ("", "/"):
        return "profile"
    if re.search(r"/(watch|shorts|reel|reels|posts|status|tweet|news|article)/?", path):
        return "content"
    return "profile"

def normalize_sns_url(url):
    parsed = urlparse(url if "://" in str(url) else f"http://{url}")
    scheme = parsed.scheme or "http"
    host = normalize_sns_host(parsed.hostname or parsed.netloc)
    path = parsed.path or "/"
    if path != "/":
        path = path.rstrip("/")
    query_items = [
        (key, value)
        for key, value in urllib.parse.parse_qsl(parsed.query, keep_blank_values=True)
        if not key.lower().startswith("utm_")
    ]
    query = urllib.parse.urlencode(sorted(query_items))
    normalized = f"{scheme}://{host}{path}"
    if query:
        normalized = f"{normalized}?{query}"
    return normalized

def record_sns_link(url, source_url=None):
    platform = identify_sns_platform(url)
    if not platform:
        return None
    link_type = classify_sns_url(url)
    normalized_url = normalize_sns_url(url)
    detail = sns_details.setdefault(
        platform,
        {
            "urls": set(),
            "pages": set(),
            "types": {},
            "representative_url": normalized_url,
        },
    )
    detail["urls"].add(normalized_url)
    if source_url:
        detail["pages"].add(source_url)
    detail["types"][link_type] = detail["types"].get(link_type, 0) + 1
    if detail.get("representative_url") in ("", None):
        detail["representative_url"] = normalized_url
    list_sns[platform] = len(detail["urls"])
    return platform

def build_sns_report_rows():
    rows = []
    for platform, detail in sorted(
        sns_details.items(),
        key=lambda item: (len(item[1]["urls"]), item[0]),
        reverse=True,
    ):
        types = ", ".join(
            f"{name}:{count}"
            for name, count in sorted(detail["types"].items())
        )
        rows.append(
            [
                platform,
                types or "unknown",
                str(len(detail["pages"])),
                str(len(detail["urls"])),
                detail.get("representative_url", ""),
            ]
        )
    if rows:
        return rows
    for platform, count in sorted(list_sns.items(), key=lambda item: (-item[1], item[0])):
        rows.append([platform, "unknown", "", str(count), ""])
    return rows

def get_varList(shtml):
    shtml=shtml.replace("==","")    
    thtml = shtml.split("=")
    varList = set()
    qm = ['"', "'", ".", ",", "/", " ", "script", "\n", "\t", "\r"]

    if len(thtml) < 2:
        return varList

    for x in range(0,len(thtml)-1,1):  ## 2020.3.1 range(0,len(thtml)-1,2):
        s1 = thtml[x]
        s2 = thtml[x+1]
        s1=s1.strip()
        n1=s1.split(" ")
        for n2 in n1:
            if len(n2) >2:
                if (check_string(n2, qm)==False):
                    varList.add(n2)
    #print(varList)
    if "function" in shtml:
        varList = varList.union(extract_functionParamater(shtml))
    return varList

def slice_before_character(url, character):
    index = url.find(character)
    if index != -1:
        return url[:index]  # index 이전까지 잘라냄
    return url  # 문자가 없는 경우 원래 문자열 반환


def get_parentPath(p1,depths):   ##../site_map/sitemap.php currURL=http://www.upcsinfor.com/index.php
    c=0
    dp = depths + 1
    retValue = p1
    p_url = urlparse(p1)
    if p_url[2]=="" or p_url[2]=="/":
        return retValue
    
    if "?" in p1:
        p1=p1.split("?")[0]

    for i in range(1,len(p1)-8, 1):  #index.php currURL="http://www.sambotruck.com/n-member/info"
        t = i * -1
        if p1[t] == "/":
            c = c + 1
            if (c==1):
                retValue = p1[:len(p1) - i]
            if (c==dp):
                retValue = p1[:len(p1) - i]
                break

    return retValue

def adjustUrl(xurl):
    retUrl = xurl
    if urlparse(xurl)[1] == "":
        retUrl = "http://" + xurl
    return retUrl

def get_domain(inpath):
    if inpath.find(':') < inpath.rfind(':'):
        r =inpath.rfind(':')
        inpath = inpath[:r]
    if inpath[-1]!="/":
        inpath=inpath+"/"
    s1 = inpath.find("//")
    if s1 < 0:
        s1 = 0
    else:
        s1 = s1 + 2
    s2 = inpath.find("/", s1)
    #print(s1, s2, inpath[s1:s2],inpath)
    return (inpath[s1:s2])

def relation_domain(u1,b1):
    u = get_domain(u1)
    b = get_domain(b1)
    p = "......"

    uc = u.count(".")
    bc = b.count(".")
    if uc == bc + 1:
        if "www." in u and "www." not in b:
            b = "www." + b
            bc = bc + 1
    if bc == uc + 1:
        if "www." not in u and "www." in b:
            u = "www." + u
            uc = uc + 1
            
    cc = bc + 1

    if uc > bc:
        dc = uc - bc
        b = p[:dc] + b
        cc = bc + 1
    if bc > uc:
        dc = bc - uc
        u = p[:dc] + u
        cc = uc + 1

    bp = b.split(".")
    up = u.split(".")

    cflag = 0
    for i in range(len(bp)):
        if bp[i] == up[i]:
            cflag = cflag + 1
    if cc == cflag:
        return 1
    else:
        return -1

def not_url_but_parameter(inpath):
    return check_string(inpath, list_para)

def add_list_script(cpath, curls, ctype):
    tmp=[]
    tmp.append(cpath)
    tmp.append(curls)
    tmp.append(ctype)
    list_script.append(tmp)  
## add extion cout   
    extion_count[ctype] = extion_count.get(ctype,0) + 1    
    
def extract_script_file(inLink, urls):
    filepath = inLink.get("src")
    inpath   = str(filepath)
    if "." in inpath:
        exf = inpath.split('.')[-1].lower()
        if len(exf) < 5:
            add_list_script(inpath, urls, exf)

def extract_function(inHtml, urls):     
    if "function " in inHtml:
        fnlist = []
        #z = inHtml.replace(" ","") 
        z=inHtml
        #print(z)
        x = z.find("function") + 9
        l = z.find("(")
        r = z.find(")")
        if x> -1 and l < r and x <l:
            fn = z[x:l]
            fp = z[l+1:r]
            fnlist.append(fn)
            fnlist.append(fp)
            if len(fp) > 0:
                list_para.add(fp)
            w = z.find("window.open")
            if w > 0:
                fu = z[z.find("(",w)+1:z.find(",",w)]
                fu=fu.replace('"','')
                fu=fu.replace("'","")
                fnlist.append(fu)
            else:
                fnlist.append("*")

            fnlist.append(len(z.split(fn)))
            fnlist.append(urls)
            list_function.append(fnlist)

def extract_functionParamater(inHtml):     
    paraList = set()
    z=inHtml
    while(1):
        if "function " in z:   
            x = z.find("function") + 7
            l = z.find("(",x)
            r = z.find(")",l)
            if x > -1 :
                fp = z[l+1:r]
                ##print("---------------",fp)
                if len(fp) > 0:
                    for u in fp.split(","):
                        paraList.add(u)
            z=z[(x+9):]
        else:
            break
            
    return paraList

def get_extion(inpath):  ###http://www.samshin.es.kr/viewer/skin/doc.html?fn=6efe1609bc80c35d245591d53b9c60de8858780636bc81b8_0.hwp&rs=/viewer/result/20210
    q = inpath.find("?")
    if q > 0:
        inpath = inpath[:q]
    exf = urllib.parse.urlparse(inpath)[4]

    if len(exf) == 0:
        exf = urllib.parse.urlparse(inpath)[2]
    if len(exf) > 0:
        if "." in exf:   #fn=6efe1609bc80c35d245591d53b9c60de8858780636bc81b8_0.hwp&rs=/viewer/result/20210   20210516
            exf = exf.split('.')[-1]
            if len(exf) > 3:
                if "&" in exf:
                    exf = exf.split("&")[0]
                else:
                    exf = exf[:3]   
        else:
            exf = "*"
    else:
        exf = "*"       
    return exf.lower()

def skip_extion(exf):
    if exf in skip_ext_list:
        return True
    else:
        return False

def get_parent_path(cpath):
    try:
        rindex = cpath.rindex("/")
        if rindex == 0:
            ppath="/"
        else:
            ppath = cpath[:rindex]
    except:
        ppath="/"
    return(ppath)


def normalize_host(host):
    return normalize_domain_host(host, strip_common_subdomain=True)


def tree_path_for_url(value, site_host):
    value = str(value or "").strip()
    if not value or value.startswith(("mailto:", "tel:", "javascript:")):
        return None

    parsed = urlparse(value)
    if parsed.scheme in {"http", "https"}:
        if normalize_host(parsed.hostname) != normalize_host(site_host):
            return None
        path = parsed.path or "/"
    elif value.startswith("/"):
        path = urlparse(value).path or "/"
    else:
        return None

    path = re.sub(r"/+", "/", path)
    if path != "/" and path.endswith("/"):
        path = path[:-1]
    return path or "/"


def build_tree_map_data(edges, visited_urls, site_host, simple_mode=True):
    paths = {"/"}
    counts = {}
    external_urls = {}

    for edge in edges:
        for value in edge:
            path = tree_path_for_url(value, site_host)
            if path is not None:
                display_path = simpleUrl2(path) if simple_mode else path
                paths.add(display_path or "/")
                parent = get_parent_path(display_path)
                while parent and parent != "/":
                    paths.add(parent)
                    parent = get_parent_path(parent)
            else:
                parsed = urlparse(str(value))
                if parsed.scheme in {"http", "https"} and parsed.hostname:
                    host = normalize_host(parsed.hostname)
                    external_urls.setdefault(host, set()).add(
                        parsed._replace(query="", fragment="").geturl()
                    )

    for value in visited_urls:
        path = tree_path_for_url(value, site_host)
        if path is None:
            continue
        display_path = simpleUrl2(path) if simple_mode else path
        display_path = display_path or "/"
        counts[display_path] = counts.get(display_path, 0) + 1

    external_counts = {
        host: len(urls)
        for host, urls in sorted(external_urls.items())
    }
    return paths, counts, external_counts


def reduce_string(instr):
    tmp=instr
    ix = tmp.find("<img")
    if ix > -1:
        dx = tmp.find("data:image/", ix)
        if dx > -1:
            tmp = tmp[:dx]
            # lx = tmp.find(">",dx)
            # if lx > -1:
            #     tmp = tmp[:dx] + tmp[lx+1:]
            #     print(ix,dx,lx, len(instr), len(tmp))
    return(tmp)

def remove_Remark_str(in_html):
    s = in_html
    x = s.find("<!--")
    if x > 0:
        y = s.find("-->",x+1)
        if y > 0:
            s = s[:x] + s[y+3:]
            s = remove_Remark_str(s)
    return s

def remove_Remark_bs(bb):
    ss = str(bb)
    if ss.find("<!--"):
        ss = remove_Remark_str(ss)
        bb = bs(ss, "html.parser")
    return bb

def remove_String(bb, pos):
    sbb = bb
    if "img src" in sbb or "background:url" in sbb:
        x = sbb.find("data:image",pos)
        #print(x, sbb[x-20:x+20])
        if x > 0:
            ch = sbb[x-1]
            x1 = sbb.find(ch,x+1)
            #print(ch, x1)
            if x1 > 0:
                sbb = sbb[:x+20] + sbb[x1:]
                sbb = remove_String(sbb, x+21)
    return sbb

def style_remove_string(str):
    return re.sub(r'<style.*?>.*?</style>', '', str, flags=re.DOTALL)

def new_remove_string(bb,pos):
    if "data:img" in bb:    #"data:img/png;base64
        s1 = bb.find("data:img")
        s2 = bb.find("base64", s1)
        s3 = bb.find(">", s2)
        ##s4 = bb.find("</img", s3)  please consider </img> tag
        r1 = bb.rfind("<",s1-123, s1)
        bb = bb[:r1-1]+bb[s3+1:]
        bb = new_remove_string(bb,r1)
    if "data:image" in bb:    #"data:img/png;base64
        s1 = bb.find("data:image")
        s2 = bb.find("base64", s1)
        s3 = bb.find(">", s2)
        ##s4 = bb.find("</img", s3)  please consider </img> tag
        r1 = bb.rfind("<",s1-123, s1)
        bb = bb[:r1-1]+bb[s3+1:]
        bb = new_remove_string(bb,r1)
    if "data:video" in bb:
        s1 = bb.find("data:video")
        s2 = bb.find("base64", s1)
        s3 = bb.find(">", s2)
        r1 = bb.rfind("<",s1-123, s1)
        bb = bb[:r1-1]+bb[s3+1:]
        bb = new_remove_string(bb,r1)
    if "data:audio" in bb:
        s1 = bb.find("data:audio")
        s2 = bb.find("base64", s1)
        s3 = bb.find(">", s2)
        r1 = bb.rfind("<",s1-123, s1)
        bb = bb[:r1-1]+bb[s3+1:]
        bb = new_remove_string(bb,r1)
    return bb

def enhanced_remove_string(bb,pos):
    
    k1 = bb.find("data:image") 
    k2 = bb.find("data:video") 
    k3 = bb.find("data:audio") 
    s1 = -28
    if k1 > 0:
        s1 = k1
    elif k2 > 0:
        s1 =k2
    elif k3 > 0:
        s1 = k3    

    if s1 > 0:
        s2 = bb.find("base64", s1)
        s3 = bb.find(">", s2)
        r1 = bb.rfind("<",s1-123, s1)
        if r1 < 0:
            r1 = bb.rfind("url",s1-30,s1)
            s3 = bb.find(")", s2)
        #print(s1,s2,s3,r1) #, bb[s1-30:s2])
        if r1 > 0 and s3 > 0:
            bb = bb[:r1-1]+bb[s3+1:]
            bb = enhanced_remove_string(bb,r1)
    return bb

def parsing_fontlist(slink):
    text = str(slink)
    declarations = re.findall(
        r"font-family\s*:\s*([^;}{]+)",
        text,
        flags=re.IGNORECASE,
    )
    shorthand_values = re.findall(
        r"(?<![-\w])font\s*:\s*([^;}{]+)",
        text,
        flags=re.IGNORECASE,
    )
    face_values = [
        next(value for value in groups if value)
        for groups in re.findall(
            r"\bface\s*=\s*(?:\"([^\"]+)\"|'([^']+)'|([^\s>]+))",
            text,
            flags=re.IGNORECASE,
        )
    ]

    for shorthand in shorthand_values:
        family_match = re.search(
            r"(?:\d*\.?\d+(?:px|pt|pc|em|rem|%|vh|vw)|"
            r"xx-small|x-small|small|medium|large|x-large|xx-large)"
            r"(?:\s*/\s*[^\s]+)?\s+(.+)$",
            shorthand,
            flags=re.IGNORECASE,
        )
        if family_match:
            declarations.append(family_match.group(1))

    fonts = set()
    for declaration in declarations + face_values:
        for font_name in re.findall(
            r'"([^"]+)"|\'([^\']+)\'|([^,]+)',
            declaration,
        ):
            name = next(value for value in font_name if value).strip()
            if name:
                fonts.add(name)
    return fonts


def select_font(font_cccc):
    ignored = {"inherit", "initial", "revert", "revert-layer", "unset"}
    return sorted(
        {
            re.sub(
                r"\s*!important\s*$",
                "",
                str(font).strip().strip("'\""),
                flags=re.IGNORECASE,
            ).strip()
            for font in font_cccc
            if re.sub(
                r"\s*!important\s*$",
                "",
                str(font).strip().strip("'\""),
                flags=re.IGNORECASE,
            ).strip().lower()
            not in ignored
        },
        key=str.lower,
    )

def formatHTTP(tag): #adding http to tag 
    if "http" in tag.lower():
        return tag
    else:
        return http_https + "://" + tag

def simpleUrl(instr):
    retUrl ="/"

    vi = instr.rfind(".")
    if vi == -1:
        if (len(instr) > 1) and (instr not in ["http:", "https:", "http:/", "https:/", "HTTP:", "HTTPS:"]):
            retUrl = instr            
    else:
        xr = instr.rfind("/")
        if instr[:4].lower() == "http":
            retUrl = instr
        else:
            if (vi > xr and xr > 1):
                retUrl = instr[:xr]
    vi = retUrl.rfind("#")
    if vi > 0:
        retUrl = retUrl[:vi]
    if retUrl[-1:] == "/":  ####  2019.11.23 silicon
        retUrl = retUrl[:-1]
        
    return retUrl

def simpleUrl2(instr):
    retUrl = simpleUrl(instr)
    vi = retUrl.rfind("?")
    if vi > 0 :
        retUrl = retUrl[:vi]       
    vi = retUrl.find("/&")      #/skin45/&boardT=v&b  2020.2.25
    if vi > 0 :
        retUrl = retUrl[:vi]     
    return retUrl

def toSimpkeList(inlist):
    retlist = list()

    for instr in inlist:
        instr = simpleUrl2(instr)
        if instr[-1:] == "/":
            instr = instr[:-1]
        if instr=="/":
            continue
        if instr not in retlist:
            retlist.append(instr)
    return (retlist)     

def toSimpkeList2(inlist):
    retlist = list()

    for instr in inlist:
        if instr[-1:]=="/":
            continue
        vi = instr.rfind(".")
        if vi == -1:
            if (len(instr) > 1) and (instr not in ["http:", "https:", "http:/", "https:/", "HTTP:", "HTTPS:"]):
                if instr not in retlist:
                    retlist.append(instr)
        else:
            xr = instr.rfind("/")
            if instr[:4].lower() == "http":
                if instr not in retlist:
                    retlist.append(instr)
            else:
                if (vi > xr and xr > 1):
                    if instr[:xr] not in retlist:
                        retlist.append(instr[:xr])
    return (retlist)                
                                #### c   2020.08.19
                                #### window.location.href=/ETJSP/index.jsp--></script>
def get_locationhref(lstr, sflag):     ###type=text/javascript>location.href=/html/00_main/main.php;</script>
    if sflag==0:
        if len(lstr) > 500:
            return None  ### redirection 인경우 script가 간단하다 복잡하면 아니다
        else:
            x = lstr.lower().rfind("location.href") 
    else:
        x = lstr.lower().find("location.href")    ##location.href=/memberCreate.do?left_menu=5;
        ###location.href = 'intent://' + o.param + '#Intent;' + o.g_proto + ';end'}, 100);
 
    tstr = None
    ep = lstr.find("=",x)
    se = lstr.find(";",x)
    # if se < ep:
    #     return None
    lh = lstr.find("location.host", ep)  ### 20220313
    if lh > ep:
        pl = lstr.rfind("+", lh)
        if pl > ep:
            ep = pl
    if ep > x:
        tstr = lstr[ep+1:]
        tstr=tstr.replace(" ","").replace('"','').replace("'","")
        y = tstr.find(";")
        if y == -1:
            y = tstr.find("-->")
        if y == -1:
            y = tstr.find("<")
        if y > 0:
            tstr = tstr[:y]
        
        #print("DEBUG===>", lstr,tstr,x,y)
        
        if tstr[0]=="#":
            return None

        if check_string(tstr, ["(", ">", "+", ":"]):
            return None
        #print("DEBUG===>", lstr,tstr,x,y)
        
    if x < 0 and y < 0:
        tstr=None
    
    return(tstr)

def get_location(lstr):         ##document.location=login.asp;
    x = lstr.lower().find("location=")      ##location=no,menubar=no,resizable=yes
    #print(lstr, x)  ##self.location=/sub/sub1-1.htm;   ##self.location=/products/products.php?cat=100;
    tstr = None
    if x > 0:
        tstr = lstr[x+9:]     
        y = tstr.find(";")
        x = tstr.find("<")
        if y > -1:
            if y < x:
                y=y    
            tstr = tstr[:y]
        if "," in tstr:
            tstr=None
    return(tstr)

def get_location2(lstr):
    x = lstr.lower().find("$(location)")
    tsrt = None
    if x > -1:
        tstr = lstr[x+11:]  #.attr('href', "/kor/index.do",'aaa',"qqq");
        lb = tstr.find("(")
        rb = tstr.find(")")
        if rb > lb and lb > 0:
            tstr = tstr[lb+1:rb]
            tstr = tstr.replace('"','')
            tstr = tstr.replace("'","")
            sd = tstr.split(",")
            sd_inx = 0       
            tsrt = None     
            for xkey in sd:
                if xkey == "href":
                    tstr = sd[sd_inx + 1]
                    break
                sd_inx = sd_inx + 1
    return(tstr)

def get_href2(lstr):    ##<a class="slider" onclick="location.href='/business/field01.asp'">
                        ##<a href="/dbimage/ipsi/WebData/data/pdf/2020.pdf" target="_blank">
           #  <use xlink:href=/modules/my-apostrophe-svg-sprites/svg/sprite-glyphs.svg#glyphs-blue-arrow-north></use></svg>
    #lstr=lstr.replace(" ","")
    #<a   href="?mode=view&amp;article_no=146881&amp;board_wrapper=%2Fko%2Fstate%2Finfo08.jsp&amp;pager.offset=0&amp;board_no=6&amp;default:category_id=170">
    x = lstr.lower().find("href=") 
    tstr = lstr[x+5:]
    y = tstr.find(">")
    z = tstr.find(" ")
    if z > -1:
        if y > z:
            y = z
    if y > -1:
        tstr = tstr[:y]
        tstr = tstr.replace('"','')
        tstr = tstr.replace("'","")
    else:
        tstr = ","   ##  None

    if check_string(tstr, ["(", ">", "+", ":", ","]):
        tsr = None
    return(tstr)


def get_href(lstr):    ##<a class="slider" onclick="location.href='/business/field01.asp'">
                        ##<a href="/dbimage/ipsi/WebData/data/pdf/2020.pdf" target="_blank">
           #  <use xlink:href=/modules/my-apostrophe-svg-sprites/svg/sprite-glyphs.svg#glyphs-blue-arrow-north></use></svg>
    #lstr=lstr.replace(" ","")
    #<a   href="?mode=view&amp;article_no=146881&amp;board_wrapper=%2Fko%2Fstate%2Finfo08.jsp&amp;pager.offset=0&amp;board_no=6&amp;default:category_id=170">
    pattern = r'href\s*=\s*"(.*?)"'
    match = re.search(pattern, lstr)
    tstr = None
    if match:
        tstr = match.group(1).strip()  # 앞뒤 공백 제거 후 반환
    if check_string(tstr, ["(", ">", "+", ":", ","]):
        tstr=None
    return tstr

def checkNOTblank(str, pos):
    for st in range(0, pos-1):
        if str[st] != " ":
            return 1
    return 0


def find_first_char_from_index(c, s, n):  ## character, string, index
    if n< 0 or n >= len(s):
        return -1
    for i in range(n, -1, -1):
        if s[i] == c:
            return i
    return -1

def get_windowlocationhref(lstr):   
    ### var current_path =window.location.pathname;
    ##window.location.origin+"/")&&(r.crossOrigin="anonymous"),document.head.appendChild(r)}(e,r,a,n)}))},i={6658:0},c.f.miniCss=function(e,a){i[e]?a.push(i[e]):0!==i[e]&&{1344:1,1407:1,1475:1,2081:1,2262:1,2395:1,3973:1,4102:1,5246:1,8138:1,8389:1,9279:1,9941:1}[e]&&a.push(i[e]=o(e).then((function(){i[e]=0}),(function(a){throw delete i[e],a})))},function(){var e={6658:0};c.f.j=function(a,n){var t=c.o(e,a)?e[a]:void 0;if(0!==t)if(t)n.push(t[2]);else if(6658!=a){var r=new Promise((function(n,r){t=e[a]=[n,r]}));n.push(t[2]=r);var o=c.p+c.u(a),i=new Error;c.l(o,(function(n){if(c.o(e,a)&&(0!==(t=e[a])&&(e[a]=void 0),t)){var r=n&&("load"===n.type?"missing":n.type),o=n&&n.target&&n.target.src;i.message="Loading chunk "+a+" failed.\n("+r+": "+o+")",i.name="ChunkLoadError",i.type=r,i.request=o,t[1](i)}}),"chunk-"+a,a)}else e[a]=0},c.O.j=function(a){return 0===e[a]};var a=function(a,n){var t,r,o=n[0],i=n[1],s=n[2],d=0;for(t in i)c.o(i,t)&&(c.m[t]=i[t]);if(s)var l=s(c);for(a&&a(n);d<o.length;d++)r=o[d],c.o(e,r)&&e[r]&&e[r][0](),e[o[d]]=0;return c.O(l)},n=self.webpackJsonp__wix_thunderbolt_app=self.webpackJsonp__wix_thunderbolt_app||[];n.forEach(a.bind(null,0)),n.push=a.bind(null,n.push.bind(n))}()}();
    # sourceMappingURL=https://static.parastorage.com/services/wix-thunderbolt/dist/webpack-runtime.e6d563b4.bundle.min.js.map</script>
    
    x = lstr.lower().find("window.location")  
    y = find_first_char_from_index('=', lstr, x)
    if (x-y) < 3:  ## = windows.location
        return None

    q = lstr[x+15:]  
    n = q.find("=")
    if q[0]==".":
        return None
    if n < 0:
        return None
    if checkNOTblank(q,n):
        return None
    # if q.find("=") > -1:
    tstr = q.split("=")[1]
    y = tstr.find(";")
    r = tstr[:y]
    r = r.replace('"','')
    r = r.replace("'","")
    # else:
    #     r = None
    if check_string(r, ["(", "+"]):
        return None
    return (r)

def set_html_from_load(tlink):
    rlist = []
    xlist = tlink.split(".load(")
    for n in range(1, len(xlist)-1):
        ustr = xlist[n].split(")")[0]
        ustr = ustr.replace("'","")
        ustr = ustr.replace('"','')
        rlist.append(ustr)
    return rlist
        
def collect_path(tmp):  ##'/userfiles/images/%ED%83%9C%EB%B8%94%EB%A1%9C3(3).png' << modi 20220120
    #print("collect url",tmp)
    xd = tmp.find(".")
    xp = tmp.find("+")
    xm = tmp.find("-")
    xb = tmp.find("[")
    
    if xd > xp and xp > 0:
        return False
    if xd < xm and xd > 0:
        return False
    if xd > xb and xb > 0:
        return False    
    if check_string(tmp, [";", "=+", "=-", "@"]):   ##! * ' ( ) ; : @ & = + $ , / ? # [ ]
        return False
    
    return True

def remove_left_blank(ux):
    if ux[:1] == " ":
        ux = ux[1:]
        ux = remove_left_blank(ux)
    return (ux)
        
def urlForm(tag, currentURL, flags):
    # 자체도메인이 아닐경우 처리 방안 필요 print(tag, currentURL)
    returl = tag
    tmp = tag
  
    if tag is None:
        return None
    if len(tag) == 0:
        return None
    if "javascript:void(0)" in tag:
        return None
    if " " in tag:
        return None
    try:
        if not collect_path(urlparse(tag)[2]):
            if tag not in incollect_path_list:
                incollect_path_list.add(tag)
                Debug_w(f"In collect url: {tag} {currentURL}")
            return 
            
    except:
        print("Invalid URL:", tag)
        return None
    # tag = tag.replace(" ", "%20")b
    # tag = tag.replace("?", "%3D")
    tag = urllib.parse.unquote(tag)

    if "=+" in tag: ##urlparse(tag)[4]:
        return None
    if "/+" in tag:
        return None
    if "+/" in tag:
        return None
    if "%(" in tag:
        return
        ##sub05/sub03.php?com_board_basic=read_form&&com_board_idx=89&com_board_id=9  currURL=http://www.magicit.co.kr/default/index.php

    # if tag[-1]=="?":  ## 2021.5.9  ## 2011.11.29 재변경
    #     tag = tag[:-1]
    if flags == 1:
        if not_url_but_parameter(tag):
            return None
    tag = remove_left_blank(tag)   ### 2021.1.31   ## http://www.daemi.net/bbs/ /ch/company/location.php
    if ("http://" in tag or "https://" in tag) and len(tag) <9:
        return None

    tag = html.unescape(tag)  ### 2021.2.24
    
    Debug_w("urlForm-01 " + tag + " currURL=" + currentURL )  
    qp = currentURL.find("?")
    if qp > -1:
        currentURL = currentURL[:qp]   
    
    if tag == ".." :
        tag = "../"
    if tag == ".//":
        tag = "./"
    if tag[:2] == "//":
        tag = http_https + ":" + tag
    if tag[:2] == "./":      ### 2021.1.30
        tag = tag[2:]  
        
    if "moovit://" in tag:
        return None
    if "mailto:" in tag:
        return None
    if "tel:" in tag:
        return None
    if len(tag) == 0:
        return None
    if 'javascript' in tag:
        return None
    if "{{" in tag: ## 20211002
        if "}}" in tag:
            return None
    if "=" in tag:
        if "?" not in tag:
            tag = tag.split(" ")[0]
        # else:
        #     return None

    # if tag[-1:] =="/" and tag[1] != "/":  ###2021.3.1 1 http://www.ubsc.or.kr
    #     tag = "/" + tag

   
    if tag[0] == "?":
        utag = tag.replace("%2F","/")   
        qp = currentURL.find("?")
        if qp > -1:
            tag = currentURL[:qp] + utag        
    
# 2021.8.9
# 144    http://allminwon.com/m/m/main/main/main/main/main/main/main/main/main/main/main/main/main/m/etc/login.html
# 145    http://allminwon.com/m/m/main/main/main/main/main/main/main/main/main/main/main/main/main/m/main/404.php
# 146    http://allminwon.com/m/m/main/main/main/main/main/main/main/main/main/main/main/main/main/m/m/voice_customer.php


    # lp = tag.find("?")    ### 2021.2.4  http://koreapnc21.com
    # np = tag.find("&")
    # if lp > 1 and np < lp:
    #     if len(tag) - lp > 20:
    #         tl = tag[:lp+20]
    #         if visitLinkDict[tl] == 1:  
    #             return None

    tag = tag.strip()
    p_sign = -1    
    #special bootstrap dir + /other/terms/  https://stapharma.com.cn/cn   20211002
    if "dir" in tag:
        p_sign = tag.find("+") 
        if p_sign > 1:
            tag = tag[p_sign+1:].rstrip().lstrip()

    if tag == "./":
        tag = baseUrl   
        returl = tag     
        Debug_w("urlForm-02" + tag)       
    elif -1 < tag.find(baseUrl) <  10:  #if baseUrl in tag:
        #print(1, baseUrl,">>>>>>>>",tag) 
        # ginnohome/main/main.aspx currURL=http://www.g-inno.com
        returl = tag
        Debug_w("urlForm-03 " + "baseUrl=" + baseUrl + " tag= " + tag + " ret=" + returl)
    elif (tag[0] == "/"):     
        if p_sign > -1:
            returl = currentURL + tag
        else:
            returl = get_domain(baseUrl) + tag
        Debug_w("urlForm-04 " + tag + " ret=" + returl + " baseurl=" + baseUrl + "currentURL=" + currentURL)
    elif '../' in tag:
        subCount = tag.count('../')
        # modiURL = currentURL
        # if "?" in modiURL:
        #     modiURL = modiURL[0:modiURL.rfind("?")]  ###  remove "/" that a rightside of ? sss/vvv/ddd?u=/vvv 
        tag = re.sub(r'[.\/][.\/][\\/]',r'', tag)
        #../site_map/sitemap.php currURL=http://www.upcsinfor.com/index.php

        # xurl = http_https + "://" + baseUrl +  "/"
        # xtag = tag
        # modiURL = modiURL[len(xurl):]
        # for i in range(0, subCount+1):
        #     modiURL = modiURL.replace(modiURL.split("/")[len(modiURL.split("/"))-1], "")
        #     modiURL = modiURL[:-1]
        # ../index.php currURL=http://www.sammisound.com/
       
        # if baseUrl not in modiURL:
        #     if len(modiURL)>0:
        #         modiURL = http_https + "://" + baseUrl + "/" + modiURL
        #     else:
        #         modiURL = http_https + "://" + baseUrl

        # returl =  modiURL +'/'+ tag
        #print(">>>>>>>> subcount", subCount)
        returl = get_parentPath(currentURL,subCount)
        #print("<<<<<<<<<<<<", returl, currentURL)
        if tag in returl and "." in tag:    ## 2020.4.19
            return None
        if returl[-1] == "/":
            returl = returl + tag
        else:
            returl = returl + "/" + tag
        Debug_w("urlForm-05 " + tag + " ret=" + returl + " currURL=" + currentURL )

    elif "./" == tag[:2]:
        di = currentURL.rfind("/")+1
        if di > 8:
            returl = currentURL[:currentURL.rfind("/")+1] + tag[2:]
        else:
            returl = currentURL + tag[1:]
        Debug_w("urlForm-06 " + tag + "baseUrl=" + baseUrl + " ret=" + returl + " currURL=" + currentURL )
    elif "#" == tag[0]:
        sp = currentURL.find("#")
        if sp > -1:
            returl = currentURL[:sp] + tag
        else:
            returl = currentURL + tag
        Debug_w("urlForm-77 " + tag + " ret=" + returl )
    else:
        # #return None #This will format the urls for recursive purpose
        if Redirection_level == 1:
            if "http" in tag.lower():
                returl =  tag
                Debug_w("urlForm-08 " + tag + " ret=" + returl )
            else:
                if baseUrl2=="":
                    returl = http_https + "://" + baseUrl + "/" + tag
                else:
                    returl = http_https + "://" + baseUrl2 + "/" + tag                
                Debug_w("urlForm-09" + tag + " ret=" + returl + " baseUrl2=" + baseUrl2 )
        else:
            if "http" in tag.lower():
                returl = tag
                Debug_w("urlForm-10 " + tag + " ret=" + returl )
            else:
                # if currentURL.rfind("/") > 7:
                #     returl = currentURL[:currentURL.rfind("/")+1]+ tag
                # else:
                # 2024.8.25 chnaged 
                #urlForm-11 ?mode=view&article_no=137017&board_wrapper=/ko/state/info03.jsp&pager.offset=100&board_no=6&default:category_id=165 ret=http://www.dtaq.re.kr/ko/state/?mode=view&article_no=137017&board_wrapper=/ko/state/info03.jsp&pager.offset=100&board_no=6&default:category_id=165 currURL=http://www.dtaq.re.kr/ko/state/info03.jsp
                if tag[0]=="?":
                    returl = currentURL + tag
                else:
                    di = currentURL.rfind("/")+1
                    if di > 8 and tag[:2] !="./":
                        returl = currentURL[:di] + tag
                    elif tag[0]!='/':
                        returl = currentURL + "/" + tag
                Debug_w("urlForm-11 " + tag + " ret=" + returl + " currURL=" + currentURL)
    
     ###2024.07.30
    ### http://www.energy-news.co.kr) ,  http://www.iso.org)를 
    if ")" in returl:
        returl = slice_before_character(returl, ")")
        if ";" in returl:
            return None

    Debug_w("urlForm-12 " + tag + " ret=" + returl + " currURL=" + currentURL)
    try:
        returl = urllib.parse.urldefrag(returl)[0]
    except:
        returl = returl
        Debug_w("urllib.parse.urldefrag"+ returl)
    return (returl)

def getheadInformation(url):
    print(url)
    try:
        rhead = HTTP_SESSION.head(
            url, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
        )
        print(rhead)
        ##rhead.status_code == 307:
        #print("++++++++",rhead,type(rhead))
        #return(rhead.headers)    
        return(rhead)    
    except:
        return(None)

def identify_website_builder_legacy(soup):
    return identify_website_builder(soup)
    stext = str(soup)
    scripts = soup.find_all('script')
    links = soup.find_all('link')
    meta_tags = soup.find_all('meta')

    for meta in meta_tags:
        if meta.get('name') == 'generator':
            if 'Hostinger' in meta.get('content', ''):
                return "Hostinger"
            elif 'Shopify' in meta.get('content', ''):
                return "Shopify"
            elif 'Webflow' in meta.get('content', ''):
                return "Webflow"
            elif 'Square' in meta.get('content', ''):
                return "Square"
            elif 'Duda' in meta.get('content', ''):
                return "Duda"
            elif 'GoDaddy' in meta.get('content', ''):
                return "GoDaddy"
            elif 'Jimdo' in meta.get('content', ''):
                return "Jimdo"
            elif 'SITE123' in meta.get('content', ''):
                return "SITE123"
            elif 'Webador' in meta.get('content', ''):
                return "Webador"
            elif 'IONOS' in meta.get('content', ''):
                return "IONOS"
            elif 'BigCommerce' in meta.get('content', ''):
                return "BigCommerce"
            elif 'site123' in meta.get('content', ''):
                return "SITE123"
        elif meta.get('property') == 'og:url' :
            if meta.get('content') == 'https://coderium-studio.com/':
                return "CODERIUM"
    
    for script in scripts:
        if script.get('src'):
            # 일반적으로 Vite 번들 파일에서 특정 패턴을 확인
            if re.search(r'/assets/index\.[a-f0-9]{8}\.js$', script['src']):
                return "Vite"
            # Vite가 생성한 소스맵 확인
            elif re.search(r'/assets/index\.[a-f0-9]{8}\.js\.map$', script['src']):
                return "Vite"
            elif 'hostinger' in script['src']:
                return "Hostinger"
            elif 'cdn.shopify.com' in script['src']:
                return "Shopify"
            elif  'webflow.com' in script['src']:
                return "Webflow"
            elif  'dudamobile.com' in script['src']:
                return "Duda"
            elif 'godaddy.com' in script['src']:
                return "GoDaddy"
            elif 'jimdo.com' in script['src']:
                return "Jimdo"
            elif 'site123.com' in script['src'] or 'static.s123-cdn-network-a.com' in script['src']:
                return "SITE123"
            elif 'webador.com' in script['src']:
                return "Webador"
            elif 'ionos.com' in script['src']:
                return "IONOS"
            elif 'ionos.com' in script['src']:
                return "IONOS"
            elif 'web.com' in script['src']:
                return "Web.com"
            elif 'bigcommerce.com' in script['src']:
                return "BigCommerce"
    
    for link in links:
        if link.get('href'):
            if  'https://websites-api.hostinger.com' in link['href']:
                return "Hostinger"
            elif link.get('rel') == ['preconnect'] and 'https://www.googletagmanager.com' in link['href']:
                return "Hostinger"
            elif 'cdn.shopify.com' in link['href']:
                return "Shopify"
            elif 'webflow.com' in link['href']:
                return "Webflow"
            elif 'squareup.com' in link['href']:
                return "Square"
            elif 'squareup.com' in link['href']:
                return "Square"
            elif  'dudamobile.com' in link['href']:
                return "Duda"
            elif 'godaddy.com' in link['href']:
                return "GoDaddy"
            elif 'jimdo.com' in link['href']:
                return "Jimdo"
            elif 'site123.me' in link['href'] or 'static.s123-cdn-network-a.com' in link['href']:
                return "SITE123"
            elif 'webador.com' in link['href']:
                return "Webador"
            elif 'ionos.com' in link['href']:
                return "IONOS"
            elif 'bigcommerce.com' in link['href']:
                return "BigCommerce"
   

    # HTML에서 특정 패턴 검색   <meta name="generator" content="Wix.com Website Builder"/>
    if soup.find('meta', attrs={'name': 'generator', 'content': 'WordPress'}):
        return "WordPress"
    elif soup.find('meta', attrs={'name': 'generator', 'content': 'Wix.com Website Builder'}):
        return "Wix"
    elif 'wp-content' in stext or 'wp-includes' in stext:
        return "WordPress"
    elif 'Wix.com' in stext:
        return "Wix"
    elif 'squarespace.com' in stext:
        return "Squarespace"
    else:
        return "Unknown Builder"

def identify_website_builder(soup):
    if not soup:
        return "Custom/Static Site"

    stext = str(soup)
    stext_lower = stext.lower()
    scripts = soup.find_all("script") if hasattr(soup, "find_all") else []
    links = soup.find_all("link") if hasattr(soup, "find_all") else []
    meta_tags = soup.find_all("meta") if hasattr(soup, "find_all") else []

    meta_patterns = [
        ("wordpress", "WordPress"),
        ("wix.com website builder", "Wix"),
        ("wix", "Wix"),
        ("shopify", "Shopify"),
        ("webflow", "Webflow"),
        ("squarespace", "Squarespace"),
        ("hostinger", "Hostinger"),
        ("square", "Square"),
        ("duda", "Duda"),
        ("godaddy", "GoDaddy"),
        ("jimdo", "Jimdo"),
        ("site123", "SITE123"),
        ("webador", "Webador"),
        ("ionos", "IONOS"),
        ("bigcommerce", "BigCommerce"),
        ("cafe24", "Cafe24"),
        ("gnuboard", "Gnuboard"),
        ("rhymix", "Rhymix"),
        ("xpressengine", "XpressEngine"),
        ("xe", "XpressEngine"),
        ("imweb", "Imweb"),
        ("sixshop", "Sixshop"),
        ("makeshop", "Makeshop"),
        ("godo", "Godo Mall"),
        ("joomla", "Joomla"),
        ("drupal", "Drupal"),
        ("magento", "Magento"),
        ("opencart", "OpenCart"),
        ("tistory", "Tistory"),
        ("modoo", "Naver Modoo"),
        ("gatsby", "Gatsby"),
        ("docusaurus", "Docusaurus"),
        ("hugo", "Hugo"),
        ("jekyll", "Jekyll"),
        ("elementor", "Elementor"),
        ("ghost", "Ghost"),
        ("hubspot", "HubSpot CMS"),
        ("blogger", "Blogger"),
        ("framer", "Framer"),
        ("tilda", "Tilda"),
        ("carrd", "Carrd"),
        ("readymag", "Readymag"),
        ("notion", "Notion/Super"),
        ("typedream", "Typedream"),
        ("unbounce", "Unbounce"),
        ("instapage", "Instapage"),
        ("landingi", "Landingi"),
        ("shopby", "Shopby"),
        ("firstmall", "Gabia Firstmall"),
        ("wisa", "WISA"),
    ]
    url_patterns = [
        ("elementor/assets", "Elementor"),
        ("elementor-frontend", "Elementor"),
        ("wp-content/plugins/elementor", "Elementor"),
        ("shopby.co.kr", "Shopby"),
        ("shopby.cloud", "Shopby"),
        ("wisa.co.kr", "WISA"),
        ("wisaimg.co.kr", "WISA"),
        ("wisaimg.com", "WISA"),
        ("wp-content", "WordPress"),
        ("wp-includes", "WordPress"),
        ("static.parastorage.com", "Wix"),
        ("wixstatic.com", "Wix"),
        ("wix.com", "Wix"),
        ("cdn.shopify.com", "Shopify"),
        ("myshopify.com", "Shopify"),
        ("webflow.com", "Webflow"),
        ("assets-global.website-files.com", "Webflow"),
        ("squarespace.com", "Squarespace"),
        ("static1.squarespace.com", "Squarespace"),
        ("dudamobile.com", "Duda"),
        ("godaddy.com", "GoDaddy"),
        ("jimdo.com", "Jimdo"),
        ("site123.com", "SITE123"),
        ("site123.me", "SITE123"),
        ("static.s123-cdn-network-a.com", "SITE123"),
        ("webador.com", "Webador"),
        ("ionos.com", "IONOS"),
        ("bigcommerce.com", "BigCommerce"),
        ("hostinger", "Hostinger"),
        ("cafe24.com", "Cafe24"),
        ("cafe24.co.kr", "Cafe24"),
        ("ecimg.cafe24img.com", "Cafe24"),
        ("img.echosting.cafe24.com", "Cafe24"),
        ("gnuboard", "Gnuboard"),
        ("/bbs/", "Gnuboard"),
        ("/theme/basic", "Gnuboard"),
        ("rhymix", "Rhymix"),
        ("xpressengine", "XpressEngine"),
        ("/modules/", "XpressEngine"),
        ("imweb.me", "Imweb"),
        ("cdn.imweb.me", "Imweb"),
        ("sixshop.com", "Sixshop"),
        ("makeshop", "Makeshop"),
        ("godo.co.kr", "Godo Mall"),
        ("godomall", "Godo Mall"),
        ("/media/system/js/", "Joomla"),
        ("/templates/system/", "Joomla"),
        ("content=\"joomla!", "Joomla"),
        ("/sites/default/files/", "Drupal"),
        ("/core/misc/drupal", "Drupal"),
        ("drupal-settings-json", "Drupal"),
        ("/static/version", "Magento"),
        ("mage/cookies", "Magento"),
        ("magento", "Magento"),
        ("catalog/view/theme", "OpenCart"),
        ("route=product/", "OpenCart"),
        ("opencart", "OpenCart"),
        ("tistory.com", "Tistory"),
        ("tistory1.daumcdn.net", "Tistory"),
        ("k.kakaocdn.net/dn/", "Tistory"),
        ("modoo.at", "Naver Modoo"),
        ("naver.me", "Naver Modoo"),
        ("cdn.modoo.at", "Naver Modoo"),
        ("bootstrap.min.css", "Bootstrap (frontend framework)"),
        ("bootstrap.bundle.min.js", "Bootstrap (frontend framework)"),
        ("react-dom", "React (frontend library)"),
        ("data-reactroot", "React (frontend library)"),
        ("vue.runtime", "Vue (frontend framework)"),
        ("data-v-", "Vue (frontend framework)"),
        ("ng-version", "Angular (frontend framework)"),
        ("svelte", "Svelte (frontend framework)"),
        ("/page-data/", "Gatsby"),
        ("gatsby", "Gatsby"),
        ("docusaurus", "Docusaurus"),
        ("hugo", "Hugo"),
        ("jekyll", "Jekyll"),
        ("/ghost/content/", "Ghost"),
        ("ghost/content", "Ghost"),
        ("static.ghost.org", "Ghost"),
        ("static.hsappstatic.net", "HubSpot CMS"),
        ("hs-scripts.com", "HubSpot CMS"),
        ("hubspot", "HubSpot CMS"),
        ("blogger.com", "Blogger"),
        ("blogblog.com", "Blogger"),
        ("framerusercontent.com", "Framer"),
        ("framer.com/m/", "Framer"),
        ("static.tildacdn.com", "Tilda"),
        ("tilda.cc", "Tilda"),
        ("assets.carrd.co", "Carrd"),
        ("carrd.co/assets", "Carrd"),
        ("readymag.com", "Readymag"),
        ("notion.site", "Notion/Super"),
        ("super.so", "Notion/Super"),
        ("typedream.com", "Typedream"),
        ("unbouncepages.com", "Unbounce"),
        ("instapage.com", "Instapage"),
        ("landingi.com", "Landingi"),
        ("shopby", "Shopby"),
        ("firstmall.kr", "Gabia Firstmall"),
        ("firstmall.co.kr", "Gabia Firstmall"),
        ("gabia.com", "Gabia Firstmall"),
    ]
    framework_patterns = [
        (r"/assets/index\.[a-f0-9]{8,}\.js(?:\.map)?$", "Vite (frontend build tool)"),
        (r"/_next/static/", "Next.js (frontend framework)"),
        (r"/_nuxt/", "Nuxt (frontend framework)"),
        (r"/assets/.*\.[a-f0-9]{8,}\.(?:js|css)$", "Vite (frontend build tool)"),
        (r"__next_data__", "Next.js (frontend framework)"),
        (r"window\.__nuxt__", "Nuxt (frontend framework)"),
        (r"webpackjsonp", "Webpack (frontend build tool)"),
    ]

    for meta in meta_tags:
        name = (meta.get("name") or "").strip().lower()
        prop = (meta.get("property") or "").strip().lower()
        content = (meta.get("content") or "").strip()
        content_lower = content.lower()
        if name == "generator":
            for pattern, builder in meta_patterns:
                if pattern in content_lower:
                    return builder
        if prop == "og:url" and content == "https://coderium-studio.com/":
            return "CODERIUM"

    asset_values = []
    for script in scripts:
        src = script.get("src") or ""
        if src:
            asset_values.append(src)
        inline_script = script.get_text(" ", strip=True)
        if inline_script:
            asset_values.append(inline_script[:5000])
    for link in links:
        href = link.get("href") or ""
        if href:
            asset_values.append(href)

    for value in asset_values:
        value_lower = value.lower()
        for pattern, builder in url_patterns:
            if pattern in value_lower:
                return builder
        for pattern, builder in framework_patterns:
            if re.search(pattern, value_lower):
                return builder

    for pattern, builder in url_patterns:
        if pattern in stext_lower:
            return builder
    for pattern, builder in framework_patterns:
        if re.search(pattern, stext_lower):
            return builder

    return "Custom/Static Site"

# import re

# def identify_website_builder(soup):
#     # 주요 키워드와 빌더를 매핑하는 딕셔너리
#     meta_generators = {
#         "Hostinger": "Hostinger",
#         "Shopify": "Shopify",
#         "Webflow": "Webflow",
#         "Square": "Square",
#         "Duda": "Duda",
#         "GoDaddy": "GoDaddy",
#         "Jimdo": "Jimdo",
#         "SITE123": "SITE123",
#         "Webador": "Webador",
#         "IONOS": "IONOS",
#         "BigCommerce": "BigCommerce",
#         "Wix.com Website Builder": "Wix",
#         "WordPress": "WordPress"
#     }

#     script_patterns = [
#         (r'/assets/index\.[a-f0-9]{8}\.js$', "Vite"),
#         (r'/assets/index\.[a-f0-9]{8}\.js\.map$', "Vite"),
#         ("hostinger", "Hostinger"),
#         ("cdn.shopify.com", "Shopify"),
#         ("webflow.com", "Webflow"),
#         ("dudamobile.com", "Duda"),
#         ("godaddy.com", "GoDaddy"),
#         ("jimdo.com", "Jimdo"),
#         ("site123.com", "SITE123"),
#         ("static.s123-cdn-network-a.com", "SITE123"),
#         ("webador.com", "Webador"),
#         ("ionos.com", "IONOS"),
#         ("web.com", "Web.com"),
#         ("bigcommerce.com", "BigCommerce"),
#     ]

#     link_patterns = [
#         ("https://websites-api.hostinger.com", "Hostinger"),
#         ("https://www.googletagmanager.com", "Hostinger"),
#         ("cdn.shopify.com", "Shopify"),
#         ("webflow.com", "Webflow"),
#         ("squareup.com", "Square"),
#         ("dudamobile.com", "Duda"),
#         ("godaddy.com", "GoDaddy"),
#         ("jimdo.com", "Jimdo"),
#         ("site123.me", "SITE123"),
#         ("static.s123-cdn-network-a.com", "SITE123"),
#         ("webador.com", "Webador"),
#         ("ionos.com", "IONOS"),
#         ("bigcommerce.com", "BigCommerce"),
#     ]


#     print(soup)
#     print("***********************************************")
#     # Step 1: Meta 태그 확인
#     meta_tags = soup.find_all('meta')
#     for meta in meta_tags:
#         # "generator" 속성 확인
#         if meta.get('name') == 'generator':
#             content = meta.get('content', '').lower()
#             for key, builder in meta_generators.items():
#                 if key.lower() in content:
#                     return builder
#         # "og:url" 속성 확인
#         if meta.get('property') == 'og:url' and meta.get('content') == 'https://coderium-studio.com/':
#             return "CODERIUM"

#     # Step 2: Script 태그 확인
#     scripts = soup.find_all('script')
#     for script in scripts:
#         src = script.get('src', '')
#         for pattern, builder in script_patterns:
#             if re.search(pattern, src):
#                 return builder

#     # Step 3: Link 태그 확인
#     links = soup.find_all('link')
#     for link in links:
#         print(link)
#         href = link.get('href', '')
#         rel = link.get('rel', [])
#         for pattern, builder in link_patterns:
#             if pattern in href:
#                 return builder
#             if rel == ['preconnect'] and 'https://www.googletagmanager.com' in href:
#                 return "Hostinger"

#     # Step 4: HTML 텍스트 검색
#     stext = str(soup)
#     if 'wp-content' in stext or 'wp-includes' in stext:
#         return "WordPress"
#     elif 'Wix.com' in stext:
#         return "Wix"
#     elif 'squarespace.com' in stext:
#         return "Squarespace"

#     # Unknown builder
#     return "Unknown Builder"

def getipInformation(ip):
    urls = "http://ip-api.com/json/" + ip
    try:
        rhead = HTTP_SESSION.get(urls, timeout=REQUEST_TIMEOUT)
        rhead.raise_for_status()
        return rhead.text
    except requests.RequestException:
        return ""

def getMyip():
    urls = "https://api.ipify.org"
    try:
        response = HTTP_SESSION.get(
            urls, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
        )
        response.raise_for_status()
        return response.text.strip()
    except requests.RequestException:
        return ""

def getsitemapInformation(url):
    root = normalize_site_root(url)
    return fetch_text(urljoin(root + "/", "sitemap.xml"))

def getsitemap2Information(url):
    root = normalize_site_root(url)
    return fetch_text(urljoin(root + "/", "sitemap.html"))

def googleSearch(query):
#     requests.headers = {
#     'Accept-Language':'en',
#     'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:57.0) Gecko/20100101 Firefox/57.0'
# }
    surl = 'https://www.google.com/search?q="' + query + '"&hl=en'
    resultStats = "xx"
    try:
        resp = HTTP_SESSION.get(
            surl, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
        )
        _, content = parse_html_response(resp)
        ##print(content)
        ##s  =  content.find('div', {'id' : 'resultStats'})   ##About 74 results
        s = str(content)
        x1 = s.find("About")
        x2 = s.find("results")
        if x2 > x1 and (x2 - x1 < 35): 
            resultStats = s[x1:x2+7]
        else:
            resultStats = ""
    except :
        pass
    return (resultStats)    


def getaCompareSite(siteflag):    
    requests.headers = {
    'Accept-Language':'en',
    #'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:57.0) Gecko/20100101 Firefox/57.0'
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
}
    url_response_time = -1
    url_contents_length = -1

    urls = "https://www.naver.com" 
    if siteflag=="us":
        urls = "https://about.google.com"
    elif siteflag=="kr":
        urls = "https://www.naver.com"
    try:
        resp = HTTP_SESSION.get(
            urls, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
        )
        url_response_time = resp.elapsed.total_seconds()
        url_contents_length = len(resp.content)
    except:
        url_response_time = -1
        url_contents_length = -1

    return (url_response_time, url_contents_length)

def abnormal_url(u):
    s = u.split("?")
    if len(s)==2:
        if "LOGIN" in s[0].upper():
            if s[1].find("&")==-1:
                return True
    return False

def modi_url(inurl):
    ctag=""
    try:
        if baseUrl in inurl:
            ctag =  inurl.split(baseUrl)[1]
            x = ctag.split("?")
            ctag = x[0]
        else:                   ### 2019.8.2 
            ctag = inurl        ### handling other domain, because the website has frameset and forwaring other site
    except:
        print(baseUrl, "<<<>>>", inurl)
        #ctag = baseUrl
    return ctag       

def not_valid(t):   ### 20240730, url 등과 같은 함수내 prameter 사용시
    v = 0
    if ":" not in t:
        v = v + 1
    if "/" not in t:
        v = v + 1
    if "." not in t:
        v = v + 1
    if len(t) < 6:
        v = v + 1
    if v >= 3:
        return True
    else:
        return False

def addgraphNode(tag,url):

    if tag is None:
        print("None=====")

    if (tag in G) == False:
        G.add_node(tag)
        ctag =  modi_url(tag)
        H.add_node(ctag)
        visitLinkDict[tag] = 0    ###  0>1>0 2019.11.2
       
    if (G.has_edge(url, tag)) == False:
        G.add_edge(url, tag)
        ctag = modi_url(tag)
        curl = modi_url(url)
        H.add_edge(curl,ctag)   

def check_iframe_location(url):
    try:
        Debug_w(url + " check iframe location")
        page = HTTP_SESSION.get(
            url, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
        )
    except Exception as exception:
        print("incomplete content %s" % url)
        print(exception)
        return url
    except (http.client.IncompleteRead) as e:
        print("Incomplete link %s" % url)
        return url
    except requests.exceptions.MissingSchema as err:
        print("MissingSchema",err,url)
        return url
    except requests.exceptions.InvalidSchema as err:
        print("InvalidSchema",err,url)
        return url
    except requests.exceptions.RequestException as err:
        print ("OOps: Something Else 1999",err)
        return url
    except requests.exceptions.HTTPError as err:
        print ("Http Error:",err)
        return url
    except requests.exceptions.ConnectionError as err:
        print ("Error Connecting:",err)
        return url
    except requests.exceptions.Timeout as err:
        print ("Timeout Error:",err)  
        return url
    except requests.exceptions.RequestException as err:
        print ("OOps: Something Else 2011",err)
        return url
    #######################################################
    Debug_w(url + " >>> frame tag")        
    _, content = parse_html_response(page)
    Debug_w(content)
    #######################################################
    Debug_w(str(content.find_all('frame')))
    for frameset in content.find_all('frame'):
        tag = frameset.get("src")   
        if tag=="" or tag == None:  ## include or None  2020.2.2
            continue 
        if (baseUrl not in tag) and tag[:4].lower()=="http":
            print("Rediection level changed:", tag, "   ", baseUrl, "===Level=", Redirection_level)         
            url = urllib.parse.unquote(tag)    
            p_url = urlparse(url)
            baseUrlX = p_url[1] + p_url[2]
            return baseUrlX
    #######################################################
    Debug_w(url + " >>> meta tag, reflash")      
    ###############2021.10.26########################################
    for meta in content.find_all('meta'):
        if meta.get('http-equiv') in ["refresh", "Refresh", "REFRESH"]:   
            mc = meta.get('content')
            if mc.upper().find('URL') > 0:
                print("2049***************************",mc)
                tag = mc.split("=")[1]
                tag=tag.replace("'","")
                tag=tag.replace('"','')
                tag = urlForm(tag, url, 0)
                if get_domain(baseUrl) in tag:
                    continue
                else:
                    return (tag)    
    #########################################
    # window.location.replace('http://www.lumensleds.com');
    #########################################
    for slink in content.find_all("script"):
        temp= slink.text
        Debug_w(f"script redirect check: {temp[:300]}")
        x = temp.find("window.location.replace")
        if x > 0:
            temp = temp[x+23:]
            lb = temp.find("(")
            rb = temp.find(")")
            if rb > lb:
                temp = temp[lb+1:rb].replace("'","").replace('"','')
                return (temp)
        x = temp.find("location.href")        
        if x > 0:
            temp = get_locationhref(temp,0)   ## <script>location.href="/new/kor/main/main.php"</script>
            if temp==None:
                continue
            temp = urlForm(temp, url, 0)
            return(temp)
 #########################################
    return url    

def addScanList(t,f,u,s):
    global scanWebSet
    global scanWebList
    if t[-1]=="?":
        t=t[:-1]

    bl = len(scanWebSet)
    scanWebSet.add(t)
    al = len(scanWebSet)
    if al > bl:       
        scanWebList.append([t,f,u])

def should_skip_scheme(url):
    return (
        url[:7] == "moovit:"
        or url[:7] == "mailto:"
        or url[:4] == "tel:"
    )

def handle_external_link(url, source_url):
    if relation_domain(url, baseUrl) != -1:
        return False
    xUrl = get_domain(url)
    record_sns_link(url, source_url)
    if url.find("blog") > -1:
        Debug_w(f"{url} {xUrl}")
    Debug_w(f"skip other domain: {url} {xUrl} {baseUrl} {get_domain(baseUrl)} {get_domain(url)}")
    return True

def scanWeb(url, Purl):

    global baseUrl
    global RestricedOnlineLinkCount
    global online_score
    global counter
    global RestricedRestCycleCount
    global font_color_count

    if "index" in url:
        Debug_w(f"urlurlurl: {url}")

    if url is None:
        return
    if online_score > 500 and counter> 1000:
        print("online_score exceed !")
        return
 
    if should_skip_scheme(url):
        return

    ## 2021.9.5 
    if multiple_url(url):   
        return        
    
    #2024.8.5
    if abnormal_url(url):
        return
    
    if url[:1]==".":
        ##
        #url = modi_url(Purl) + url[1:]   ## 2019.10.20
        url = f'{modi_url(Purl)}{url[1:]}' ## 2022.01.12
        print("00000", Purl, url)
    elif url[:1]=="/":
        #url = baseUrl + url
        url = f'{baseUrl}{url}' ## 2022.01.12
        
    ##332021.5.2 check
    if yearUrlskip == True:  ##isnumeric()  
        if url[-4:].isnumeric():
            if url[-4:] not in year_list:
                return
   
    try:
        url = urllib.parse.unquote(url)    
        p_url = urlparse(url)
        if len(p_url[0]) <2:
            #v_url = "http://" + p_url[1] + p_url[2]
            v_url = f'http://{p_url[1]}{p_url[2]}' ## 2022.01.12
        else:
            #v_url = p_url[0] + "://" + p_url[1] + p_url[2]
            v_url = f'{p_url[0]}://{p_url[1]}{p_url[2]}' ## 2022.01.12

        ### onl;ine path restricted 2022.2.6 
        s_url=""
        if online_score > 100:            
            s_url = p_url[2].split("/")
            if len(s_url) > 4:
                #o_url = str(s_url[1]) + str(s_url[2]) + str(s_url[3]) 
                o_url = f'{str(s_url[1])}{str(s_url[2])}{str(s_url[3])}'
                onlineLink[o_url] = onlineLink.get(o_url,0) + 1
                if onlineLink[o_url] > RestricedOnlineLinkCount:
                    return
    except:
        print("Unexpected error:", sys.exc_info()[0], url, s_url)
        return
    
    if cmp_domain(p_url[1], baseUrl) < 2:
        return
    
    try:
        visitLinkDict[url] = visitLinkDict.get(url,0) + 1
        visitUrlDict[v_url] = visitUrlDict.get(v_url,0) + 1
        if visitUrlDict[v_url] > RestricedRestCycleCount:
            Debug_w("skip: already scaned... " + v_url)
            return
        if visitLinkDict[url] > 1:  
            Debug_w("scanded..." + url)
            return
    except:
        print("794 error", url)   # goetheunibator.dehttp://dev.goetheunibator.de/wordpress/apply1
        #print("except===", visitLinkDict)    
        #print(url)   # goetheunibator.dehttp://dev.goetheunibator.de/wordpress/apply1   
        return
    if url.find("blog") > -1:
        print ("?????????????????????????????????????????????",url)

    #print(baseUrl, url, "<<<<<<")
    try:
        if handle_external_link(url, Purl):  ## 2021.08.22
            return
    except:
        print("805 : url=", url)
        print("Unexpected error:", sys.exc_info()[0])
        print("baseUrl=", baseUrl)
    
    counter = counter + 1
    html_string = ""
    ####content = ""
    if print_toggle == 1 :
        print("%-6d %s %-5d %s" % (counter, "os:", online_score, url))    

    ####################################
    # stop line debug
    #################################
    if StopLine > 0 and StopLine < counter:
        sys.exit()
    ###########################################################
    skip_Current_url=False
    ###########################################################
    # calculation time reduce routine, 
    # if customer want full scan then remove it. 2019.7.24
    Debug_w("get_extion: "+url)
    if "." in url:
        ext_name = get_extion(url)
        if ext_name != "*":
            extion_count[ext_name] = extion_count.get(ext_name,0) + 1  
            if costSavingMode == True:
                if skip_extion(ext_name) == True:      
                    print("skip ",url)
                    skip_Current_url = True
                    return
    ###########################################################
    try:        
        if "http://" == url[:7].lower():
            url = url
        elif "https://" == url[:8].lower():
            url = url
        else:
            #url = http_https + "://" + url   
            url = f'{http_https}://{url}'  # 2022.1.12

        validate_public_url(url)
        if not robots_allows(url):
            print("skip by robots.txt:", url)
            return

        Debug_w("2427-requests.head"+url)
        rhead = HTTP_SESSION.head(
            url, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
        )
        if rhead.status_code in [302]:
            if "Location" in rhead.headers:
                purl = rhead.headers["Location"]    

                print(purl)

                if "http" not in purl:
                    if purl[0]=='/':
                        purl = f'{baseUrl}{purl}'
                    else:
                        purl = f'{baseUrl}/{purl}'
                    if "http" not in purl:
                        purl = f'{http_https}://{purl}'
   
                if purl in list_302:
                    return
                else:
                    list_302.add(purl)
                    url = purl
                    print(2237, "change url", url)
        else:
            if costSavingMode == True:
                if skip_costsavingMode(rhead.headers):    
                    skip_Current_url = True
                    print("1 cs mode =", url)
                    return
    except:
         print("Unexpected error:", sys.exc_info()[0], sys.exc_info()[1])
        #######################################################
        #        if rhead.status_code in [302]:
        #   print("302===>")
        #   purl = rhead.headers["Location"]                    
        #   print(purl)
        #######################################################
        # target(url) page read
        #######################################################
    try:
        if skip_Current_url == True:
            page = ""
            url_response_time = -1
            url_contents_length = -1
        else:
            Debug_w(url + " begin to read for head")

            #'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:57.0) Gecko/20100101 Firefox/57.0'
            page = HTTP_SESSION.get(
                url, timeout=REQUEST_TIMEOUT, verify=VERIFY_TLS
            )
            page_text = decode_response_text(page)
            
            if page.status_code in [400,404,403,408,409,501,502,503]:
                url_response_time = -1
                url_contents_length = -1
                bstr = url + " [" + str(page.status_code) + "] in " + Purl
                brokenLink.append(bstr)
                return
            else:
                url_response_time = page.elapsed.total_seconds()
                url_contents_length = len(page.content)
            #print(page, page.content)

            if skip_costsavingMode(page.headers):    
                skip_Current_url = True
                print("2 cs mode =", url)
                return

            Debug_w(url + " end of read for head")     
 
        #######################################################
        # update performance data for each page
        #######################################################        
        ut = []
        ut.append(url)
        ut.append(url_contents_length)
        ut.append(url_response_time)
        urltimelist.append(ut)      

    except (http.client.IncompleteRead) as e:
        print("Incomplete link %s" % url)
        return
    except requests.exceptions.MissingSchema as err:
        print("MissingSchema",err,url)
        return
    except requests.exceptions.InvalidSchema as err:
        print("InvalidSchema",err,url)
        return
    except requests.exceptions.RequestException as err:
        print (url, "OOps: Something Else 2304",err)
        return
    except requests.exceptions.HTTPError as err:
        print ("Http Error:",err)
        return
    except requests.exceptions.ConnectionError as err:
        print ("Error Connecting:",err)
        return
    except requests.exceptions.Timeout as err:
        print ("Timeout Error:",err)  
        return
    except requests.exceptions.RequestException as err:
        print (url,"OOps: Something Else 2316",err)
        return
    except :
        print("incomplete link %s" % url)    ##urllib3.exceptions.LocationParseError
        brokenLink.append(url)
        print ("Others Error:")  
        return

    #html_string = str(page.content)  ##20241020
    #print(page.content)
    content = bs(page_text, "html.parser")
    #print("&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&")
   
    word_html_string = extract_text(content)
    Debug_w(url + " begin to read for body")
    if skip_Current_url == True:
        return
    # try:    ### 20241129 dupilicate code
    #     x = page.content
    #     content = bs(x, 'html.parser', from_encoding="utf-8")  ##//decode(“utf8”).iso-8859-1        
    # except Exception as exception:
    # # Output unexpected Exceptions.
    #     print("incomplete content %s" % url)
    #     return
    ######################################################
    # basic html varidation check 2722
     ######################################################
    tags = extract_clean_tags(page_text, html_tag_list)

    for tag in tags:
        if "/" in tag:
            tag = tag.replace("/","")
            end_tag[tag] = end_tag.get(tag, 0) + 1
        else:
            start_tag[tag] = start_tag.get(tag, 0) + 1

   # cmp_tags = merge_dicts_as_tuples(start_tag, end_tag)  
    ######################################################
    global list_html5_tag
    list_html5_tag = merge_dicts(list_html5_tag, check_html5_elements(content))  ## 2024.07.21
    ######################################################
    global web_builder
    if web_builder =="0":
       # print(content)
        print("web site builder check...")
        web_builder = identify_website_builder(content)
    #######################################################
    # select of function in page
    #######################################################
    Debug_w(url + " begin to findall-script")
    #print("(0)",content)
    for slink in content.find_all("script"):
        #extract_function(str(slink), url)
        print("2746.&&&&&&&&&&&&&&&&&&&&&&", slink)
        extract_script_file(slink, url)
        str_slink=str(slink)
        if ".load" in str_slink:  ## 2021.1.24
            html_temp_list = set_html_from_load(str_slink)
            for temp_list in html_temp_list:
                if " " in temp_list:
                    continue
                print("2160>>>>>>>>>>>>>>>>>>>>>>", temp_list)
                tag = urlForm(temp_list, url, 0)
                if tag is not None:
                    addgraphNode(tag,url)
                    ##scanWeb(tag, url)
                    print("2149", tag, url)
                    addScanList(tag,1,url,227)
    Debug_w(url + " end of findall-script")
    #print("1794", type(content))

    try:
        html_string = str(content)  
    except ValueError as m:
        print(m)
    finally:
        pass

    if counter == 1:
        first_page_css = "\n".join(
            style.get_text(" ", strip=True)
            for style in content.find_all("style")
        )
        first_page_css += "\n" + "\n".join(
            tag.get("style", "")
            for tag in content.find_all(style=True)
        )
        first_page_css += "\n" + "\n".join(
            f"color:{tag.get('color')};"
            for tag in content.find_all(color=True)
        )
        merge_color_counts(
            font_color_count,
            extract_font_colors(first_page_css),
        )
    
    # Debug_w(url + "finished convert string from bs4")

    Debug_w(url + "swf checking")
    if html_string.find(".swf"):
        xx=html_string.count(".swf")
        if xx > 0:
            for x in range(1, xx):
                list_flash.append(url)

    if len(Search_string) > 1:    #if check_string(ct,
        # if html_string.find(Search_string) > 1:
        #     print (Search_string , " in ", url) 
        #     list_search.add(url)
        scheck = check_stringAll(html_string.upper(), list_Search_string)
        if len(scheck) > 0:
            tmp=[]
            tmp.append(url)
            tmp.append(scheck)
            #list_search.add(url)
            #print(scheck)
            list_search.append(tmp)

    #######################################################
    Debug_w(url + " begin to parsing for font")
    global tmp_cccc
    tmp_cccc = tmp_cccc | parsing_fontlist(html_string)
    Debug_w(url + " end of the font parcing procedure")
    #######################################################
    #print(content)
    #print(html_string)

    global cssCount
    global cssCount2
    global scriptCount    
    global styleCount
    global AddFavoriteCount
    Debug_w(url + " >>>title, head")
    try:
        tmp_list = []
        head_tag = content.find('head')  
        if head_tag != None:
            # for meta in head_tag.find_all('meta', content=True):
            #     mtext = meta.text.strip()
            #     if len(mtext) <  100 and len(mtext) > 1:
            #         #print("2324,,,meta content", meta.name,mtext)
            #         page_word_list.append(mtext)

            tmp_title = head_tag.find('title')   
            if tmp_title != None:
                str_title = str(tmp_title)

                # page_word_list.append(str_title)  ##2022.2.26
                
                if ("403 " not in str_title) and ("404 " not in str_title):
                    tmp_list.append(url)
                    tmp_list.append(str_title)

                    list_title.append(tmp_list)
    except:
        print("not found head tag", head_tag)
    #print(content)

    Debug_w(url + " >>>link, rel, href")
    first_page_stylesheets = []
    for frameset in content.find_all('link'):  #rel="stylesheet" 
        tag = frameset.get("rel")
        href= frameset.get("href")
        ctt = frameset.get("content")

        # if tag == None:
        #     continue
        # if href==None:
        #     continue
        if not tag or not href:
            continue
        try:            
            if tag[0].upper()=="STYLESHEET":
                add_list_script(href, url, "css")
                if counter == 1:
                    first_page_stylesheets.append(urljoin(url, href))
                if "http" in href.lower():
                    if baseUrl not in href:
                        cssCount2 = cssCount2 + 1
            global FaviconUrl
            if ("icon" in tag or "ICON" in tag):
                href = href.split(";")[0]
                FaviconUrl = urlForm(href, url, 0)
                Debug_w("FaviconUrl" + FaviconUrl)

                #print("FaviconUrl===",FaviconUrl)
                # if "http" is not FaviconUrl[:4]:
                #     FaviconUrl = http_https + "://" + FaviconUrl
        except:
            print("56-error", frameset, "*****", tag, "*******", href, "*****", type(href))

    if counter == 1:
        for stylesheet_url in first_page_stylesheets[:10]:
            stylesheet_text = fetch_text(stylesheet_url)
            if stylesheet_text:
                tmp_cccc = tmp_cccc | parsing_fontlist(stylesheet_text)
                merge_color_counts(
                    font_color_count,
                    extract_font_colors(stylesheet_text),
                )
          
    scriptCount = scriptCount + len(content.find_all('script'))
    styleCount = styleCount + len(content.find_all('<style'))

    Debug_w("Begin html_adjust step")
    #######################################################
    # div style="background:url('data:image/png;base64,iVBO
    # <img src="data:image/png;base64,, data:image/jpg;base64
    ######################################################
    #print("Check zzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz")
    # if "img src" in html_string or "background:url" in html_string:
    #     html_string = remove_String(html_string,1)
    Debug_w("   Html_adjust step-2-1")
    html_string = style_remove_string(html_string)    
    Debug_w("   Html_adjust step-2-2")
    html_string = enhanced_remove_string(html_string,1)
    Debug_w("End html_adjust step")
    #######################################################
    ## checking the window.external.AddFavorite code
    #    javascript:window.external.AddFavorite('http://www.dc87.co.kr','기프트수림');
    #######################################################
    Debug_w(url + " >>> window.external.AddFavorite") ##  2022.1.15
    paf = html_string.find("window.external.AddFavorite")
    if paf > 0:         
        saf = html_string.find("(", paf)
        eaf = html_string.find(")", saf)
        afLink = html_string[paf+1:saf]
        AddFavoriteCount = AddFavoriteCount + 1
        AddFavoriteSet.add(afLink)
   
    #######################################################
    global Redirection_level
    ######################################################
    # checking alt message od image tag
    ######################################################
    # Debug_w(url + " >>> h1, h2, h3 tag")  ## 20241130, # basic html varidation check 2722
    # global h1Count, h2Count, h3Count
    # #print(heading.name + ' ' + heading.text.strip())
    # ##for h123 in content.findAll(["h1", "h2", "h3"]):
    # for h123 in content.find_all(["h1", "h2", "h3"]):
    #     #h123 = str(h123)
    #     if h123.name=='h1':
    #         h1Count = h1Count + 1
    #     elif h123.name == 'h2' :
    #         h2Count = h2Count + 1
    #     else:
    #         h3Count = h3Count + 1
        ##print("2485", h123.text.strip(), type(h123), h1Count,"******************************************************")
        # page_word_list.append(h123.text.strip())
    ######################################################
    # checking alt message od image tag
    ######################################################
    global imgCount, imgAlt
    Debug_w(url + " >>> img tag")
    ##<img src="data:image/png;base64,iVBORw0K
    # for img in content.find_all('img', alt=True):        
    #     page_word_list.append( img['alt'])
    try:
        for imgTag in content.find_all('img'):
            imgCount = imgCount + 1
            imgTags  = str(imgTag)[:128]
            srcExt = "*"
            if len(imgTag.get('alt','')) > 0:
                imgAlt = imgAlt + 1
            src = imgTag.get('src')
            if src == None:
                src = imgTag.get('srcset')
                if src == None:
                    src = imgTag.get('longdesc')
                    if src == None:
                        src = imgTag.get('data-src')
                        if src == None:
                            src = imgTag.get('data-lazy')                        
                        else:
                            src =" "
            else:
                if src[:10] =="data:image":   #data:image/png;base64,iV
                    srcExt = src[11:14]
                    #print(srcExt)
                else:
                    x = src.rfind(".")
                    if x > 0:
                        srcExt = get_extion(src) ##src[x+1:x+4]
                    else:
                        srcExt = "*"
                if srcExt != "*":
                    extion_count[srcExt] = extion_count.get(srcExt,0) + 1  
                    if srcExt in ["jpg", "png"]:
                        if webPageColorAnalysis == True and counter == 1:
                            #print (counter, src)
                            full_src = urlForm(src, url, 0)
                            full_src = formatHTTP(full_src)
                            list_img_analysis.add(full_src)
                    if srcExt in skip_ext_list:
                        img_list.append(srcExt)
                        if "http" in src.lower():
                            if baseUrl not in src:
                                img_list2.append(srcExt)
    except:
        print(">>>>>>>",src[:10] ,srcExt, imgTags, src[:100], type(src))
    ######################################################
    # checking alt message od amp-image tag
    ######################################################
    for imgTag in content.find_all('amp-img'):
        imgCount = imgCount + 1
        if len(imgTag.get('alt','')) > 0:
            imgAlt = imgAlt + 1

        src = imgTag.get('src')
        if src == None:
            src = imgTag.get('srcset')
            if src == None:
                src = imgTag.get('longdesc')
                if src == None:
                    src = imgTag.get('data-src')
                    if src == None:
                        src = imgTag.get('data-lazy')
                    else:
                        src =" "
        try:
            x = src.rfind(".")
            if x > 0:
                srcExt = get_extion(src) ##src[x+1:x+4]
                if srcExt != "*":
                    extion_count[srcExt] = extion_count.get(srcExt,0) + 1  
                if srcExt in skip_ext_list:
                    img_list.append(srcExt)
                    if "http" in src.lower():
                        if baseUrl not in src:
                            img_list2.append(srcExt)
        except:
            print(imgTag, src[:100], type(src))
    #######################################################
    # checking the background:url('data:imag
    ##<div style="background:url('data:image/png;base64, 2020.7.4
    # $(".goog-te-gadget-icon").css("background", "url(http://i02.i.aliimg.com/wimg/buyer/single/google-translate-logo.png) 0px 0px no-repeat");
    #######################################################    
    if "background:url(" in content:
        x = content.find("data:image")
        if x > 0:
            srcExit = content[x+11:x+14]
            extion_count[srcExt] = extion_count.get(srcExt,0) + 1  
            if srcExt in skip_ext_list:
                img_list.append(srcExt)
                if "http" in src.lower():
                    if baseUrl not in src:
                        img_list2.append(srcExt)
    #######################################################
    # checking a link in the frame 
    #######################################################
    global frameCount, iframeCount
    Debug_w(url + " >>> frame tag")
    tmp_frameset = content.find_all('frame')
    frameCount = frameCount + len(tmp_frameset)

    tmp_frameset = content.find_all('iframe')
    iframeCount = iframeCount + len(tmp_frameset)
    ##print(content)
    #######################################################
    #global baseUrl
    for frameset in content.find_all('frame'):
        tag = frameset.get("src")    
        if tag=="" or tag == None:  ## include or None  2020.2.2
            continue  
        ####if (tag not in baseUrl) and tag[:4]=="http":  2019.11.25
        print(">>>>",frameset, baseUrl, tag)
     
        if (baseUrl not in tag) and tag[:4].lower()=="http":
            Redirection_level += 1   
            print("Rediection level changed:", tag, "   ", baseUrl, "===Level=", Redirection_level)         
            if Redirection_level > 1:
                return
            else:
                baseUrl2 = urlparse(tag)
                baseUrl2 = baseUrl2.netloc 
                print("2640 Base Url is ", baseUrl2)
                #baseUrl = baseUrl2
                url = baseUrl2
                baseUrl = baseUrl2
                # tmpUrl = baseUrl
                # baseUrl = baseUrl2
                # url = baseUrl2
                # baseUrl2 = tmpUrl
        tag = urlForm(tag, url, 0)    
        if tag==None:
            return
  
        addgraphNode(tag,url)
        #scanWeb(tag, url)
        Debug_w("2490> " + url + "    " + tag)
        addScanList(tag,1,url,2566)
    #######################################################
    # checking a http-equiv property of the meta tag 
    #######################################################
    Debug_w(url + " >>> meta tag")
    for meta in content.find_all('meta'):
        if meta not in meta_list:
            meta_list.append(meta)  
        ##<meta content="0; URL=default/" http-equiv="Refresh"/> =
        ##<meta content="0; url=http://www.samwoogroup.co.kr/kr/index.php" http-equiv="refresh">   
        # <meta content="0;url='/kor/intro.do'" http-equiv="refresh"/> 
        #<meta http-equiv="refresh"content="2;url=http://ulsancoding.com/?"/><
        if meta.get('http-equiv') in ["refresh", "Refresh", "REFRESH"]:   
            mc = meta.get('content')
            if mc.upper().find('URL') > 0:
                tag = mc.split("=")[1]
                tag=tag.replace("'","")
                tag=tag.replace('"','')
                tag = urlForm(tag, url, 0)

                try:
                    if 'http' not in tag.lower():
                        if tag[0:1]=="/":
                            ##tag = http_https + ":/" + tag
                            tag = f'{http_https}:/{tag}'
                        else:
                            ##tag = http_https + "://" + tag
                            tag = f'{http_https}://{tag}'
                    addgraphNode(tag,url)
                    Debug_w("2519> " + url + "    " + tag)
                    addScanList(tag,1,url,2595)
                except:
                    print(tag)
                    print("type:" , type(tag) )
                    print("error")
    #######################################################
    # checking a location.href or 
    #   location.replace property of the script area 
    #######################################################  
    #### 944 location in script
    Debug_w(url + " >>> script 1167")
    varList=set()
    for ilink in content.find_all('script'):  ##.find("location.href")
        slink = str(ilink)
        if "application/json" in slink:
            continue
        varList = varList.union(get_varList(slink))        
        #print(varList)
        # slink = slink.replace('\n','').replace('\r','').replace('"','').replace("'","").replace("}"," ").replace("{"," ")
        # slink = slink.replace(" = ","=").replace("= ","=").replace(" =","=")
        # slink = slink.replace("\t"," ").replace(" ,",",").replace(", ",",")
        # llink = slink.split(" ")

        slink = re.sub(r'[\n\r\'"{}]', ' ', slink)
        # ' = ', '= ', ' ='와 같은 패턴을 모두 '='로 치환
        slink = re.sub(r'\s*=\s*', '=', slink)
        # 공백 처리: 여러 개의 공백을 하나로 줄이고, 콤마 주변의 공백 제거
        slink = re.sub(r'\s+', ' ', slink)        # 여러 공백을 하나로
        slink = re.sub(r'\s*,\s*', ',', slink)    # 콤마 앞뒤의 공백 제거
        llink = slink.split()

        for xx in llink:
            if xx[:2]=="//":
                continue
            if check_string(xx, ["window.open", "location"]) == False:
                continue
            skip_flag = 0
            ###$(location).attr('href', "/kor/index.do");  2020.6.25 
            #### window.location.href = "/etjsp/index.jsp"
            ## window.location.replace(sUrl);
            try:  
                cx = xx.find("href")
                if cx > -1:
                    cxs = xx[cx:]
                    if should_skip_href(cxs):
                        print("skip")
                        continue
                   
                if xx.find("location.href") > -1:   ###document.location.href = '/www';
                    tag = get_locationhref(xx,1)   ## <script>location.href="/new/kor/main/main.php"</script>
                    if tag != None :
                        skip_flag = 1
                if xx.find("location=") > -1:
                    tag = get_location(xx)
                    if tag != None:
                        skip_flag = 1
                if xx.find("window.location.replace") > -1:
                    continue
                if xx.find("$(location)") > -1:
                    tag = get_location2(xx)
                    if tag != None:
                        skip_flag = 1                                       
                if xx.find("location.replace") > -1:
                    tag = xx.split("(")[1].split(")")[0]
                    if "+" in tag:    ##2020.2.29
                        skip_flag = 0
                    else:
                        skip_flag = 1
                if xx.find("window.location") > -1 and skip_flag ==0:   ###window.location = "/doc/info01.php";
                    tag = get_windowlocationhref(xx)
                    if tag != None:
                        skip_flag = 1
                    else:
                        continue
                ff = xx.find("window.open")                 
                if ff > -1:  # 	window.open(http://www.upinfor.com/sub3.php,xxx  ##window.open("about:blank","",wopt);
                    ss = xx[ff:]
                    bl = ss.find("(")
                    el = ss.find(")",bl)
                    if bl < el and bl > -1:
                        ss = ss[bl+1:el]
                        if ss.find(",") > -1:
                            tag = ss.split(",")[0]  
                        else:
                            tag = ss
                        skip_flag = 1
                        if "=+" in tag:
                            skip_flag = 0
                        delete_chars = "+-\"'"
                        # maketrans 함수로 삭제할 문자들의 매핑을 생성
                        translation_table = str.maketrans('', '', delete_chars)
                        # translate 함수로 한 번에 문자들을 삭제
                        tag = tag.translate(translation_table)
                        # tag=tag.replace("+","")
                        # tag=tag.replace("-","")    
                        # tag=tag.replace("'","")
                        # tag=tag.replace('"','') 
                  
                        if tag in varList:
                            skip_flag = 0
                #print("skip_flah:", skip_flag)
                if skip_flag == 0:
                    continue 
                if tag==None:
                    continue
                if tag[:2] in ["=+", "=-"]:
                    continue
                if tag[:6] == "mailto":
                    continue                

                if tag in varList:  ### variable
                    continue
                
                if tag.find("this.href") > -1:
                    continue
                # if "을" in tag or ";" in tag or ")" in tag:
                #     print(url, "2675+++++++", tag)
                #     sys.exit()
                xtag = tag
                #print("@@@@@@" ,xx)
                if not_valid(tag) :  ### 20240730, url 등과 같은 함수내 prameter 사용시
                    continue
                tag = urlForm(tag, url, 0)
                if tag==None:       ##2020.2.25
                    continue
                if 'http' not in tag.lower():
                    ##tag = http_https + "://" + tag
                    tag = f'{http_https}://{tag}'
                addgraphNode(tag,url)
                #scanWeb(tag, url)
                Debug_w("2634> " + url + "    " + tag)

                addScanList(tag,1,url,"::"+xtag+"2714")
            except:
                print("944 error:", tag, xx, skip_flag)  ##type=text/javascript>location.href=/html/00_main/main.php;</script>
                continue  
    #######################################################
    # checking the plugin tag  
    #######################################################
    for link in content.find_all('object'):
        list_plugin.append(link)
    for link in content.find_all('embed'):
        list_plugin.append(link)
    #######################################################
    # checking a href property of the anchor tag  
    #######################################################
    global anchorCount, anchorTCount
    Debug_w(url + " >>> anchor 2169")

    for link in content.find_all('a'):  ##<a class="slider" onclick="location.href='/business/field01.asp'">
        anchorCount = anchorCount + 1
        Debug_w(url + " >>> in anchor 2172")
        # atext = link.get_text()
        # if len(atext) > 1:
        #     page_word_list.append(atext)

        if len(link.text)<1:
            anchorTCount = anchorTCount + 1
        tag = None
        if 'href' in link.attrs:
            tag = link.get('href')
            #print()
            #print(url, "*******************",tag)
            # if "을" in tag or ";" in tag or ")" in tag:
            #     print(url, "2713+++++++", tag)
            #     sys.exit()
            tag = urlForm(tag, url, 0)
        else:
            xlink = str(link)
            tag = None
            if "href" in xlink:
                tag = get_href(xlink)
                # if "을" in tag or ";" in tag or ")" in tag:
                #     print(url, "2722+++++++", tag)
                #     sys.exit()
                tag = urlForm(tag, url, 0)

        if tag is None:
            continue
        if tag in ["#","/"]:    ## 2019.7.10
            continue
        if 'http' not in tag.lower():
            ##tag = http_https + "://" + tag    
            tag = f'{http_https}://{tag}'  
        Debug_w(url + " >>> anchor finish 2192")

        addgraphNode(tag,url)
        #scanWeb(tag, url)
        Debug_w("2678> " + url + "    " + tag)

        addScanList(tag,1,url,2765)
    #######################################################
    # checking a href property of the area tag  
    #######################################################
    for link in content.find_all('area'):
        tag = link.get('href')    
        #print(link,tag,url)
        # if "을" in tag or ";" in tag or ")" in tag:
        #     print(url, "2746+++++++", tag)
        #     sys.exit()
        tag = urlForm(tag, url, 0)
        #print(tag)
        #xmessage()
        if tag is None:
            continue
        if tag in ["#","/"]:    ## 2019.7.10
            continue
        if 'http' not in tag.lower():
            ##tag = http_https + "://" + tag    
            tag = f'{http_https}://{tag}'     
        addgraphNode(tag,url)
        #scanWeb(tag, url)
        Debug_w("2697> " + url + "    " + tag)
        addScanList(tag,1,url,2787)
    ##############
    # paragraph tag
    # ####################################
    # for p in content.find_all("p"):
    #     #print("2866", p.text.strip(), type(p), "******************************************************")
    #     page_word_list.append(p.text.strip())

 #######################################################
    ## Special condition..please change or remove next line
    if len(html_string) > 750000:
        html_string = html_string[:750000]
    else:
        Debug_w(url + "\nContnts: " + html_string + " >>>" + str(len(html_string)) +">>>" + url)
    #print("Check 00000000000000000000000000000000000000000000000000000000")
    ####
    # key word routine
    ####
    # text = content.find_all(text=True)
    # blacklist = ['[document]', 'noscript', 'header', 'html', 'meta', 'head', 'input', 'script']
    # for x in text:
    #     if x.parent.name not in blacklist:
    #         if len(x) > 2:
    #             page_word_list.append(x)

    if keyword_path in url:     ## default "/"
        #texts = html_string     #20241020
        texts = word_html_string
        # texts=""
        # for w in page_word_list:
        #     ##texts=texts + " " + w
        #     texts += '{} '.format(w)
        # texts = texts.replace("\\","")
        # 2021.6.11   bs 활용 테스트 시작
        try:          
        #######################################################
        # keyword collection
        #######################################################
            if online_business in texts:
                online_score = online_score + 10
           ## print("check_string2", len(texts))
            if check_string2(texts, online_word_list)==True:
                online_score = online_score + 1

            counting_esg_word(texts)

            x_text = nltk.word_tokenize(texts)
    
            x_list = nltk.pos_tag(x_text)

            #print("check_string2 finish", len(texts))
            #print("ntlk=", len(x_list))
            sharp_flag = 0
            for lxw in x_list:                       
                if sharp_flag == 1:
                    if(len(lxw[0]) == 6):
                        sharp_flag = 0
                        continue
                if lxw[0]== "#":
                    sharp_flag = 1
                else:
                    sharp_flag = 0

                if ":" in lxw[0]:
                    continue
                
                if ord(lxw[0][0]) == 92:
                    continue

                kw = lxw[0]      
                try: 
                    if forbiddenWordEnable:
                        if nplt_forbiddenword.check_nplt_forbiddenword(kw):
                            forbiddenList.append(kw)
                except:
                    #print(lxw)
                    continue                   
                #print(lxw)
                if check_string(lxw[0], word_filter) == False:   

                    if lxw[1] in ["NN", "NNS", "NNP", "NNPS"] :  ##SL, "JJ",
                        #print(kw)
                        kw = lxw[0]                             
                        if len(kw) < 31 and len(kw) > 1:                 
                            if caseSensitiveMode == False:
                                kw = kw.upper()
                            # if kw[1] in remove_prefix_list1:
                            #     kw = kw[1:]
                            
                            if check_Prefixstring(kw, skip_prefix_list) == True:                                    
                                continue
                            b12 = test_12Byte(kw)
                            if b12 > 0:
                                    #print("online=>", kw[-1])
                                kw = kw.replace("ㆍ", "")
                                if kw[-3:] in remove_suffix_list3:
                                    kw = kw[:-3]
                                if kw[-2:] in remove_suffix_list2:
                                    kw = kw[:-2]
                                if kw[-1:] in remove_suffix_list1:                                        
                                    if check_string(kw[-1], remove_excetion1)==False:
                                        kw = kw[:-1]
                                if kw and kw[-1].isdigit():
                                    kw = remove_right_number(kw)
                                if ( (len(kw) > 1 and b12 == 2) or (len(kw) > 4 and b12 == 1) ) :
                                    #("word-list: ",kw)
                                    if skipWordMode == False:
                                        word_count[kw] = word_count.get(kw,0) + 1  
                                    else:
                                        if check_string(kw.upper(), skip_word_list2)==False:
                                        ##if kw.upper() not in skip_word_list:
                                            #print(kw, len(kw), kw.strip())
                                            word_count[kw] = word_count.get(kw,0) + 1          
        except:
            print("occured a erro on keyword exctract procedure 1017", texts[:50], "===>lxw", lxw)   
            Debug_w(url + " >>>end of keyword collections.")

    #######################################################
    # checking a http link for specially using {"externalLink":"http://www.dmk-korea.com/...".....}  
    #######################################################
    s_list = extract_s_list(html_string)

    for u,w in s_list:
        tag = html_string[u:w]
        # if "을" in tag or ";" in tag or ")" in tag:
        #     print(url, "2793+++++++", tag)
        #     sys.exit()
        tag = urlForm(tag, url, 0)
        if tag == None:
            continue
        addgraphNode(tag,url)
        Debug_w(tag + " >>> special link 2680" + html_string)
        # xx = 0
        # for key, val in visitLinkDict.items():
        #     xx=xx+1
        #     print(xx,key, val)
        # print(url)
        if visitLinkDict.get(url,0) > 0:
            continue

        Debug_w("2735> " + url + "    " + tag)
        addScanList(tag,1,url,2845)

def extract_meta_tag(instr):
    tstr = instr
    tinx = tstr.find("<meta ")
    if tinx == -1:
        return
    cinx = tstr.find(">")

    if cinx > (tinx + 5) :
        metaListset.add(tstr[tinx:cinx+1])
        extract_meta_tag(tstr[cinx+1:])
    else:
        return

def get_meta_content(m):
    mi = m.find("content")
    ms = m[mi+7:].split('"')
    return(ms[1])

def re_make_title(xt):
    ts = set()
    ss = set()
    xx = set()
    duf_flag = False

    for tu,tt in xt:
        xu = tu
        if tu[-2:]=="/#":
            xu = tu[:-2]
        if tu[-1:]=="/":
            xu = tu[:-1]
        xu = simpleUrl2(xu)
        ss.add(tt)
        if xu not in xx:
            xx.add(xu)
            xm = tt[tt.find('>')+1:tt.rfind('<')]  + " , " + xu
            ts.add(xm)
    if len(ss) != len(ts):
        duf_flag = True
    else:
        duf_flag = False
    ts = sorted(ts)
    return ts, duf_flag

def colorChange(G):
    nodeCount = {'0': ['default']}
    for nlist in G.nodes():
        fsc = nlist.count("/")
        if str(fsc) in nodeCount:
            nodeCount[str(fsc)].append(nlist)
        else:
            nodeCount[str(fsc)] = [nlist]
    #print(len(nodeCount))
    return nodeCount

def page_Analisys(plist):
    #pageAnalisys
    maxLength = 0
    minLength = 99999999
    maxTime = 0
    maxT_Length = 0
    minT_Length = 0
    maxL_time = 0
    minTime = 999999
    minL_time = 0
    minT_page = ""
    maxT_page = ""
    maxL_page = ""
    minL_page = ""
    pageAnalisys= []

    for ur in urltimelist:
        page_url = ur[0]
        page_len = ur[1]
        page_tim = ur[2]

        if page_len < 0:
            continue

        if page_len > maxLength:
            maxLength = page_len
            maxL_page = page_url
            maxL_time = page_tim
        if page_len < minLength:
            minLength = page_len
            minL_page = page_url
            minL_time = page_tim
        if page_tim > maxTime:
            maxTime = page_tim
            maxT_page = page_url
            maxT_Length = page_len
        if page_tim < minTime:
            minTime = page_tim
            minT_page = page_url
            minT_Length = page_len

    pageAnalisys.append([maxL_time, maxL_page, maxLength])
    pageAnalisys.append([minL_time, minL_page, minLength])
    pageAnalisys.append([maxTime,   maxT_page, maxT_Length])
    pageAnalisys.append([minTime,   minT_page, minT_Length])
    
    return pageAnalisys

def getdomainInformation(url):
    try:
        w = nplt_whois2.nplt_whois(url)

        if "Domain_Name" in w:
            progress_make(1, "domain name is ", w["Domain_Name"])
        if "Lookup_Source" in w:
            progress_make(1, "lookup source is ", w["Lookup_Source"])
        if "Registrar" in w:
            progress_make(1, "Registrar is ", w["Registrar"].replace("'",""))
            list_domain.append(["Registrar", w["Registrar"].replace("'","")])
        if "Agency" in w:
            progress_make(1, "Agency name is ", w["Agency"])
        if "Creation_Date" in w:
            progress_make(1, "creation date is ", w["Creation_Date"])
            list_domain.append(["creation_date", w["Creation_Date"]])
        ex_date = ""
        if "Expiration_Date" in w:
            progress_make(1, "expiration date is ", w["Expiration_Date"])
            list_domain.append(["expiration_date", w["Expiration_Date"]])
        if "Updated_Date" in w:
            progress_make(1, "updated date is ", w["Updated_Date"])
            list_domain.append(["updated_date", w["Updated_Date"]])
        if w.get("NoDRfexpire") is not None:
            di_days = int(w["NoDRfexpire"])
            if di_days < 365 * 2:  ##86  rogress_make(2, MessageList[5][rptLang],"")
                progress_make(6, MessageList[86][rptLang], str(di_days))
        if w.get("Name_Server"):
            progress_make(1, "name_servers:", "")
            for ns in w["Name_Server"]:
                progress_make(1, "  ", ns)
        if "Lookup_Error" in w:
            progress_make(1, "lookup note is ", w["Lookup_Error"])
        return w
    except (TypeError, ValueError, requests.RequestException) as error:
        progress_make(1, "domain lookup failed: ", str(error))
        return {"Lookup_Error": str(error)}


def make_piChart(recipe, title, savefile):  ##(datas, "The rate of 'alt tag' used in the image.", "p1.png")
    import numpy as np
    fig, ax = plt.subplots(figsize=(6, 3), subplot_kw=dict(aspect="equal"))

    #recipe = ["29.95 % used-tag",
    #      "70.05 % none-used-tag"]

    data = [float(x.split()[0]) for x in recipe]
    ingredients = [x.split()[-1] for x in recipe]

    def func(pct, allvals):
        absolute = int(pct/100.*np.sum(allvals))
        return "{:.1f}%\n({:d} %)".format(pct, absolute)

    wedges, texts, autotexts = ax.pie(data, autopct=lambda pct: func(pct, data),
                                  textprops=dict(color="w"))

    ax.legend(wedges, ingredients,
          title=title,
          loc="best",
          bbox_to_anchor=(1, 0, 0.5, 1))

    plt.setp(autotexts, size=8, weight="bold")
    ax.set_title("The rate of 'alt tag' used in the image.")
    savefilename = str(TEMP_DIR / savefile)
    plt.savefig(savefilename)
    plt.close()
    return savefilename

def improved_progress_make(print_status, tag_number, tag_value):
    tag_name = MessageList[tag_number][rptLang]
    field_name = MessageList[tag_number][3]
    if tag_name[-2:] != "_c" and type(tag_value) == str:
        if "," in tag_value:
            tag_value = tag_value.replace(",","")
    if print_status > 0:
        progress_make(print_status, tag_name, tag_value)

    if len(field_name) > 1 :
        list_field=[]
        list_field.append(field_name)
        list_field.append(tag_value)
        list_record.append(list_field)
        #print(list_field)

def improved_progress_make5(tag_number, v1, n2, v2, n3, v3):
    tag_name = MessageList[tag_number][3]
    list_page.append([tag_name, v1])

    tag_name = MessageList[tag_number][3] + "_" + n2
    list_page.append([tag_name, v2])
    
    tag_name = MessageList[tag_number][3] + "_" + n3
    list_page.append([tag_name, v3])   

def progress_make(print_status, tag_name, tag_value):
    sline = tag_name + str(tag_value)
    mline = []
    if print_status in [1, 4, 5, 7]:
        print(sline)
    elif print_status == 2:
        print("\n" + sline)            
    mline.append(print_status)
    mline.append(sline)
    report_list.append(mline)


def add_favicon_report(favicon_url, favicon_path):
    print(f"Favicon information: {favicon_url},{favicon_path}")
    report_list.append(
        [
            7,
            {
                "label": "Favicon information",
                "url": favicon_url,
                "path": str(favicon_path),
            },
        ]
    )


def resolve_favicon_path(favicon_path):
    path = Path(str(favicon_path).strip().strip('"'))
    candidates = [path]

    if path.name.lower().endswith(".png.png"):
        candidates.insert(0, path.with_name(path.name[:-4]))

    candidates.extend(
        [
            IMAGE_DIR / candidate.name
            for candidate in list(candidates)
            if candidate.parent != IMAGE_DIR
        ]
    )

    for candidate in candidates:
        if candidate.is_file():
            return candidate
    return path


def add_picture_fitted(run, image_path, max_width, max_height=None):
    shape = run.add_picture(str(image_path))
    width_ratio = max_width / shape.width if shape.width > max_width else 1.0
    height_ratio = (
        max_height / shape.height
        if max_height is not None and shape.height > max_height
        else 1.0
    )
    scale = min(width_ratio, height_ratio, 1.0)
    if scale < 1.0:
        shape.width = int(shape.width * scale)
        shape.height = int(shape.height * scale)
    return shape


def progress_make2(print_status, tag_name, tag_value1, tag_value2):
    mstr = str(tag_value1) + " " + str(tag_value2)
    progress_make(print_status, tag_name, mstr)

def progress_make3(print_status, tag_name, tag_value1, tunit, tag_value2):
    if tunit == "sec":
        mstr = format(tag_value1,".6f") + " " + tunit + " " + str(tag_value2)
    else:
        mstr = "{:,}".format(tag_value1) + " " + tunit + " " + str(tag_value2)
    progress_make(print_status, tag_name, mstr)

def progress_make33(print_status, tag_name, tag_value1, tunit1, tag_value2, tunit2, tag_value3):
    if tunit1 == "sec":
        mstr1 = format(tag_value1,".6f") + " " + tunit1 
    else:
        mstr1 = "{:,}".format(tag_value1) + " " + tunit1
    if tunit2 == "sec":
        mstr2 = format(tag_value2,".6f") + " " + tunit2
    else:
        mstr2 = "{:,}".format(tag_value2) + " " + tunit2
    mstr = mstr1 + "(" + mstr2 + ") " + tag_value3 
    
    progress_make(print_status, tag_name, mstr)

def progress_make_table(headers, rows):
    report_list.append([8, {"headers": headers, "rows": rows}])

def get_header_value(headers, header_name):
    for key, value in headers.items():
        if key.lower() == header_name.lower():
            return str(value).strip()
    return ""

def is_yes_option(value):
    return str(value or "").strip().upper() in ("YES", "Y", "1", "TRUE")

def collect_security_header_recommendations(headers):
    recommendations = []
    x_xss = get_header_value(headers, "X-XSS-Protection")
    x_content_type = get_header_value(headers, "X-Content-Type-Options")
    x_frame = get_header_value(headers, "X-Frame-Options")
    csp = get_header_value(headers, "Content-Security-Policy")
    hsts = get_header_value(headers, "Strict-Transport-Security")
    referrer = get_header_value(headers, "Referrer-Policy")
    permissions = get_header_value(headers, "Permissions-Policy")
    cache_control = get_header_value(headers, "Cache-Control")
    expires = get_header_value(headers, "Expires")

    if not x_xss:
        recommendations.append(
            "X-XSS-Protection is missing. Modern browsers rely mainly on Content-Security-Policy; use CSP and, for legacy compatibility, consider 'X-XSS-Protection: 0' or '1; mode=block' according to policy."
        )
    elif x_xss.lower() in ("0", "disabled"):
        recommendations.append(
            "X-XSS-Protection disables legacy browser XSS filtering. Confirm this is intentional and that Content-Security-Policy is configured."
        )

    if x_content_type.lower() != "nosniff":
        recommendations.append(
            "X-Content-Type-Options should be set to 'nosniff' to reduce MIME sniffing risk."
        )

    if not x_frame:
        recommendations.append(
            "X-Frame-Options is missing. Add 'DENY' or 'SAMEORIGIN', or use CSP 'frame-ancestors' to reduce clickjacking risk."
        )
    elif x_frame.upper() not in ("DENY", "SAMEORIGIN") and not x_frame.upper().startswith("ALLOW-FROM"):
        recommendations.append(
            f"X-Frame-Options value '{x_frame}' is not a common safe value. Review DENY/SAMEORIGIN or CSP frame-ancestors."
        )

    if not csp:
        recommendations.append(
            "Content-Security-Policy is missing. Add a CSP, at minimum with default-src and frame-ancestors directives."
        )
    elif "frame-ancestors" not in csp.lower():
        recommendations.append(
            "Content-Security-Policy exists but has no frame-ancestors directive. Add it for clickjacking protection."
        )

    if not hsts:
        recommendations.append(
            "Strict-Transport-Security is missing. For HTTPS sites, add HSTS after confirming HTTPS is stable."
        )

    if not referrer:
        recommendations.append(
            "Referrer-Policy is missing. Consider 'strict-origin-when-cross-origin' or stricter."
        )

    if not permissions:
        recommendations.append(
            "Permissions-Policy is missing. Restrict unused browser features such as camera, microphone, geolocation, and payment."
        )

    cache_lower = cache_control.lower()
    if cache_control and (
        "no-store" not in cache_lower
        and ("no-cache" in cache_lower or "max-age=0" in cache_lower)
    ):
        recommendations.append(
            "Cache-Control disables freshness but does not prevent storage. For sensitive pages, use 'no-store, no-cache, must-revalidate'."
        )
    if expires and expires.strip() == "0":
        recommendations.append(
            "Expires is set to 0. Prefer explicit Cache-Control directives for predictable browser and proxy behavior."
        )

    return recommendations

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), relationship_id)
    run = docx.oxml.shared.OxmlElement("w:r")
    properties = docx.oxml.shared.OxmlElement("w:rPr")
    color = docx.oxml.shared.OxmlElement("w:color")
    color.set(docx.oxml.shared.qn("w:val"), "0000FF")
    underline = docx.oxml.shared.OxmlElement("w:u")
    underline.set(docx.oxml.shared.qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text_node = docx.oxml.shared.OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink

def report_write(outfile,graphfile):
    doc = docx.Document()

    doc.core_properties.author = "np-solitions"
    doc.core_properties.subject = baseUrl
    doc.core_properties.title = "Web Anaslisis Report"
    doc.core_properties.comments = "Generate by nplt(www.6p-solutions.com)"
    doc.add_heading('6p Website AnalysisReport v1.0', 0)
    para = doc.add_paragraph()
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin
    content_height = section.page_height - section.top_margin - section.bottom_margin
    header = section.header
    parah = header.paragraphs[0]
    parah.text = "www.6p-solutions.com\t" "\t" + str(datetime.now())
    parah.style = doc.styles["Header"]

    bold_type = False

    for xstyle, sline in report_list:
        try:
            if xstyle == 2:
                run = para.add_run("\n"+sline+ "\n")
                font = run.font
                font.size = Pt(12)
                #font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.bold = True
            elif xstyle == 4:
                run = para.add_run(sline + "\n")
                font = run.font
                font.size = Pt(11)
                font.color.rgb = RGBColor(0x00, 0x00, 0xff)
                run.bold = True
            # elif xstyle == 44:   #### hhtp heaser warrig
            #     run = para.add_run(sline)
            #     font = run.font
            #     font.size = Pt(11)
            #     font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            #     run.bold = True
            elif xstyle == 6:
                run = para.add_run("    " + sline + "\n")
                font = run.font
                font.size = Pt(11)
                font.color.rgb = RGBColor(0xc4, 0x40, 0x36)
                #run.bold = True
            elif xstyle == 5:
                # run = para.add_run("\n\nnode Graph" + "\n")
                # run.font.name = 'Arial'
                # run.bold = True
                ##doc.add_picture(graphfile)
                image_path = Path(sline)
                if not image_path.is_absolute():
                    image_path = BASE_DIR / image_path
                if not image_path.is_file():
                    run = para.add_run(f"    [image unavailable: {image_path}]\n")
                    run.font.size = Pt(10)
                    continue
                add_picture_fitted(
                    para.add_run(),
                    image_path,
                    max_width=content_width,
                    max_height=int(content_height * 0.72),
                )
            elif xstyle == 7:  ## Favicon    
                if isinstance(sline, dict):
                    label = sline.get("label", "Favicon information")
                    favicon_url = sline.get("url", "")
                    favicon_path = sline.get("path", "")
                else:
                    label, _, payload = sline.partition(": ")
                    favicon_url, separator, favicon_path = payload.rpartition(",")
                    if not separator:
                        raise ValueError("Invalid favicon report entry")
                run = para.add_run()
                run.add_text("    " + label + ": " + favicon_url)
                font = run.font
                font.size = Pt(11)
                try:
                    add_picture_fitted(
                        run,
                        resolve_favicon_path(favicon_path),
                        max_width=Inches(0.25),
                        max_height=Inches(0.25),
                    )
                except (OSError, ValueError, TypeError, docx.opc.exceptions.PackageNotFoundError) as error:
                    run.add_text(f" [favicon image unavailable: {error}]")
                run.add_text("\n")   
                run = para.add_run("\n")         
            elif xstyle == 8:
                para = doc.add_paragraph()
                headers = sline.get("headers", [])
                rows = sline.get("rows", [])
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for index, header_text in enumerate(headers):
                    table.rows[0].cells[index].text = str(header_text)
                for row in rows:
                    cells = table.add_row().cells
                    for index, value in enumerate(row):
                        text_value = str(value)
                        if text_value.startswith(("http://", "https://")):
                            add_hyperlink(cells[index].paragraphs[0], text_value, text_value)
                        else:
                            cells[index].text = text_value
                para = doc.add_paragraph()
            else:
                #### x code check 할것 2021.2.6 ValueError: All strings must be XML compatible: Unicode or ASCII, no NULL bytes or control characters

                run = para.add_run("   " + sline + "\n")
                run.font.name = 'Arial'
                run.bold = bold_type
        except (OSError, ValueError, TypeError, docx.opc.exceptions.PackageNotFoundError) as error:
            print("report write error....", sline, xstyle, error)
            continue
    # run = para.add_run("\n\nnode Graph" + "\n")
    # run.font.name = 'Arial'
    # run.bold = True
    # ##doc.add_picture(graphfile)
    # run.add_picture(graphfile)

    outfilename = DEFAULT_REPORT_PATH + outfile
    doc.save(outfilename)


def make_word_cloud(dtext, sfile):
    font_path = BASE_DIR / "fonts" / "NanumGothic.ttf"
    wordcloud = WordCloud(font_path=str(font_path),background_color='white',relative_scaling=1,normalize_plurals=False).generate_from_frequencies(dtext)
                                                                                                    
    filename = "wc_" + sfile

    plt.figure(figsize=(6,5)) #이미지 사이즈 지정
    plt.figure(linewidth=1)
    plt.imshow(wordcloud, interpolation='lanczos') #이미지의 부드럽기 정도
    plt.tight_layout(pad=0)
    plt.axis('off') #x y 축 숫자 제거
    filename = DEFAULT_TEMP_PATH + filename
    plt.savefig(filename, facecolor='white', bbox_inches='tight', edgecolor='blue')
    return (filename)

def draw_node_relation(savefilename, fig_title, N):
    min_degree_percentile = 50
    max_nodes = 100
    nodes_to_remove = [node for node in N.nodes if '?' in str(node) or '.jpg' in str(node) or '.png' in str(node) or '.css' in str(node) or '.svg' in str(node)]
    N.remove_nodes_from(nodes_to_remove)
    try:
        degrees = dict(N.degree())
        #significant_nodes = [node for node, degree in degrees.items() if degree >= 0]
        # 최대 노드 수 제한
        # if len(significant_nodes) > max_nodes:
        #     # 연결 정도가 높은 순으로 정렬하여 선택
        #     significant_nodes.sort(key=lambda x: degrees[x], reverse=True)
        #     significant_nodes = significant_nodes[:max_nodes]

        # Create a subgraph with only significant nodes
        #subgraph = N.subgraph(significant_nodes)
        nodeSize = [(degree * 7 + 33) for _, degree in N.degree()]
        #pos = nx.spring_layout(N, k=0.5, seed=4572321)
        
        options = {
            'node_color': 'red',
            'edge_color': 'blue',
            'with_labels': True,
            'font_size': 6,
            'node_size' : nodeSize ,
            'font_family' : 'Malgun Gothic' ##'arial',
        }
        fig, ax = plt.subplots()
        pos = nx.spring_layout(N, k=0.15,iterations=20)  # `subgraph`를 사용합니다.

        nx.draw(N, pos=pos, ax=ax, **options)
        plt.suptitle(fig_title,fontsize=11)
        
        plt.savefig(savefilename)
        plt.close()

        # fig = plt.figure(figsize=(12, 8))
        # print("2222222222222222222222")
        # ax = fig.add_subplot(111)
        # #fig, ax = plt.subplots()
        # print("333333333333333333333")
        # #plt.figure(figsize=(15, 12))
        # print("4444444444444444444444444444")
        # nx.draw(subgraph, pos, ax=ax, **options)
        # fig.suptitle(baseUrl, fontsize=11)
        # print("666666666666666666666")
        # fig.savefig(filename, dpi=300, bbox_inches='tight')
        # plt.close(fig)

    except Exception as e:
        print(f"Error while saving plot: {e}")
        plt.close()
        return None

if __name__ == "__main__":
    ensure_output_directories()
    parser = argparse.ArgumentParser()
    parser.add_argument('-url', help='website address')
    parser.add_argument('-cost', help='processing mode [Yes/No]', default = "Yes") 
    parser.add_argument('-tree', help='tree structure depth mode [Simple/Full]', default = "Simple")     
    parser.add_argument('-wl', help='disply count of word list [number]', default = '50')   
    parser.add_argument('-wp', help='specific path, only check the url that contains specific path [string]', default = '/')       
    parser.add_argument('-sw', help='use inner word_list [Yes/No]', default = 'Yes')   
    parser.add_argument('-cs', help='apply case-sentive function to the word_list [No/Yes]', default = 'No') 
    parser.add_argument('-dm', help='debug mode [Yes/No]', default='No')  
    parser.add_argument('-rl', help='report language [EN/KR]', default='EN')  
    parser.add_argument('-ml', help='count of the list for the lowest linked page [number]', default = '5') 
    parser.add_argument('-ss', help='search specific string in each pages', default ='')
    parser.add_argument('-ca', help='webpage color analysis for the first 3 pages ', default ='Yes')
    parser.add_argument('-db', help='database update option ;Yes/No]', default ="No")
    parser.add_argument('-sl', help='stop line for debug; linr_number]', default ="No")
    parser.add_argument('-ys', help='year list skip [Yes/No]', default ="Yes")
    parser.add_argument('-fw', help='Word Forbidden enable [No/Yes]', default ="No")
    parser.add_argument('-id', help='compamy id if db_connect Yes [0/Number]', default ="0")
    parser.add_argument('-fd', help='down file [pdf/dox/hwpx/....]', default ="No")
    parser.add_argument(
        '-robots',
        help='respect robots.txt rules [Yes/No]',
        default='Yes',
    )
    
    args = parser.parse_args()

    counter = 0
    brokenLink = []
 
    #K = nx.Graph()  2022.1.20

    url = args.url

    if url is None:
        print("usage : nplt -url website_url [-cost Yes/[No]] [-tree Simple/[Full] [-id 0/[digit]\
                                             [-wl 82/[digit]] [-wp specific path] [-sw Yes/[No]] [-cs No/[Yes]] [-ca No/[Yes] \
                                             [-rl KR/[EN]] [-ml 5/[digit]] [-sl [digit]] [-ys No/[Yes] [-fw No/[Yes] \
                                             [-ss String or Data file] [-sl No/[digit] [-db Yes/[No]] [-dm Yes/[No]] \
                                             [-db Yes/[No]] [-fd No/[file extension]]")
        sys.exit()

    if args.cost.upper() in ["NO", "N", "0"]:
        costSavingMode = False        
    else:
        costSavingMode = True
    
    if args.tree.upper() in ["FULL", "F", "0"]:
        fullTreeMode = True        
    else:
        fullTreeMode = False

    if args.sw.upper() in ["NO", "F", "0"]:
        skipWordMode = False        
    else:
        skipWordMode = True

    if args.cs.upper() in ["NO", "F", "0"]:
        caseSensitiveMode = False
    else:
        caseSensitiveMode = True

    if args.ca.upper() in ["NO", "F", "0"]:
        webPageColorAnalysis = False
    else:
        webPageColorAnalysis = True
        
    if args.dm.upper() in ["YES", "Y", "1"]:
        Debug_mode = True
    else:
        Debug_mode = False
    
    if args.ys.upper() in ["YES", "Y", "1"]:
        yearUrlskip = True
    else:
        yearUrlskip = False

    if args.fw.upper() in ["YES", "Y", "1"]:
        forbiddenWordEnable = True
    else:
        forbiddenWordEnable = False

    if args.rl.upper() in ["EN"]:
        rptLang = 2
    else:
        rptLang = 1    
    if is_yes_option(args.db):
        dbConnection = 1
    else:
        dbConnection = 0
    
    if dbConnection == 1:
        dbConnection_id = args.id
        print(f"Database update: enabled (company id: {dbConnection_id})")
        print(
            "Database connection check: OK "
            f"(next id: {get_lastnumber()})"
        )
    else:
        print("Database update: disabled")
            
    if len(args.wp) > 1:
        keyword_path = args.wp
    else:
        keyword_path= "/"

    if len(args.ss) > 1:
        Search_string = args.ss
    else:
        Search_string=""

    try:
        keyWordList = int(args.wl)
    except:
        keyWordList = 50


    if args.fd.upper() =="No":
        File_Download = []
    else:
        File_Download = Extract_extion_fd(args.fd)
    print(File_Download)

    try:
        StopLine = int(args.sl)
    except:
        StopLine = 0

    try:
        MinLinkLimit = int(args.ml)
    except:
        MinLinkLimit = 5

    print("target url:",url)
    print("costSavingMode:", costSavingMode)
    print("fullTreeMode:", fullTreeMode)
    print("keyWordList:", keyWordList)
    print("StopLine:", StopLine)
    print("Use Inner Word_list:", skipWordMode)

    list_Search_string = make_Search_string(Search_string)
    print(">>>>>>>>>>", list_Search_string)
        
    ##visitLinkDict = {url:0}
    input_url = url
    url = adjustUrl(url)
    #baseUrl = getbaseUrl(url)
    baseUrl = urlparse(url)
    baseUrl = baseUrl.netloc 

    print("baseUrl: ", baseUrl)

    try:
        validate_public_url(url)
    except ValueError as error:
        parser.error(str(error))

    respect_robots = args.robots.upper() not in ["NO", "N", "0"]
    configure_robots(url, respect=respect_robots)

    http_https = getHttp_Https(url)
    savefile =  baseUrl + ".png"
    outfile  =  baseUrl + ".docx"
    #######################################################
    # General parameter display
    #######################################################
    myIpAddress = getMyip()
    start_time = time.time()    ### datetime.now()
    ip_address = "."

    # if dbConnection == 1:
    #     npltIndex = getStandardIndex()   #### 2021.7.19
    #     for x in npltIndex:
    #         print(x, npltIndex[0])
    # else:
    #     npltIndex=[0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
    npltIndex=[0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
    #npltIndex = getStandardIndex()   #### 2021.7.19
    Debug_w(f"npltIndex: {npltIndex}")

    head_tableX = getheadInformation(url)
    print(head_tableX)
    jump_url = 0
    if head_tableX != None:
        #head_table = getheadInformation(url)
        head_table = head_tableX.headers       
        print(head_table) 
        #progress_make(2, MessageList[1][rptLang],"")
        improved_progress_make(2, 1,"")
        #progress_make(1, MessageList[2][rptLang], datetime.now())
        improved_progress_make(1, 2, datetime.now())
        hostname = baseUrl
        if ":" in hostname:
            hostname = hostname.split(":")[0]

        ip_address = socket.gethostbyname(hostname)

        #progress_make(1, MessageList[3][rptLang], baseUrl)
        improved_progress_make(1, 3, baseUrl)
        #progress_make(1, MessageList[4][rptLang], ip_address)
        improved_progress_make(1, 4, ip_address)
        
        #head_table Location tag 2021.1.23
        print(type(head_tableX.status_code), head_tableX)
        print("Before Url=", url, baseUrl)
        
        if head_tableX.status_code in [300, 301, 302, 307]:
            if "Location" in head_table:
                jump_url = 1
                xurl = head_table["Location"]
                print("xurl=", xurl)
                if "./" in xurl:
                    xurl = xurl[1:]
                if baseUrl in xurl:
                    url = xurl
                else:
                    if 'http' in xurl:
                        url = xurl
                        baseUrl=xurl
                    else:
                        if baseUrl[-1]!="/" and xurl[0]!="/":
                            url = baseUrl + "/" + xurl
                        else:
                            url = baseUrl + xurl
                print(url)
    if jump_url == 0:
        xUrl = check_iframe_location(url)
        if get_domain(baseUrl) in xUrl:
            url = xUrl
        else:
            baseUrl=xUrl
            url = xUrl
    if url[-1]=="?":
        url = url[:-1]
    print("After  Url=", url, baseUrl)

    ######################################################
    scanWeb(url, url)
    print("FFFF",len(scanWebList))
    #print(scanWebList)
    while 1:
        #print(len(scanWebList), sys.getsizeof(scanWebList))
        if len(scanWebList) < 1:
            break
        x=scanWebList.pop()
        scanWeb(x[0], x[2])
        ######################################################
    end_time = time.time()   ## datetime.now()
    FaviconFile = ""
    try:
        if len(FaviconUrl) > 5:
            FaviconFile = save_Favicon(FaviconUrl)
            print(FaviconUrl,"<<<<<")
            add_favicon_report(FaviconUrl, FaviconFile)
    except:
        print(type(FaviconUrl), "FaviconUrl error", FaviconUrl, FaviconFile)
    
    progress_make(1, "Host ip address(System run): ", myIpAddress)

    ######################################################
    # make a grah for relation of the each contents
    ######################################################
    savefilename = DEFAULT_REPORT_PATH + savefile
    draw_node_relation(savefilename, baseUrl, G)

    #######################################################
    # checking the image color analisys option
    #######################################################
    #print(webPageColorAnalysis)
    #print(len(list_img_analysis))
    color_bar_path = None
    dominant_color_path = None
    font_color_path = None
    try:
        if webPageColorAnalysis == True and len(list_img_analysis) > 0:
            #print("Image for the color analysis")
            file_list = image_copy(list_img_analysis)
            list_img_analysis = file_list
            color_bar_path = make_image_analysis_bar(file_list)
            dominant_color_path = make_dominant_color_chart(file_list)
        font_color_path = make_font_color_chart(font_color_count)
    except (OSError, ValueError, TypeError, cv2.error) as error:
        print(f"Image color analysis failed: {error}")
    #######################################################
    # checking the bbs style page
    #######################################################
    urllist = dict()
    urltree = []
    rownum = 0
    bbs_page = 0
    
    for nlist in G.nodes():
        if nlist is not None:
            x = nlist.split("=")
            # if len(x)>1:
            #     print(x, x[1])
            u = x[0]
            qc = u.find("?")
            if qc > 1:
                u = u[:qc]                
            if u in urllist:
                urllist[u] += 1
            else:
                urllist[u] = 1
                urltree.append(u)

    for x,y in urllist.items():
        rownum += 1
        if y > 1 :
            bbs=" <<< bbs page"
            bbs_page += 1
        else:
            bbs=""
        print(rownum, x, y, bbs)

    print()
    #######################################################
    # create chart for each table
    datas=[]
    if imgCount > 0:
        imgUtilization = float("{0:.2f}".format((imgAlt * 100.0 / imgCount)))
        imgUtilization2 = 100 - imgUtilization
        ##d1 = str(imgUtilization) + " % " + "used-tag"
        d1 = f'{str(imgUtilization)} % used-tag'
        datas.append(d1)
        ##d1 = str(imgUtilization2) + " % " + "none-used-tag"
        d1 = f'{str(imgUtilization2)} % none-used-tag'
        datas.append(d1)
        make_piChart(datas, "The rate of 'alt tag' used in the image.", "p1.png")
    #######################################################
    # Web server header information display
    #######################################################
    #  list_header []
    #  
    progress_make(2, MessageList[5][rptLang],"")
    progress_make(1, "Site address is ", baseUrl)
    list_header.append(["url",baseUrl ])
    if "server" in head_table:
        progress_make(1, "Server software is ", head_table["server"])
        list_header.append(["server", head_table["server"]])
    if "Location" in head_table:
        progress_make(1, "http Redriect to ", head_table["Location"])
        ####list_header.append(["Location", head_table["Location"]])    
    if "Date" in head_table:
        progress_make(1, "Server Date is ", head_table["Date"])
        list_header.append(["server_date", head_table["Date"]])
    if "Last-Modified" in head_table:
        progress_make(1, "Last-Modified date is ", head_table["Last-Modified"]) 
        list_header.append(["last_modified", head_table["Last-Modified"]])
    if "Connection" in head_table:
        progress_make(1, "Connection type is ", head_table["Connection"])
        list_header.append(["connection", head_table["Connection"]])
    if "X-Forwarded-For" in head_table:
        progress_make(1, "X-Forwarded-For ", head_table["X-Forwarded-For"]) 
        list_header.append(["x_forward", head_table["X-Forwarded-For"]])
    if "X-Backend-Server" in head_table:
        progress_make(1, "X-Backend-Server id ", head_table["X-Backend-Server"]) 
        list_header.append(["x_backend", head_table["X-Backend-Server"]])
    if "X-Powered-By" in head_table:
        if "ARR" in head_table["X-Powered-By"]:
            arr_string = "(Application Request Routing)"
        else:
            arr_string = ""
        progress_make(1, "X-Powered-By ", head_table["X-Powered-By"]+ arr_string)    
        list_header.append(["x_powered", head_table["X-Powered-By"]])    
    if "Via" in head_table:
        progress_make(1, "via ", head_table["Via"])
        list_header.append(["via", head_table["Via"]])    
    if "Content-Encoding" in head_table:
        progress_make(1, "Content-Encoding type is ", head_table["Content-Encoding"] +" compress algorithm")
        list_header.append(["content_encoding", head_table["Content-Encoding"]])    
        if head_table["Content-Encoding"] == "default":
            progress_make(1, " no compress", "")    
    if "Content-Type" in head_table:
        progress_make(1, "Content-Type is ", head_table["Content-Type"])
        list_header.append(["content_type", head_table["Content-Type"]]) 
    if "Content-Language" in head_table:
        progress_make(1, "Content-Language is ", head_table["Content-Language"])    
        list_header.append(["content_language", head_table["Content-Language"]]) 
    if "Location" in head_table:
        progress_make(1, "Location is ", head_table["Location"])    
        list_header.append(["location", head_table["Location"]]) 
    if "Transfer-Encoding" in head_table:
        progress_make(1, "Transfer-Encoding is ", head_table["Transfer-Encoding"])  
        list_header.append(["transfer_encoding", head_table["Transfer-Encoding"]]) 
    if "Cache-Control" in head_table:
        progress_make(1, "Cache-Control is ", head_table["Cache-Control"])  
        list_header.append(["cash_control", head_table["Cache-Control"]]) 
    if "expires" in head_table:
        progress_make(1, "expires is ", head_table["expires"]) 
        list_header.append(["expires", head_table["expires"]]) 
    if "P3P" in head_table:
        progress_make(1, MessageList[71][rptLang], "") 
        p3pl = head_table["P3P"].split(" ")
        p3p_result = get_p3p_label(p3pl)
        for key, value in p3p_result.items():
            progress_make(1, "  - " + key + ": ", value) 
        list_header.append(["p3p", head_table["P3P"]]) 
    else:
        progress_make(6, MessageList[72][rptLang], "")         
    if "X-XSS-Protection" in head_table:
        progress_make(1, " X-XSS(Cross-site scripting)-Protection is ", head_table["X-XSS-Protection"]) 
        list_header.append(["xss_protection", head_table["X-XSS-Protection"]]) 
    else:
        progress_make(6, "X-XSS(Cross-site scripting)-Protection is ", "not defined [W]") 
    if "X-Content-Type-Options" in head_table:
        progress_make(1, "X-Content-Type-Options is ", head_table["X-Content-Type-Options"]) 
        list_header.append(["x_contenttype", head_table["X-Content-Type-Options"]]) 
    else:
        progress_make(6, "X-Content-Type-Options is ", "not defined [W]")         
    if "X-Frame-Options" in head_table:
            progress_make(1, " X-Frame-Options is ", head_table["X-Frame-Options"]) 
            list_header.append(["x_frameoption", head_table["X-Frame-Options"]]) 
    else:
        progress_make(6, "X-Frame-Options is ", "not defined [W]")   
        progress_make(6, MessageList[73][rptLang], "")  
    if "ETag" in head_table:
        list_header.append(["Etag", head_table["Etag"]]) 
    security_header_recommendations = collect_security_header_recommendations(head_table)
    if security_header_recommendations:
        progress_make(2, "Security Header Recommendations", "")
        for recommendation in security_header_recommendations:
            progress_make(6, " - ", recommendation)
    #######################################################
    # web site builder tool
    #######################################################
    if web_builder != "0":
        progress_make(2, MessageList[88][rptLang],web_builder)    ##web site builder tool
    #######################################################
    # page performance report
    #######################################################
    progress_make(2, MessageList[6][rptLang],"")    ##Page information
    if costSavingMode == True:
        progress_make(1, "   ", MessageList[46][rptLang])
    improved_progress_make(0, 69, start_time)  
    improved_progress_make(0, 70, end_time)  
    
    improved_progress_make(0, 7, end_time - start_time)  
    if counter > 0:
        improved_progress_make(0, 8, (end_time - start_time) / counter)        

    progress_make3(1, MessageList[7][rptLang], end_time - start_time, "sec", " (include display time)")
    if counter > 0:
        progress_make3(1, MessageList[8][rptLang],  (end_time - start_time) / counter, " sec", "")

    progress_make(1, MessageList[75][rptLang], npltIndex[0]) 
    #progress_make(1,  MessageList[9][rptLang], counter)
    #progress_make(1,  MessageList[10][rptLang], bbs_page)
    improved_progress_make(1,  9, counter)
    progress_make(1, MessageList[76][rptLang], npltIndex[1]) 
    improved_progress_make(1,  10, bbs_page)
    pa = page_Analisys(urltimelist)

    #  list_field=[]
    #     list_field.append(field_name)
    #     list_field.append(tag_value)
    #     list_record.append(list_field)
    # list_field=[]
    # list_field.append("max_cl_url")
    # list_field.append(pa[0][1])
    list_record.append(["max_cl_url", pa[0][1]])
    list_record.append(["max_cl_byte", pa[0][2]])
    list_record.append(["max_cl_time", pa[0][0]])

    list_record.append(["min_cl_url", pa[1][1]])
    list_record.append(["min_cl_byte", pa[1][2]])
    list_record.append(["min_cl_time", pa[1][0]])

    list_record.append(["max_rt_url", pa[2][1]])
    list_record.append(["max_rt_byte", pa[2][2]])
    list_record.append(["max_rt_time", pa[2][0]])

    list_record.append(["min_rt_url", pa[3][1]])
    list_record.append(["min_rt_byte", pa[3][2]])
    list_record.append(["min_rt_time", pa[3][0]])
    
    progress_make(1, MessageList[11][rptLang],"")
    progress_make33(1, MessageList[12][rptLang], pa[0][2], "byte", pa[0][0], "sec", pa[0][1])    
    progress_make33(1, MessageList[13][rptLang], pa[1][2], "byte", pa[1][0], "sec", pa[1][1])
    progress_make33(1, MessageList[14][rptLang], pa[2][0], "sec",  pa[2][2], "byte",pa[2][1])
    progress_make33(1, MessageList[15][rptLang], pa[3][0], "sec",  pa[3][2], "byte",pa[3][1])

    improved_progress_make5(12, pa[0][2], "time", pa[0][0], "url", pa[0][1])
    improved_progress_make5(13, pa[1][2], "time", pa[1][0], "url", pa[1][1])
    improved_progress_make5(14, pa[2][0], "length", pa[2][2], "url", pa[2][1])
    improved_progress_make5(15, pa[3][0], "length", pa[3][2], "url", pa[3][1])

    progress_make(2, MessageList[52][rptLang],"")
    rtime, rlength = getaCompareSite('kr')
    if rtime > 0:
        progress_make33(1, "Naver main page: ",rtime,"sec", rlength, "byte","http://www.naver.com")    
    rtime, rlength = getaCompareSite('us')
    if rtime > 0:
        progress_make33(1, "Google about page: ",rtime,"sec", rlength, "byte","http://about.google.com") 

    ##insert image of a page responsetine index 
    progress_make(5, "./reponsetime_index.JPG","")
    progress_make(1, "","\n")
    #######################################################
    # page link count report
    #######################################################
    progress_make(1, MessageList[47][rptLang],"")
 
    MaxVisitCnt = 0  
    MinVisitCnt = 9999999
    InnerLink = 0
    OuterLink = 0
    
    for key, val in visitLinkDict.items():
        if val > MaxVisitCnt:
            MaxVisitCnt = val
        if val < MinVisitCnt:
            MinVisitCnt = val
        if key.find(baseUrl) == -1:
            OuterLink = OuterLink + 1
        else:
            InnerLink = InnerLink + 1

    TotalLink = InnerLink + OuterLink
    if TotalLink > 0:
        strInner = str("{:,}".format(InnerLink)) + " (" + str(float("{0:.2f}".format((InnerLink * 100.0 / TotalLink)))) + "%)"
        strOuter = str("{:,}".format(OuterLink)) + " (" + str(float("{0:.2f}".format((OuterLink * 100.0 / TotalLink)))) + "%)"
        progress_make(1, MessageList[60][rptLang], strInner)
        progress_make(1, MessageList[61][rptLang], strOuter)
        improved_progress_make(0, 60, InnerLink)
        improved_progress_make(0, 61, OuterLink)
    
    #progress_make(1, MessageList[48][rptLang], str("{:,}".format(MaxVisitCnt + 1)))
    improved_progress_make(1, 48, str("{:,}".format(MaxVisitCnt + 1)))
    
    for key, val in visitLinkDict.items():
        if val == MaxVisitCnt:
            progress_make(1, "    ", key)

    #progress_make(1, MessageList[49][rptLang], str("{:,}".format(MinVisitCnt + 1)))
    improved_progress_make(1, 49, str("{:,}".format(MinVisitCnt + 1)))
    
    minll = 0
    for key, val in visitLinkDict.items():
        if val == MinVisitCnt:
            if minll < MinLinkLimit:
                progress_make(1, "    ", key)
                minll = minll + 1   
    #######################################################
    # Broken link report
    #######################################################
    progress_make(2, MessageList[16][rptLang], " (currently not reponsed)")
    improved_progress_make(0, 16, len(brokenLink))
    progress_make(1,MessageList[77][rptLang], npltIndex[2])
    if len(brokenLink) == 0:
        progress_make(1,MessageList[17][rptLang], "")
    else:
        bl_row = 0
        for bl in brokenLink:
            bl_row += 1
            if bl_row > MinLinkLimit:
                progress_make(1, "more than " , str(bl_row - MinLinkLimit))
                break
            else:
                progress_make(1, str(bl_row) + ") ", bl)
            
    #######################################################
    # Check SEO
    ####################################################### 
    progress_make(2, MessageList[18][rptLang],"")
    mobile_support = 0
    meta_weight = 0
    
    for ml in meta_list:
        extract_meta_tag(str(ml))
    for xx in metaListset:
        yy =  bs(xx,'html.parser')
        if yy.find("meta",{"name":"robots"}) != None:
            progress_make(1, "robots info: ",xx)
            meta_weight = meta_weight + 1
        if yy.find("meta",{"name":"description"}) != None:
            progress_make(1, "description: ", get_meta_content(xx))
            meta_weight = meta_weight + 1
        if yy.find("meta",{"name":"keywords"}) != None:
            progress_make(1, "keywords: ", get_meta_content(xx))   
            meta_weight = meta_weight + 1
        if yy.find("meta",{"name":"HandheldFriendly"}) != None:
            mobile_support = 1
        if yy.find("meta",{"name":"viewport"}) != None:
            mobile_support = 1
    if mobile_support == 1 :
        progress_make(1, MessageList[19][rptLang], "")
        list_record.append(["mobile_tag", "Y"])
    else:
        progress_make(1, MessageList[20][rptLang], "")
        list_record.append(["mobile_tag", "N"])

    if meta_weight < 3:
        progress_make(1, MessageList[21][rptLang],"")
        progress_make(1," "," ")
        progress_make(4, "Note) ", MessageList[22][rptLang]) 
        progress_make(4, "--- ", MessageList[23][rptLang])
    # robots.txt
    progress_make(2, "robots.txt ", MessageList[24][rptLang])
    progress_make(1, robots_status_message(ROBOTS_INFO), "")

    progress_make(2, "Sitemap ", MessageList[24][rptLang])
    if "SITEMAP" in word_count:
        siteMap_flag = True
        progress_make(1, MessageList[26][rptLang], "")
    else:
        sitemap_xml = getsitemapInformation(baseUrl)
        sitemap_html = None if sitemap_xml else getsitemap2Information(baseUrl)
        if sitemap_xml:
            progress_make(1, MessageList[29][rptLang], "")
        elif sitemap_html:
            progress_make(1, MessageList[28][rptLang], "")
        else:
            progress_make(1, MessageList[27][rptLang], "")
    ############################################        
    # AddFavorite function
    ############################################
    if AddFavoriteCount > 0:
        progress_make(1, MessageList[87][rptLang],str("{:,}".format(AddFavoriteCount)))
        for afs in AddFavoriteSet:
            progress_make(1, " ", afs)
    ############################################        
    # Title of each pages
    ############################################
    title_, duf_flag = re_make_title(list_title)
    progress_make(2, MessageList[50][rptLang],"")
    if duf_flag == True:
        progress_make(4, "    ", MessageList[51][rptLang])
    for st in title_:
        progress_make(1, " " , st)
    ############################################
    # Anchor text information
    ############################################
    progress_make(2,"Anchor Text ", MessageList[24][rptLang])        
    #progress_make(1, MessageList[66][rptLang], str("{:,}".format(anchorCount)))
    improved_progress_make(1, 66, str("{:,}".format(anchorCount)))
    #progress_make(1, MessageList[67][rptLang], str("{:,}".format(anchorCount - anchorTCount)))
    improved_progress_make(1, 67, str("{:,}".format(anchorCount - anchorTCount)))

    if anchorCount > 0:
        imgUtilization = str(float("{0:.2f}".format(((anchorCount-anchorTCount) * 100.0 / anchorCount)))) + " %"
        progress_make(1, MessageList[68][rptLang], imgUtilization)
    ############################################
    # Alt attribute of image tag, "{:,}".format(imgCount)
    ############################################
    progress_make(2,"Image ALT Text ", MessageList[24][rptLang])        
    progress_make(1, MessageList[81][rptLang], npltIndex[6])
    #progress_make(1, MessageList[31][rptLang], str("{:,}".format(imgAlt)))
    improved_progress_make(1, 30, str("{:,}".format(imgCount)))
    improved_progress_make(1, 31, str("{:,}".format(imgAlt)))
    if imgCount > 0:
        imgUtilization = str(float("{0:.2f}".format((imgAlt * 100.0 / imgCount)))) + " %"
        progress_make(1, MessageList[32][rptLang], imgUtilization)
        p1_png = str(TEMP_DIR / "p1.png")
        progress_make(5, p1_png, "")
    ############################################
    # image file extion analysis 
    ############################################
    if len(img_list) > 0:
        progress_make(2,MessageList[33][rptLang], MessageList[24][rptLang])   #Image file extension
        imgSet = set()
        imgno = 0
        img_string = ""
        for ext in img_list:
            imgSet.add(ext)
        for ext in imgSet:
            imgno = imgno + 1     
            progress_make(1, " " + str(imgno) + ") "+ ext + ": ", str("{:,}".format(img_list.count(ext)))) 
            img_string = img_string + ext + ":"  +  str(img_list.count(ext)) + " "
        list_record.append(["img_string", img_string])
    
    if len(img_list2) > 0:
        progress_make(2,MessageList[53][rptLang], MessageList[24][rptLang])   #Image file extension2
        imgSet = set()
        imgno = 0
        img_string2 = ""
        for ext in img_list2:
            imgSet.add(ext)
        for ext in imgSet:
            imgno = imgno + 1     
            progress_make(1, " " + str(imgno) + ") "+ ext + ": ", str("{:,}".format(img_list2.count(ext)))) 
            img_string2 = img_string2 + ext + ":"  +  str(img_list2.count(ext)) + " "
        list_record.append(["img_string2", img_string2])
    ############################################
    # H1, H2, H2 tag analysis 
    ############################################
    tags_dict = merge_dicts_as_tuples(start_tag, end_tag)
    h1Count =0
    h2Count = 0
    h3Count = 0

    for k,v in tags_dict.items():  ##counter
        v1, v2 = v
        # print(k,v)
        # if v1 != v2:
        #     print(f"{k} not matched. ({v1},{v2})")
        if k=="<h1>": 
            h1Count = v1
        elif k=="<h2>":
            h2Count = v1
        elif k=="<h3>":
            h3Count = v1

    progress_make(2,"Head line tag ", MessageList[24][rptLang])   
    progress_make(1,"   ",MessageList[34][rptLang])
    progress_make(1,"   H1 tag: " + str(format(h1Count,"9.0f")) , MessageList[35][rptLang])  ##  times used in this site
    progress_make(1, MessageList[82][rptLang],npltIndex[7])    
    progress_make(1,"   H2 tag: " + str(format(h2Count,"9.0f"))  , MessageList[35][rptLang])  
    progress_make(1, MessageList[83][rptLang],npltIndex[8])  
    progress_make(1,"   H3 tag: " + str(format(h3Count,"9.0f"))  , MessageList[35][rptLang])  
    progress_make(1, MessageList[84][rptLang],npltIndex[9])  
    list_record.append(["hl1", h1Count])
    list_record.append(["hl2", h2Count])
    list_record.append(["hl3", h3Count])

    if h1Count < 2 and h2Count <2 :
       progress_make(1, "    ", MessageList[36][rptLang]) ##"No h1, h2, or h3 tags were detected on this site."
       progress_make(1, "    ", MessageList[37][rptLang])
    #########################################################
    # font list
    flcnt = 0    
    #fcount = len(list_font)
    slstr = ""
    progress_make(2, MessageList[59][rptLang],"")
    list_font = select_font(tmp_cccc)
    for fl in list_font:
        slstr = slstr + fl
        flcnt = flcnt + 1
        if flcnt % 5 == 0:            
            progress_make(1, "   ", slstr)
            slstr = ""
        else:
            slstr = slstr + ",  "
    if flcnt % 5 > 0:
        progress_make(1, "   ", slstr[:-3])
    #########################################################
    # html5 list
    flcnt = 0    
    slstr = ""
    progress_make(2, MessageList[90][rptLang],"")
    if len(list_html5_tag) > 0:
        for fl in list_html5_tag:
            slstr = slstr + fl
            flcnt = flcnt + 1
            if flcnt % 5 == 0:            
                progress_make(1, "   ", slstr)
                slstr = ""
            else:
                slstr = slstr + ",  "
        if flcnt % 5 > 0:
            progress_make(1, "   ", slstr[:-3])
    else:
        progress_make(1,  "    ",MessageList[25][rptLang])
    ########################################################
    # flash file information
    progress_make(2, "Flash file ", MessageList[24][rptLang])
    #progress_make(1, MessageList[38][rptLang], str("{:,}".format(flashCount)))
    flashCount = len(list_flash)
    improved_progress_make(1, 38, str("{:,}".format(flashCount)))   
    if flashCount > 0:
        progress_make(4, "    ", MessageList[39][rptLang]) ##"We do not recommend using flash file on your website.")
    else:
        progress_make(1, "    ", MessageList[40][rptLang]) ##"Can not found Flash file in this - Good")
    progress_make(1, MessageList[85][rptLang],npltIndex[10])  
    #######################################################
    #  scriptCount
    #######################################################
    js_tmp=set()
    cs_tmp=set()
    js_cnt = 0
    cs_cnt = 0
    for x,y,z in list_script:
        if z=="js":
            js_cnt = js_cnt + 1
            js_tmp.add(x)
        elif z=="css":
            cs_cnt = cs_cnt + 1
            cs_tmp.add(x)

    ######################################################
    print("script, link list")
    for x in js_tmp:
        print(x)
    for x in cs_tmp:
        print(x)       
    #######################################################  
    progress_make(2, "Script ", MessageList[24][rptLang])
    #progress_make(1, MessageList[41][rptLang], str("{:,}".format(scriptCount)))   ##"count of script tag: "
    improved_progress_make(1, 41, str("{:,}".format(scriptCount)))   ##"count of script tag: "
    
    progress_make(1, MessageList[42][rptLang], str("{:,}".format(cs_cnt)) + "/" + str("{:,}".format(cssCount2)) )  ##"count of stylesheet tag: "
    progress_make(1, MessageList[43][rptLang], str("{:,}".format(cs_cnt)))  ##"count of stylesheet file: "
    progress_make(1, MessageList[54][rptLang], str("{:,}".format(js_cnt)))  ##"count of javascript file: "
    #######################################################
    progress_make(2, "Frame ", MessageList[56][rptLang])
    #progress_make(1, MessageList[57][rptLang], str("{:,}".format(frameCount)))   
    #progress_make(1, MessageList[58][rptLang], str("{:,}".format(iframeCount)))  
    improved_progress_make(1, 57, str("{:,}".format(frameCount)))   
    progress_make(1, MessageList[78][rptLang],npltIndex[3])  
    improved_progress_make(1, 58, str("{:,}".format(iframeCount)))  
    if iframeCount > 0:
        progress_make(4, MessageList[74][rptLang], " ")
    #######################################################
    # Plugin information
    #######################################################
    progress_make(2, MessageList[62][rptLang],"")
    #progress_make(1, MessageList[63][rptLang], str("{:,}".format(len(list_plugin))))   
    improved_progress_make(1, 63, str("{:,}".format(len(list_plugin))))   
    #######################################################
    # SNS Site information
    #######################################################
    progress_make(2, MessageList[64][rptLang],"")
    sns_rows = build_sns_report_rows()
    if sns_rows:
        progress_make_table(
            ["Platform", "Type counts", "Pages", "Unique URLs", "Representative URL"],
            sns_rows,
        )
    else:
        progress_make(1, " - No SNS links were found.", "")
    #######################################################
    progress_make(2, MessageList[55][rptLang], "")
    gr = googleSearch(baseUrl)
    progress_make(1, gr, "")
    if ' ' in gr:
        gr_cnt = gr.split(' ')[1]
    else:
        gr_cnt = 0
    improved_progress_make(0, 55, gr_cnt) 
    progress_make(1, MessageList[79][rptLang],npltIndex[4])  
    #######################################################
    # display for the information map tree
    #######################################################
    xt = {}
    tree_paths, tree_counts, external_domain_counts = build_tree_map_data(
        H.edges(),
        visitLinkDict.keys(),
        baseUrl,
        simple_mode=not fullTreeMode,
    )
    thelist = sorted(tree_paths)

    # for xx in thelist:
    #     print(xx)

    rnode = Node("/")
    xt["/"] = rnode
    root_tree = rnode

    for cset in thelist:
        if cset!="/":
            ppath = get_parent_path(cset)
            try:
                pnode = xt[ppath]
            except:
                pnode = root_tree
            cnode = Node(cset, parent=pnode)
            xt[cset]=cnode
    progress_make(2, "Tree Map","")
    try:
        for pre, fill, node in RenderTree(root_tree):
            if node.name in tree_counts:
                strNodename = node.name + " - " + str(tree_counts[node.name])
            else:
                strNodename = node.name
            progress_make(1, pre, strNodename)
    except:
        progress_make(1,"Tree Structure does not support this version","")    

    progress_make(2, "External Link Domains", "")
    if external_domain_counts:
        for domain, count in external_domain_counts.items():
            progress_make(1, f"{domain}: ", count)
    else:
        progress_make(1, "No external link domains detected.", "")
  
  
    ########################################################
    # Report RSG keyword   20240725
    ########################################################
    if len(esg_count) > 0:
        word_list=[]            
        for word,count in esg_count.items():
            xtr = '{:09d}'.format(count) + word
            #print(xtr, ":",word,":", len(word))
            word_list.append(xtr)
        word_list.sort(reverse=True)

        progress_make(2, "Pages containing ESG keywords", " (" + str(keyWordList) + ")")  
        record_no = 1
        for wl in word_list:
            ws = wl[9:]
            
            wc = int(wl[:9])
            sline = str(record_no) + ")  " + ws + " : " + str(wc) 
            progress_make(1, sline,"")     
            record_no += 1
            if record_no > keyWordList:
                break
  
    ########################################################
    # Report Top n frequency word in this pages
    ########################################################
    if len(word_count) > 0:
        wc_file = make_word_cloud(word_count, savefile)

        word_list=[]            
        for word,count in word_count.items():
            xtr = '{:09d}'.format(count) + word
            #print(xtr, ":",word,":", len(word))
            word_list.append(xtr)
        word_list.sort(reverse=True)

        progress_make(2, MessageList[44][rptLang], " (" + str(keyWordList) + ")")  ##" Frequency word list in website"
        record_no = 1
        for wl in word_list:
            ws = wl[9:]
            
            wc = int(wl[:9])
            sline = str(record_no) + ")  " + ws + " : " + str(wc) 
            progress_make(1, sline,"")     
            record_no += 1
            if record_no > keyWordList:
                break

        progress_make(5, wc_file,"")
    #######################################################
    # Node Graph
    #######################################################
    progress_make(2, "Node Graph","")
    savefilename = DEFAULT_REPORT_PATH + savefile
    progress_make(5, savefilename,"")
    #######################################################
    # Website color analysis
    #######################################################
    if color_bar_path or dominant_color_path or font_color_path:
        progress_make(2, "Website Color Analysis", "")
        progress_make(1, "Analyzed image count: ", len(list_img_analysis))
        if dominant_color_path:
            progress_make(1, "Dominant color palette: ", "")
            progress_make(5, dominant_color_path, "")
        if color_bar_path:
            progress_make(1, "RGB and neutral distribution: ", "")
            progress_make(5, color_bar_path, "")
        if font_color_path:
            progress_make(
                1,
                "Declared font/text colors on the first page: ",
                len(font_color_count),
            )
            progress_make(5, font_color_path, "")
    #######################################################
    # IP address information report
    #######################################################
    data_table = getipInformation(ip_address)
    if data_table:
        progress_make(2, "IP ", MessageList[24][rptLang])   ##"IP information",
        if len(data_table) > 2:
            ip_data = data_table[1:-1].split(",") 
            for tags in ip_data:
                tags = tags.replace('"','')
                tagi = tags.replace(":", ": ")
                try:
                    tag_split = tags.split(":")
                    if len(tag_split[1]) > 1:
                        progress_make(1,tagi,"")    
                    if tag_split[0]=="isp":
                        list_domain.append(["isp", tag_split[1]])
                except:
                    continue
    #######################################################
    # Domain information report from whois
    #######################################################
    progress_make(2, "Domain ", MessageList[24][rptLang])   ###"Domain information"
    getdomainInformation(baseUrl)
  
    #######################################################
    # converting the report data to word file
    #######################################################
    
    report_write(outfile, savefile)
    if dbConnection == 1:
        print("Database update: writing report data...")
        db_result = report_to_db(savefile)
        if db_result == 1:
            print("Database update: completed")
        else:
            print("Database update: failed")
 #####################
 #
 #   https://www.booksr.co.kr
 #
 ###################

    for sx, sc in extion_count.items():
         print(sx, sc)
    
    for x in list_plugin:
        print(x)
    
    for x in list_sns:
        print(x)
    
    print(anchorCount, anchorTCount)
    
    n=0
    
    for x in tmp_cccc:
        n = n + 1
        print(n,x)
    
    n=0
    for x in list_font:
        n = n + 1
        print(n,x)
    
    print(FaviconUrl)
    
    for x in list_function:
        print(x)
    
    for xx in list_record:
        print(xx)
    
    for w,c in esg_count.items():
        print(w,c)
    
    
    if len(Search_string) > 1:
        print("Special Search string is ", Search_string)
        for x in list_search:
            print(x)
    
    if len(list_flash) > 0:
        for x in list_flash:
            print(x)
    
    for x in forbiddenList:
        print(x)
    
    for w,c in esg_count.items():
        print(w,c)
    
    print(list_html5_tag)
    
    # incollect_path_list\\
    print("incollect_path_list")
    for paths in incollect_path_list:
        print(paths)
    
    # html validation check
    print("html validation check list")
    
    tags_dict = merge_dicts_as_tuples(start_tag, end_tag)
    
    for k,v in tags_dict.items():  ##counter
        v1, v2 = v
        print(k,v)
        if v1 != v2:
            print(f"{k} not matched. ({v1},{v2})")
    
    if False and len(list_img_analysis) > 0:
        print("Image for color analysis")
        all_pixels = []
        b=baseUrl.replace(".", "_")
        s_path = IMAGE_DIR / f"{b}_CCAI.jpg"
       
        for ifile in list_img_analysis:
            i_path = Path(ifile)
            if not i_path.is_absolute():
                i_path = IMAGE_DIR / f"{b}_{ifile}"
            try:
                img = cv2.imread(str(i_path))
                if img is not None:
                    img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
                    all_pixels.append(img_rgb.reshape(-1, 3))
                else:
                    print(f"Failed to read image: {ifile}")
            except Exception as e:
                print(f"Error processing {ifile}: {e}")
    
        if all_pixels:
            # 모든 이미지를 하나의 배열로 합침
            all_pixels = np.vstack(all_pixels)
    
            # KMeans를 사용하여 주요 색상 클러스터링
            n_clusters = min(5, len(all_pixels))
            kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
            kmeans.fit(all_pixels)
            
            # 각 클러스터의 중심 색상 추출
            dominant_colors = np.round(kmeans.cluster_centers_).astype(int)
    
            # 시각화
            plt.figure(figsize=(6, 4))
            plt.title(f"Common Colors Across Images {baseUrl}")
    
            for i, color in enumerate(dominant_colors):
                plt.bar(i, 1, color=color / 255.0, edgecolor="black")
    
            plt.xticks(range(n_clusters), [f"Color {i+1}" for i in range(n_clusters)])
            plt.xlabel("Dominant Colors")
            plt.ylabel("Proportion")
            #plt.show()
            plt.savefig(s_path)
            plt.close()
        else:
            print("No valid images to process.")
    elif False:
        print("No images specified for analysis.")
       
